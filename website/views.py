from flask import Blueprint, jsonify, request, flash, redirect, url_for, send_from_directory, Response
from flask_login import login_required, current_user
from . import mysql
from .models import (
    db,
    User,
    Admin,
    Garden,
    Plant,
    PlantTracking,
    GridSpace,
    Feedback,
    LearningPathContent,
    ActivityLog,
    UserPlantUpdateUsage,
    UserSharedConcept,
    UserSubscription,
    AIUsageTracking,
    AIAnalysisUsage,
    SoilAnalysisUsage,
    Notification
)
from datetime import datetime, timedelta, timezone
import os
import base64
import requests
import re
import time
import json
import io
import numpy as np
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Simple in-memory cache for AI enrichment to reduce rate and cost
_AI_CACHE = {}
_AI_CACHE_TTL_SECONDS = 60 * 60  # 1 hour

# Weather cache to prevent API spam
_WEATHER_CACHE = {}
_WEATHER_CACHE_TTL_SECONDS = 10 * 60  # 10 minutes

# Activity logging cache
_ACTIVITY_LOGS = []
_SUBSCRIPTION_LOGS = []
_HISTORY_LOGS = []


def _normalize_tags(raw_tags):
    """Convert incoming tags (list or comma-separated string) into a standardized comma string."""
    if isinstance(raw_tags, list):
        tags_list = [str(tag).strip() for tag in raw_tags if str(tag).strip()]
    elif isinstance(raw_tags, str):
        tags_list = [tag.strip() for tag in raw_tags.split(',') if tag.strip()]
    else:
        tags_list = []
    return ','.join(tags_list)

def _get_or_create_ai_usage(user_id):
    """Ensure an AIAnalysisUsage record exists for the given user (plant analyses)."""
    usage = AIAnalysisUsage.query.filter_by(user_id=user_id).first()
    if not usage:
        usage = AIAnalysisUsage(user_id=user_id, free_analyses_used=0, purchased_credits=0)
        db.session.add(usage)
        db.session.flush()
    return usage

def _get_or_create_soil_usage(user_id):
    """Ensure a SoilAnalysisUsage record exists for the given user (soil analyses)."""
    from website.models import SoilAnalysisUsage
    usage = SoilAnalysisUsage.query.filter_by(user_id=user_id).first()
    if not usage:
        usage = SoilAnalysisUsage(user_id=user_id, free_analyses_used=0, purchased_credits=0)
        db.session.add(usage)
        db.session.flush()
    return usage

def _check_ai_usage_limit(user_id, is_premium=False):
    """Check if user can perform a plant AI analysis. Returns (can_proceed, remaining, needs_payment)"""
    # Premium users get 10 free analyses, basic users get 3
    free_allocation = 10 if is_premium else 3
    
    usage = _get_or_create_ai_usage(user_id)
    db.session.flush()  # Ensure usage record is in session
    free_used = usage.free_analyses_used or 0
    purchased = usage.purchased_credits or 0
    remaining = usage.total_remaining(free_allocation)
    
    print(f"🔍 Plant usage check for user {user_id}: is_premium={is_premium}, free_allocation={free_allocation}, free_used={free_used}, purchased={purchased}, remaining={remaining}")
    
    if remaining > 0:
        return True, remaining, False
    else:
        print(f"⚠️ Plant limit reached for user {user_id}: free_used={free_used} (limit={free_allocation}), purchased={purchased}")
        return False, 0, True

def _check_soil_usage_limit(user_id, is_premium=False):
    """Check if user can perform a soil AI analysis. Returns (can_proceed, remaining, needs_payment)"""
    # Premium users get 10 free analyses, basic users get 3
    free_allocation = 10 if is_premium else 3
    
    usage = _get_or_create_soil_usage(user_id)
    db.session.flush()  # Ensure usage record is in session
    free_used = usage.free_analyses_used or 0
    purchased = usage.purchased_credits or 0
    remaining = usage.total_remaining(free_allocation)
    
    print(f"🌱 Soil usage check for user {user_id}: is_premium={is_premium}, free_allocation={free_allocation}, free_used={free_used}, purchased={purchased}, remaining={remaining}")
    
    if remaining > 0:
        return True, remaining, False
    else:
        print(f"⚠️ Soil limit reached for user {user_id}: free_used={free_used} (limit={free_allocation}), purchased={purchased}")
        return False, 0, True

def _track_ai_usage(user_id, usage_type, is_free=True, cost=0.00, image_path=None, analysis_result=None):
    """Track a plant AI analysis usage"""
    # Update usage counter
    usage = _get_or_create_ai_usage(user_id)
    if is_free:
        usage.free_analyses_used = (usage.free_analyses_used or 0) + 1
    else:
        # Deduct from purchased credits
        if usage.purchased_credits > 0:
            usage.purchased_credits = usage.purchased_credits - 1
    
    # Create tracking record
    tracking = AIUsageTracking(
        user_id=user_id,
        usage_type=usage_type,
        image_path=image_path,
        analysis_result=analysis_result,
        cost=float(cost),
        is_free_usage=is_free
    )
    db.session.add(tracking)
    db.session.flush()
    
    return tracking

def _track_soil_usage(user_id, is_free=True, cost=0.00, image_path=None, analysis_result=None):
    """Track a soil AI analysis usage"""
    # Update usage counter
    usage = _get_or_create_soil_usage(user_id)
    if is_free:
        usage.free_analyses_used = (usage.free_analyses_used or 0) + 1
    else:
        # Deduct from purchased credits
        if usage.purchased_credits > 0:
            usage.purchased_credits = usage.purchased_credits - 1
    
    # Create tracking record (using same AIUsageTracking table but with soil_analysis type)
    tracking = AIUsageTracking(
        user_id=user_id,
        usage_type='soil_analysis',
        image_path=image_path,
        analysis_result=analysis_result,
        cost=float(cost),
        is_free_usage=is_free
    )
    db.session.add(tracking)
    db.session.flush()
    
    return tracking

def _get_or_create_update_usage(user_id):
    """Ensure a UserPlantUpdateUsage record exists for the given user."""
    usage = UserPlantUpdateUsage.query.filter_by(user_id=user_id).first()
    if not usage:
        usage = UserPlantUpdateUsage(user_id=user_id)
        db.session.add(usage)
        # Flush so the object has an ID but defer commit to caller
        db.session.flush()
    return usage

def log_activity(user_id, user_name, action, description, ip_address=None, user_agent=None, status='success'):
    """Log user activity"""
    try:
        from datetime import datetime
        activity = {
            'id': len(_ACTIVITY_LOGS) + 1,
            'user_id': user_id,
            'user_name': user_name,
            'action': action,
            'description': description,
            'ip_address': ip_address or request.remote_addr,
            'user_agent': user_agent or request.headers.get('User-Agent', ''),
            'timestamp': datetime.now().isoformat(),
            'status': status
        }
        _ACTIVITY_LOGS.append(activity)
        print(f"📝 ACTIVITY LOG: {user_name} ({user_id}) - {action}: {description}")
    except Exception as e:
        print(f"Error logging activity: {str(e)}")

def log_subscription_activity(user_id, user_name, action, plan_name, amount, currency, payment_method, status, subscription_id=None):
    """Log subscription-related activity"""
    try:
        from datetime import datetime
        subscription_log = {
            'id': len(_SUBSCRIPTION_LOGS) + 1,
            'user_id': user_id,
            'user_name': user_name,
            'action': action,
            'plan_name': plan_name,
            'amount': amount,
            'currency': currency,
            'payment_method': payment_method,
            'status': status,
            'start_date': datetime.now().isoformat(),
            'end_date': (datetime.now() + timedelta(days=30)).isoformat() if status == 'active' else None,
            'timestamp': datetime.now().isoformat(),
            'subscription_id': subscription_id
        }
        _SUBSCRIPTION_LOGS.append(subscription_log)
        print(f"💳 SUBSCRIPTION LOG: {user_name} ({user_id}) - {action}: {plan_name} - {currency} {amount}")
    except Exception as e:
        print(f"Error logging subscription activity: {str(e)}")

def log_history_change(table_name, record_id, action, old_values, new_values, changed_by):
    """Log database changes"""
    try:
        from datetime import datetime
        # Determine which fields changed
        field_changes = []
        if old_values and new_values:
            for key in new_values:
                if key in old_values and old_values[key] != new_values[key]:
                    field_changes.append(key)
        elif new_values:
            field_changes = list(new_values.keys())
        elif old_values:
            field_changes = list(old_values.keys())
        
        history_log = {
            'id': len(_HISTORY_LOGS) + 1,
            'table_name': table_name,
            'record_id': record_id,
            'action': action,
            'old_values': old_values,
            'new_values': new_values,
            'changed_by': changed_by,
            'timestamp': datetime.now().isoformat(),
            'field_changes': field_changes
        }
        _HISTORY_LOGS.append(history_log)
        print(f"📋 HISTORY LOG: {table_name} - {action} - Record {record_id} by {changed_by}")
    except Exception as e:
        print(f"Error logging history change: {str(e)}")

def create_grid_spaces_for_garden(garden_id, grid_size):
    """Create grid spaces for a garden based on its grid size"""
    try:
        # Parse grid size (e.g., "3x3" -> rows=3, cols=3)
        rows, cols = map(int, grid_size.split('x'))
        
        # Create grid spaces
        for row in range(1, rows + 1):
            for col in range(1, cols + 1):
                grid_space = GridSpace(
                    garden_id=garden_id,
                    grid_position=f"{row},{col}",
                    plant_id=None,
                    planting_date=None,
                    notes='',
                    is_active=True
                )
                db.session.add(grid_space)
        
        db.session.flush()
    except Exception as e:
        print(f"Error creating grid spaces: {e}")
        raise e

def expand_garden_grid_to_premium(garden_id):
    """Expand a garden from 3x3 to 6x6 when user subscribes, preserving existing plants"""
    try:
        from website.models import GridSpace
        
        # Get existing grid spaces
        existing_spaces = GridSpace.query.filter_by(garden_id=garden_id, is_active=True).all()
        
        # Create a map of existing positions
        existing_positions = {}
        for space in existing_spaces:
            existing_positions[space.grid_position] = space
        
        # Create all 6x6 grid spaces (36 total)
        new_spaces_count = 0
        for row in range(1, 7):  # 6 rows
            for col in range(1, 7):  # 6 cols
                position = f"{row},{col}"
                
                # If this position doesn't exist, create it
                if position not in existing_positions:
                    grid_space = GridSpace(
                        garden_id=garden_id,
                        grid_position=position,
                        plant_id=None,
                        planting_date=None,
                        notes='',
                        is_active=True
                    )
                    db.session.add(grid_space)
                    new_spaces_count += 1
        
        db.session.flush()
        print(f"✅ Expanded garden {garden_id} from 3x3 to 6x6: added {new_spaces_count} new spaces, preserved {len(existing_spaces)} existing spaces")
        return new_spaces_count
        
    except Exception as e:
        print(f"Error expanding garden grid: {e}")
        raise e

def create_additional_grid_spaces(garden_id, additional_spaces, grid_size):
    """Create additional grid spaces for purchased spaces"""
    try:
        # Parse grid size to get columns
        rows, cols = map(int, grid_size.split('x'))
        
        # Get the current highest position in the grid
        existing_spaces = GridSpace.query.filter_by(garden_id=garden_id).all()
        max_row = 0
        max_col = 0
        
        for space in existing_spaces:
            try:
                row, col = map(int, space.grid_position.split(','))
                max_row = max(max_row, row)
                max_col = max(max_col, col)
            except:
                continue
        
        # Start creating additional spaces from the next available position
        current_row = max_row + 1
        current_col = 1
        
        for i in range(additional_spaces):
            # Create grid space at the calculated position
            grid_space = GridSpace(
                garden_id=garden_id,
                grid_position=f"{current_row},{current_col}",
                plant_id=None,
                planting_date=None,
                notes='',
                is_active=True
            )
            db.session.add(grid_space)
            
            # Move to next position
            current_col += 1
            if current_col > cols:
                current_col = 1
                current_row += 1
        
        db.session.flush()
        print(f"✅ Created {additional_spaces} additional grid spaces for garden {garden_id}")
        
    except Exception as e:
        print(f"Error creating additional grid spaces: {e}")
        raise e
_WEATHER_PENDING_REQUESTS = {}  # Track pending requests to prevent duplicates

views = Blueprint('views', __name__)

@views.route('/')
def home():
    from flask_login import current_user
    if current_user.is_authenticated:
        if current_user.is_admin():
            return redirect(url_for('views.admin_dashboard'))
        else:
            return redirect(url_for('views.user_dashboard'))
    return jsonify({"message": "Welcome to eGrowtify API"})

@views.route('/test_db')
def test_db():
    try:
        cur = mysql.connection.cursor()
        cur.execute("SELECT DATABASE()")
        data = cur.fetchone()
        cur.close()
        return jsonify({"database": data[0], "status": "connected"})
    except Exception as e:
        return jsonify({"error": str(e), "status": "error"})

@views.route('/user/dashboard')
@login_required
def user_dashboard():
    if current_user.is_admin():
        return jsonify({"error": "Admin users should use the admin dashboard."}), 403
    
    return jsonify({"message": "User dashboard", "user_id": current_user.id})

@views.route('/admin/dashboard')
@login_required
def admin_dashboard():
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    # Get some basic stats
    total_users = User.query.count()
    active_users = User.query.filter_by(is_active=True).count()
    total_admins = Admin.query.count()
    
    return jsonify({
        "message": "Admin dashboard",
        "stats": {
            "total_users": total_users,
            "active_users": active_users,
            "total_admins": total_admins
        }
    })

@views.route('/admin/stats')
@login_required
def admin_stats():
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    # Get detailed statistics
    users = User.query.all()
    admins = Admin.query.all()
    
    return jsonify({
        "users": [{"id": u.id, "email": u.email, "full_name": u.full_name, "is_active": u.is_active} for u in users],
        "admins": [{"id": a.id, "username": a.username, "email": a.email, "full_name": a.full_name} for a in admins]
    })

@views.route('/api/ai-recognition', methods=['POST'])
@login_required
def ai_plant_recognition():
    if current_user.is_admin():
        return jsonify({"error": "Admins do not have access to the AI recognition feature."}), 403

    try:
        # Check usage limit (premium users get 10 free tries, basic users get 3)
        is_premium = getattr(current_user, 'subscribed', False)
        can_proceed, remaining, needs_payment = _check_ai_usage_limit(current_user.id, is_premium)
        
        if not can_proceed:
            return jsonify({
                "error": "Free analysis limit reached. Please purchase additional analyses or subscribe to Premium.",
                "limit_reached": True,
                "needs_payment": True,
                "remaining": 0,
                "price_per_analysis": 20.00
            }), 402  # 402 Payment Required
        
        api_key = os.getenv('PLANT_ID_API_KEY')
        if not api_key:
            # Return graceful message with 200 so frontend doesn't throw
            return jsonify({"error": "Missing PLANT_ID_API_KEY. Add it to .env and restart backend."}), 200

        file = request.files.get('image')
        if not file:
            return jsonify({"error": "Image file is required (field name: 'image')."}), 200

        # Read and base64 encode image
        image_bytes = file.read()
        image_b64 = base64.b64encode(image_bytes).decode('utf-8')

        payload = {
            "images": [image_b64],
            "modifiers": ["similar_images"],
            "plant_language": "en",
            "plant_details": [
                "common_names",
                "edible_parts",
                "url",
                "wiki_description"
            ]
        }
        headers = {"Content-Type": "application/json", "Api-Key": api_key}
        resp = requests.post("https://api.plant.id/v2/identify", json=payload, headers=headers, timeout=30)

        if resp.status_code != 200:
            return jsonify({"error": f"Plant.id error {resp.status_code}: {resp.text}"}), 200

        data = resp.json()
        suggestions = data.get('suggestions', [])
        if not suggestions:
            return jsonify({"error": "No match found. Try a clearer photo."}), 200

        # Enhanced common name mapping for better user recognition
        def extract_names(s):
            nm = s.get('plant_name') or s.get('name') or ''
            sci = (s.get('plant_details') or {}).get('scientific_name') or ''
            commons = ((s.get('plant_details') or {}).get('common_names') or []) + (s.get('common_names') or [])
            return nm.lower(), sci.lower(), [c.lower() for c in commons]

        # Comprehensive mapping of scientific names to common names (Philippine-friendly)
        common_name_mapping = {
            # Vegetables
            'daucus carota': 'Carrot',
            'raphanus sativus': 'Radish',
            'solanum lycopersicum': 'Tomato',
            'cucumis sativus': 'Cucumber',
            'capsicum annuum': 'Bell Pepper',
            'solanum melongena': 'Eggplant',
            'allium cepa': 'Onion',
            'allium sativum': 'Garlic',
            'solanum tuberosum': 'Potato',
            'lactuca sativa': 'Lettuce',
            'brassica oleracea': 'Cabbage',
            'spinacia oleracea': 'Spinach',
            'beta vulgaris': 'Beetroot',
            'zea mays': 'Corn',
            'phaseolus vulgaris': 'Green Bean',
            'pisum sativum': 'Pea',
            'cucurbita pepo': 'Zucchini',
            'cucurbita maxima': 'Pumpkin',
            'brassica rapa': 'Turnip',
            'apium graveolens': 'Celery',
            'petroselinum crispum': 'Parsley',
            'coriandrum sativum': 'Cilantro',
            'ocimum basilicum': 'Basil',
            'mentha': 'Mint',
            'thymus vulgaris': 'Thyme',
            'rosmarinus officinalis': 'Rosemary',
            'origanum vulgare': 'Oregano',
            'allium schoenoprasum': 'Chives',
            'brinjal': 'Eggplant',
            'aubergine': 'Eggplant',
            'wild radish': 'Carrot',
            'wild carrot': 'Carrot',
            
            # Philippine-specific mappings
            'guelder-rose': 'Grape',
            'viburnum opulus': 'Grape',
            'viburnum': 'Grape',
            'snowball tree': 'Grape',
            'water elder': 'Grape',
            'cramp bark': 'Grape',
            'guelder rose': 'Grape',
            'bittersweet': 'Grape',
            'solanum dulcamara': 'Grape',
            'climbing nightshade': 'Grape',
            'woody nightshade': 'Grape',
            
            # More common plant mappings for Philippines
            'talong': 'Eggplant',
            'mangga': 'Mango',
            'dalandan': 'Orange',
            'suka': 'Orange',
            'dayap': 'Lemon',
            'saging': 'Banana',
            'niyog': 'Coconut',
            'pinya': 'Pineapple',
            'bayabas': 'Guava',
            'papaya': 'Papaya',
            'pakwan': 'Watermelon',
            'ubas': 'Grape',
            'mansanas': 'Apple',
            'peras': 'Pear',
            
            # Additional Philippine fruits and vegetables
            'atis': 'Sugar Apple',
            'langka': 'Jackfruit',
            'jackfruit': 'Jackfruit',
            'durian': 'Durian',
            'lansones': 'Lansones',
            'lanzones': 'Lansones',
            'rambutan': 'Rambutan',
            'chico': 'Chico',
            'sapodilla': 'Chico',
            'kalamansi': 'Calamansi',
            'calamansi': 'Calamansi',
            'citrus microcarpa': 'Calamansi',
            'santol': 'Santol',
            'duhat': 'Java Plum',
            'java plum': 'Java Plum',
            'syzygium cumini': 'Java Plum',
            'annona squamosa': 'Sugar Apple',
            'sugar apple': 'Sugar Apple',
            'sweetsop': 'Sugar Apple',
            'citrus aurantium': 'Orange',
            'sour orange': 'Orange',
            'malus domestica': 'Apple',
            'star apple': 'Apple',
            'chrysophyllum cainito': 'Apple',
            'pyrus': 'Pear',
            'prunus persica': 'Peach',
            'prunus avium': 'Cherry',
            'fragaria': 'Strawberry',
            'garden strawberry': 'Strawberry',
            'fragaria x ananassa': 'Strawberry',
            'rubus': 'Raspberry',
            'vitis vinifera': 'Grape',
            'musa': 'Banana',
            'persea americana': 'Avocado',
            'mangifera indica': 'Mango',
            'citrullus lanatus': 'Watermelon',
            'cucumis melo': 'Cantaloupe',
            'ananas comosus': 'Pineapple',
            
            # Herbs and Spices
            'lavandula': 'Lavender',
            'salvia officinalis': 'Sage',
            'mentha spicata': 'Spearmint',
            'mentha piperita': 'Peppermint',
            'cinnamomum verum': 'Cinnamon',
            'zingiber officinale': 'Ginger',
            'curcuma longa': 'Turmeric',
            'capsicum frutescens': 'Chili Pepper',
            
            # Flowers
            'rosa': 'Rose',
            'tulipa': 'Tulip',
            'lilium': 'Lily',
            'chrysanthemum': 'Chrysanthemum',
            'gerbera': 'Gerbera Daisy',
            'helianthus annuus': 'Sunflower',
            'petunia': 'Petunia',
            'impatiens': 'Impatiens',
            'begonia': 'Begonia',
            'marigold': 'Marigold',
            'zinnia': 'Zinnia',
            
            # Trees
            'quercus': 'Oak',
            'acer': 'Maple',
            'betula': 'Birch',
            'pinus': 'Pine',
            'picea': 'Spruce',
            'abies': 'Fir',
            'cedrus': 'Cedar',
            'juniperus': 'Juniper',
            'magnolia': 'Magnolia',
            'acer palmatum': 'Japanese Maple',
            
            # Succulents
            'aloe vera': 'Aloe Vera',
            'echeveria': 'Echeveria',
            'crassula ovata': 'Jade Plant',
            'sedum': 'Sedum',
            'kalanchoe': 'Kalanchoe',
            'haworthia': 'Haworthia',
            'agave': 'Agave',
            'cactus': 'Cactus',
            'succulent': 'Succulent'
        }

        crop_synonyms = {
            'carrot': {'carrot', 'daucus carota', 'wild carrot', 'wild radish'},
            'tomato': {'tomato', 'solanum lycopersicum'},
            'cucumber': {'cucumber', 'cucumis sativus'},
            'pepper': {'pepper', 'capsicum annuum', 'bell pepper', 'chili'},
            'radish': {'radish', 'raphanus sativus'},
            'onion': {'onion', 'allium cepa'},
            'garlic': {'garlic', 'allium sativum'},
            'potato': {'potato', 'solanum tuberosum'},
            'lettuce': {'lettuce', 'lactuca sativa'},
            'eggplant': {'eggplant', 'brinjal', 'aubergine', 'solanum melongena'}
        }

        def crop_match_score(s):
            name_l, sci_l, commons_l = extract_names(s)
            score = 0
            for synonyms in crop_synonyms.values():
                if any(k in name_l or k in sci_l or any(k in c for c in commons_l) for k in synonyms):
                    score = 1
                    break
            return score

        # Base probability
        def prob(s):
            try:
                return float(s.get('probability', 0.0))
            except Exception:
                return 0.0

        # Function to get the best common name (Philippine-friendly)
        def get_best_common_name(plant_name, scientific_name, common_names_list):
            # First, try to find a mapping for the scientific name
            if scientific_name and scientific_name.lower() in common_name_mapping:
                return common_name_mapping[scientific_name.lower()]
            
            # Then try the plant name
            if plant_name and plant_name.lower() in common_name_mapping:
                return common_name_mapping[plant_name.lower()]
            
            # Check if any common names are in our mapping
            for common in common_names_list:
                if common and common.lower() in common_name_mapping:
                    return common_name_mapping[common.lower()]
            
            # Extract simpler names from compound names in common names list
            simple_names = ['strawberry', 'orange', 'apple', 'pear', 'peach', 'cherry', 
                          'raspberry', 'grape', 'banana', 'avocado', 'mango', 'watermelon', 
                          'cantaloupe', 'pineapple', 'lemon', 'lime', 'tomato', 'pepper',
                          'cucumber', 'carrot', 'radish', 'onion', 'garlic', 'potato',
                          'lettuce', 'cabbage', 'spinach', 'corn', 'bean', 'pea']
            
            # Check common names list for compound names
            for common in common_names_list:
                if common:
                    common_lower = common.lower()
                    for simple_name in simple_names:
                        pattern = r'\b' + re.escape(simple_name) + r'\b'
                        if re.search(pattern, common_lower):
                            return simple_name.title()
            
            # Extract simpler names from compound names (e.g., "garden strawberry" -> "strawberry")
            if plant_name:
                plant_lower = plant_name.lower()
                for simple_name in simple_names:
                    # Check if the plant name contains the simple name (e.g., "garden strawberry" contains "strawberry")
                    if simple_name in plant_lower:
                        # Make sure it's not part of a larger word
                        pattern = r'\b' + re.escape(simple_name) + r'\b'
                        if re.search(pattern, plant_lower):
                            return simple_name.title()
            
            # Special handling for grape-like plants
            all_text = f"{plant_name or ''} {scientific_name or ''} {' '.join(common_names_list)}".lower()
            if any(keyword in all_text for keyword in ['grape', 'vitis', 'guelder', 'viburnum', 'snowball']):
                return 'Grape'
            
            # If no mapping found, prefer the first common name, or the plant name
            if common_names_list:
                # Filter out scientific-sounding names and prefer common ones
                common_filtered = [c for c in common_names_list if not any(char.isdigit() for char in c) and len(c.split()) <= 2]
                if common_filtered:
                    return common_filtered[0]
                return common_names_list[0]
            
            # If plant name looks scientific, try to extract a common name
            if plant_name and len(plant_name.split()) > 1:
                # Try to find a simple word in the name
                words = plant_name.split()
                for word in words:
                    if len(word) > 3 and not any(char.isdigit() for char in word):
                        return word.title()
            
            return plant_name or 'Unknown'

        # Re-rank: boost crop matches slightly so carrots beat radish if close
        ranked = sorted(suggestions, key=lambda s: (prob(s) + 0.07 * crop_match_score(s)), reverse=True)
        top = ranked[0]
        name = top.get('plant_name') or top.get('name')
        scientific_name = (top.get('plant_details') or {}).get('scientific_name') or ''
        probability = float(top.get('probability', 0)) * 100.0
        details = top.get('plant_details', {})
        common_names = details.get('common_names') or []
        wiki = details.get('wiki_description', {})

        # Get the best common name using our mapping
        display_name = get_best_common_name(name, scientific_name, common_names)
        
        # Simplify compound names to common names (e.g., "Garden Strawberry" -> "Strawberry")
        if display_name:
            display_lower = display_name.lower()
            simple_fruit_veg = {
                'strawberry': 'Strawberry', 'orange': 'Orange', 'apple': 'Apple', 
                'pear': 'Pear', 'peach': 'Peach', 'cherry': 'Cherry',
                'raspberry': 'Raspberry', 'grape': 'Grape', 'banana': 'Banana',
                'avocado': 'Avocado', 'mango': 'Mango', 'watermelon': 'Watermelon',
                'cantaloupe': 'Cantaloupe', 'pineapple': 'Pineapple', 'lemon': 'Lemon',
                'lime': 'Lime', 'tomato': 'Tomato', 'pepper': 'Pepper',
                'cucumber': 'Cucumber', 'carrot': 'Carrot', 'radish': 'Radish',
                'onion': 'Onion', 'garlic': 'Garlic', 'potato': 'Potato',
                'lettuce': 'Lettuce', 'cabbage': 'Cabbage', 'spinach': 'Spinach',
                'corn': 'Corn', 'bean': 'Bean', 'pea': 'Pea'
            }
            for simple_name, display_value in simple_fruit_veg.items():
                pattern = r'\b' + re.escape(simple_name) + r'\b'
                if re.search(pattern, display_lower):
                    display_name = display_value
                    break
        
        # Force common names for known problematic cases
        all_text = f"{name or ''} {scientific_name or ''} {' '.join(common_names)}".lower()
        if any(keyword in all_text for keyword in ['guelder', 'viburnum', 'snowball', 'water elder', 'bittersweet', 'solanum dulcamara', 'climbing nightshade', 'woody nightshade']):
            display_name = 'Grape'
        elif any(keyword in all_text for keyword in ['vitis', 'grape']):
            display_name = 'Grape'
        
        # Additional grape detection based on image characteristics
        try:
            # Check for grape-like characteristics in the image
            img = Image.open(io.BytesIO(image_bytes)).convert('RGB').resize((224, 224))
            arr = np.asarray(img, dtype=np.float32)
            
            # Look for purple/red colors typical of grapes
            purple_mask = (arr[:, :, 0] > 100) & (arr[:, :, 1] < 100) & (arr[:, :, 2] > 100)  # Red channel high, green low, blue high
            purple_ratio = float(np.mean(purple_mask))
            
            # Look for round, clustered shapes (typical of grapes)
            gray = np.mean(arr, axis=2)
            edges = np.abs(np.gradient(gray)[0]) + np.abs(np.gradient(gray)[1])
            edge_density = float(np.mean(edges > 30))
            
            # If image has grape-like characteristics and current name is not grape-related, force it to Grape
            if (purple_ratio > 0.1 or edge_density > 0.3) and not any(keyword in all_text for keyword in ['grape', 'vitis', 'ubas']):
                if any(keyword in all_text for keyword in ['bittersweet', 'nightshade', 'solanum', 'guelder', 'viburnum']):
                    display_name = 'Grape'
        except Exception:
            pass

        # Heuristic: detect strong orange coloration suggesting carrot and adjust when radish/beet chosen
        try:
            hsv = Image.open(io.BytesIO(image_bytes)).convert('HSV').resize((224, 224))
            arr = np.asarray(hsv, dtype=np.uint8)
            h = arr[:, :, 0].astype(np.float32) * (360.0 / 255.0)
            s = arr[:, :, 1].astype(np.float32) / 255.0
            v = arr[:, :, 2].astype(np.float32) / 255.0
            orange_mask = (h >= 20) & (h <= 45) & (s >= 0.45) & (v >= 0.25)
            orange_ratio = float(np.mean(orange_mask))
        except Exception:
            orange_ratio = 0.0

        if orange_ratio > 0.12:
            # If the winner is radish/beet but image is strongly orange, prefer carrot
            lowered = f"{name} {' '.join(common_names)}".lower()
            if any(k in lowered for k in ['radish', 'raphanus', 'beta', 'beet']):
                name = 'Daucus carota'
                display_name = 'Carrot'
                # Increase probability a bit to reflect confidence
                suggestions[0]['probability'] = max(prob(top), 0.5)
        # Provide top alternatives to let UI allow manual correction
        alternatives = []
        for s in ranked[:5]:
            n = s.get('plant_name') or s.get('name')
            sci = (s.get('plant_details') or {}).get('scientific_name') or ''
            commons = (s.get('plant_details') or {}).get('common_names') or []
            p = round(float(s.get('probability', 0)) * 100.0, 1)
            if n:
                # Use common name mapping for alternatives too
                alt_display_name = get_best_common_name(n, sci, commons)
                alternatives.append({'name': alt_display_name, 'confidence': p})

        # Simple heuristics from wiki description to provide dynamic care if AI is unavailable
        def infer_care_from_text(text: str):
            if not isinstance(text, str) or not text.strip():
                return {}
            lower = text.lower()
            care = {}
            # Sunlight
            if re.search(r"full\s+sun|direct\s+sun", lower):
                care['sunlight'] = 'Full sun'
            elif re.search(r"partial\s+shade|part\s+shade|filtered\s+light", lower):
                care['sunlight'] = 'Partial shade'
            elif re.search(r"shade\b|indirect\s+light", lower):
                care['sunlight'] = 'Shade or bright indirect light'
            # Watering
            if re.search(r"drought[-\s]?tolerant|tolerates\s+drought", lower):
                care['watering'] = 'Low; drought-tolerant once established'
            elif re.search(r"keep\s+moist|consistently\s+moist|regular\s+watering", lower):
                care['watering'] = 'Moderate; keep soil evenly moist'
            elif re.search(r"waterlogged|avoid\s+overwatering", lower):
                care['watering'] = 'Allow topsoil to dry; avoid overwatering'
            # Soil
            if re.search(r"well[-\s]?drain(ed|ing)", lower):
                care['soil'] = 'Well-draining soil'
            elif re.search(r"rich\s+soil|fertile\s+soil|organic\s+matter", lower):
                care['soil'] = 'Rich, fertile soil with organic matter'
            elif re.search(r"sandy\s+soil", lower):
                care['soil'] = 'Sandy, free-draining soil'
            return care

        wiki_text = (wiki.get('value') if isinstance(wiki, dict) else None) or ''
        inferred = infer_care_from_text(wiki_text)

        # Rule-based enrichment for common categories and crops (works offline)
        def apply_rule_based_enrichment(scientific: str, display: str, wiki_txt: str):
            name = (scientific or display or '').lower()
            text = (wiki_txt or '').lower()
            def has(*keywords):
                return any(k in name or k in text for k in keywords)

            # Library of rules; expand as needed
            if has('daucus carota', 'carrot'):
                return {
                    'growth_stage': 'Cool‑season root crop',
                    'care_recommendations': {
                        'watering': 'Even moisture; 2.5 cm/week; avoid crusting',
                        'sunlight': 'Full sun',
                        'soil': 'Loose, deep, stone‑free; avoid fresh manure'
                    },
                    'common_issues': ['Forked roots (compaction)', 'Carrot fly', 'Bitter taste from stress'],
                    'estimated_yield': '1–3 kg per m² in 70–80 days'
                }
            if has('solanum lycopersicum', 'tomato'):
                return {
                    'health_status': 'Unknown',
                    'growth_stage': 'Vegetative/fruiting (warm-season annual)',
                    'care_recommendations': {
                        'watering': 'Deeply 1–2x/week; keep evenly moist, avoid wet foliage',
                        'sunlight': 'Full sun (6–8+ hours)',
                        'soil': 'Rich, well‑drained soil with compost; pH 6.0–6.8'
                    },
                    'common_issues': ['Blossom end rot', 'Early blight', 'Aphids', 'Split fruit'],
                    'estimated_yield': '3–10 kg per plant in season'
                }
            if has('cucumis sativus', 'cucumber'):
                return {
                    'growth_stage': 'Vining (warm-season annual)',
                    'care_recommendations': {
                        'watering': 'Consistent moisture; 2.5 cm/week minimum',
                        'sunlight': 'Full sun',
                        'soil': 'Loose, fertile, well‑drained; pH 6.0–6.8'
                    },
                    'common_issues': ['Powdery mildew', 'Cucumber beetles', 'Bitter fruit from stress'],
                    'estimated_yield': '3–5 fruits per vine weekly in peak'
                }
            if has('capsicum annuum', 'pepper', 'bell pepper', 'chili'):
                return {
                    'growth_stage': 'Vegetative/flowering (warm-season annual)',
                    'care_recommendations': {
                        'watering': 'Even moisture; allow topsoil to dry slightly',
                        'sunlight': 'Full sun',
                        'soil': 'Fertile, well‑drained; avoid excess nitrogen'
                    },
                    'common_issues': ['Sunscald', 'Aphids', 'Blossom drop in heat'],
                    'estimated_yield': '5–15 peppers per plant'
                }
            if has('lactuca sativa', 'lettuce'):
                return {
                    'growth_stage': 'Cool‑season leafy',
                    'care_recommendations': {
                        'watering': 'Frequent light watering; keep evenly moist',
                        'sunlight': 'Full sun to partial shade in heat',
                        'soil': 'Loose, fertile, high organic matter'
                    },
                    'common_issues': ['Bolting in heat', 'Slugs', 'Aphids'],
                    'estimated_yield': 'Heads/leaves harvested 30–60 days'
                }
            if has('ocimum basilicum', 'basil'):
                return {
                    'growth_stage': 'Warm‑season herb',
                    'care_recommendations': {
                        'watering': 'Moderate; allow top 2–3 cm to dry',
                        'sunlight': 'Full sun to bright light',
                        'soil': 'Well‑draining, moderately fertile'
                    },
                    'common_issues': ['Downy mildew', 'Root rot from overwater'],
                    'estimated_yield': 'Regular cuttings every 1–2 weeks'
                }
            if has('mentha', 'mint'):
                return {
                    'growth_stage': 'Hardy perennial herb',
                    'care_recommendations': {
                        'watering': 'Keep consistently moist',
                        'sunlight': 'Partial shade to full sun',
                        'soil': 'Moist, rich soil; consider container to contain spread'
                    },
                    'common_issues': ['Aggressive spread', 'Rust'],
                    'estimated_yield': 'Frequent harvest through season'
                }
            if has('citrus', 'lemon', 'orange', 'lime'):
                return {
                    'growth_stage': 'Evergreen fruit tree',
                    'care_recommendations': {
                        'watering': 'Deep, infrequent; let top 3–5 cm dry',
                        'sunlight': 'Full sun',
                        'soil': 'Well‑draining, slightly acidic; avoid wet feet'
                    },
                    'common_issues': ['Scale', 'Leaf miner', 'Nutrient chlorosis'],
                    'estimated_yield': 'Varies by cultivar and age'
                }
            if has('malus domestica', 'apple'):
                return {
                    'growth_stage': 'Deciduous fruit tree',
                    'care_recommendations': {
                        'watering': 'Deep weekly for young trees; adjust by rainfall',
                        'sunlight': 'Full sun',
                        'soil': 'Well‑draining loam; mulch to conserve moisture'
                    },
                    'common_issues': ['Scab', 'Codling moth', 'Fire blight'],
                    'estimated_yield': '10–50+ kg per mature tree'
                }
            if has('rosa', 'rose'):
                return {
                    'growth_stage': 'Flowering shrub',
                    'care_recommendations': {
                        'watering': 'Deep soak 1–2x/week; water soil not foliage',
                        'sunlight': 'Full sun (5–6+ hours)',
                        'soil': 'Rich, well‑drained; regular feeding during bloom'
                    },
                    'common_issues': ['Black spot', 'Powdery mildew', 'Aphids'],
                    'estimated_yield': 'Blooms spring–fall depending on variety'
                }
            if has('cactus', 'succulent', 'crassula', 'echeveria', 'aloe'):
                return {
                    'growth_stage': 'Succulent',
                    'care_recommendations': {
                        'watering': 'Sparse; soak then dry fully (2–4 weeks indoors)',
                        'sunlight': 'Bright light; acclimate to full sun',
                        'soil': 'Gritty, fast‑draining cactus mix'
                    },
                    'common_issues': ['Root rot from overwatering', 'Etiolation in low light'],
                    'estimated_yield': 'N/A'
                }
            if has('tulipa', 'tulip'):
                return {
                    'growth_stage': 'Bulb (cool season)',
                    'care_recommendations': {
                        'watering': 'Moderate during growth; dry summer dormancy',
                        'sunlight': 'Full sun',
                        'soil': 'Well‑drained; avoid waterlogged bulbs'
                    },
                    'common_issues': ['Bulb rot', 'Aphids', 'Deer/rabbit browsing'],
                    'estimated_yield': 'Flowers spring; replant/refresh bulbs as needed'
                }
            if has('helianthus annuus', 'sunflower'):
                return {
                    'growth_stage': 'Annual flowering',
                    'care_recommendations': {
                        'watering': 'Deep watering; drought tolerant once established',
                        'sunlight': 'Full sun',
                        'soil': 'Well‑draining; stake tall varieties'
                    },
                    'common_issues': ['Birds/squirrels on seeds', 'Downy mildew'],
                    'estimated_yield': 'Seeds at maturity (variety dependent)'
                }
            return None

        rule_enrichment = apply_rule_based_enrichment(name, display_name, wiki_text)
        # Determine plant type based on scientific name and common names
        def determine_plant_type(scientific_name, common_names, display_name):
            name_lower = (scientific_name or '').lower()
            common_lower = ' '.join(common_names).lower()
            display_lower = (display_name or '').lower()
            all_text = f"{name_lower} {common_lower} {display_lower}"
            
            # Vegetables
            if any(keyword in all_text for keyword in ['carrot', 'tomato', 'cucumber', 'pepper', 'eggplant', 'onion', 'garlic', 'potato', 'lettuce', 'cabbage', 'spinach', 'beet', 'corn', 'bean', 'pea', 'zucchini', 'pumpkin', 'turnip', 'celery', 'parsley', 'cilantro', 'basil', 'mint', 'thyme', 'rosemary', 'oregano', 'chives']):
                return 'vegetable'
            # Fruits
            elif any(keyword in all_text for keyword in ['orange', 'lemon', 'apple', 'pear', 'peach', 'cherry', 'strawberry', 'raspberry', 'grape', 'banana', 'avocado', 'mango', 'watermelon', 'cantaloupe', 'pineapple']):
                return 'fruit'
            # Herbs
            elif any(keyword in all_text for keyword in ['basil', 'mint', 'thyme', 'rosemary', 'oregano', 'sage', 'cilantro', 'parsley', 'chives', 'lavender', 'ginger', 'turmeric', 'chili']):
                return 'herb'
            # Flowers
            elif any(keyword in all_text for keyword in ['rose', 'tulip', 'lily', 'chrysanthemum', 'gerbera', 'sunflower', 'petunia', 'impatiens', 'begonia', 'marigold', 'zinnia']):
                return 'flower'
            # Trees
            elif any(keyword in all_text for keyword in ['tree', 'oak', 'maple', 'pine', 'cedar', 'birch', 'willow', 'elm', 'ash', 'poplar']):
                return 'tree'
            else:
                return 'flower'  # Default fallback
        
        # Generate comprehensive care guide
        def generate_care_guide(rule_enrichment, inferred, display_name, scientific_name):
            care_parts = []
            
            # Add watering info
            watering = (rule_enrichment or {}).get('care_recommendations', {}).get('watering', inferred.get('watering', 'Water as needed; keep soil appropriate for species'))
            care_parts.append(f"💧 Watering: {watering}")
            
            # Add sunlight info
            sunlight = (rule_enrichment or {}).get('care_recommendations', {}).get('sunlight', inferred.get('sunlight', 'Provide suitable sun exposure for species'))
            care_parts.append(f"☀️ Sunlight: {sunlight}")
            
            # Add soil info
            soil = (rule_enrichment or {}).get('care_recommendations', {}).get('soil', inferred.get('soil', 'Well-draining soil is recommended'))
            care_parts.append(f"🌱 Soil: {soil}")
            
            # Add growth stage info
            growth_stage = (rule_enrichment or {}).get('growth_stage', 'Unknown')
            if growth_stage != 'Unknown':
                care_parts.append(f"📈 Growth Stage: {growth_stage}")
            
            # Add common issues if available
            common_issues = (rule_enrichment or {}).get('common_issues', [])
            if common_issues:
                care_parts.append(f"⚠️ Common Issues: {', '.join(common_issues)}")
            
            return '\n\n'.join(care_parts)
        
        plant_type = determine_plant_type(scientific_name, common_names, display_name)
        care_guide = generate_care_guide(rule_enrichment, inferred, display_name, scientific_name)
        
        result = {
            'plant_name': display_name,  # This now uses our common name mapping
            'plant_type': plant_type,  # Auto-determined plant type
            'care_guide': care_guide,  # Generated care guide
            'scientific_name': scientific_name,  # Keep the scientific name for reference
            'original_name': name,  # Keep the original name from API
            'common_names': common_names,
            'confidence': round(probability, 1),
            'health_status': (rule_enrichment or {}).get('health_status', 'Unknown'),
            'care_recommendations': {
                'watering': (rule_enrichment or {}).get('care_recommendations', {}).get('watering', inferred.get('watering', 'Water as needed; keep soil appropriate for species')),
                'sunlight': (rule_enrichment or {}).get('care_recommendations', {}).get('sunlight', inferred.get('sunlight', 'Provide suitable sun exposure for species')),
                'soil': (rule_enrichment or {}).get('care_recommendations', {}).get('soil', inferred.get('soil', 'Well-draining soil is recommended'))
            },
            'common_issues': [],
            'growth_stage': (rule_enrichment or {}).get('growth_stage', 'Unknown'),
            'estimated_yield': (rule_enrichment or {}).get('estimated_yield', 'N/A'),
            'info_url': details.get('url'),
            'ai_enriched': False
        }
        if alternatives:
            result['alternatives'] = alternatives

        # If rule-based provided issues, set them
        if rule_enrichment and isinstance(rule_enrichment.get('common_issues'), list):
            result['common_issues'] = rule_enrichment['common_issues']
        # Mark whether rules contributed
        result['rules_enriched'] = bool(rule_enrichment) or bool(inferred)

        # Enhanced OpenAI integration for more accurate analysis (PREMIUM ONLY)
        # Basic users get rule-based enrichment only, Premium users get AI-enhanced analysis
        is_premium = getattr(current_user, 'subscribed', False)
        
        try:
            openai_key = os.getenv('OPENAI_API_KEY')
            # Only run AI enrichment for Premium users
            if openai_key and is_premium:
                # Check cache first (key by scientific name or display name)
                cache_key = (result.get('scientific_name') or result.get('plant_name') or '').lower().strip()
                cache_hit = False
                # Only check cache if we have a valid key
                if cache_key and len(cache_key) > 0 and cache_key in _AI_CACHE:
                    cached = _AI_CACHE[cache_key]
                    # Only use cache if it's fresh and has valid AI data
                    if time.time() - cached['ts'] < _AI_CACHE_TTL_SECONDS and isinstance(cached.get('data'), dict):
                        ai = cached['data']
                        # Only use cache if it has meaningful AI data
                        if ai.get('health_status') or ai.get('growth_stage') or ai.get('care_recommendations'):
                            result['health_status'] = ai.get('health_status') or result['health_status']
                            result['growth_stage'] = ai.get('growth_stage') or result['growth_stage']
                            if isinstance(ai.get('care_recommendations'), dict):
                                result['care_recommendations'].update({
                                    'watering': ai['care_recommendations'].get('watering') or result['care_recommendations']['watering'],
                                    'sunlight': ai['care_recommendations'].get('sunlight') or result['care_recommendations']['sunlight'],
                                    'soil': ai['care_recommendations'].get('soil') or result['care_recommendations']['soil']
                                })
                            if isinstance(ai.get('common_issues'), list):
                                result['common_issues'] = ai.get('common_issues')
                            if isinstance(ai.get('possible_issues'), list):
                                result['possible_issues'] = ai.get('possible_issues')
                            if ai.get('estimated_yield'):
                                result['estimated_yield'] = ai.get('estimated_yield')
                            result['ai_enriched'] = True
                            cache_hit = True
                            print(f"✅ Using cached AI data for: {cache_key}")
                            # If using cached data, skip OpenAI API call (usage tracking happens at the end)
                        else:
                            print(f"⚠️ Cached data for '{cache_key}' is invalid, will call OpenAI")
                    else:
                        print(f"⚠️ Cache expired for '{cache_key}', will call OpenAI")
                else:
                    if not cache_key:
                        print(f"⚠️ No cache key available (plant_name: {result.get('plant_name')}), will call OpenAI")
                    else:
                        print(f"🔄 No cache found for '{cache_key}', will call OpenAI")
                
                # Only run OpenAI if cache miss
                if not cache_hit:
                    print(f"🔄 Cache miss for '{cache_key}', calling OpenAI API...")
                    # Cache miss - proceed with OpenAI enhancement
                    # Enhanced system prompt for image-based analysis
                    system_prompt = (
                        "You are an expert horticulturist and plant pathologist with 20+ years of experience. "
                        "Analyze the provided plant image and identification data to give accurate, image-specific guidance. "
                        "Look at the actual condition of the plant in the image - examine leaves, stems, fruits, flowers, and overall health. "
                        "IMPORTANT: Always use common, everyday plant names that people recognize (like 'Carrot', 'Eggplant', 'Tomato') "
                        "instead of scientific names or regional names (like 'Daucus carota', 'Brinjal', 'Solanum lycopersicum'). "
                        "Return your analysis as strict JSON with these exact keys: "
                        "health_status (detailed assessment based on what you see in the image - mention specific visual indicators), "
                        "growth_stage (specific stage based on what's visible in the image - flowers, fruits, leaves, etc.), "
                        "care_recommendations (object with watering, sunlight, soil, fertilizing, pruning based on current condition), "
                        "common_issues (array of specific problems you can see or likely issues based on the image), "
                        "possible_issues (array of objects, each with: issue (string - disease/problem name), likelihood (string - 'low', 'medium', 'high'), causes (string or array - what causes this issue), actions (string or array - what to do about it), watch_for (string or array - symptoms to monitor)), "
                        "estimated_yield (realistic expectations based on current plant condition), "
                        "seasonal_notes (seasonal care tips), pest_diseases (common threats and prevention). "
                        "CRITICAL: For possible_issues, carefully examine the image for ANY visual signs of disease, pests, nutrient deficiencies, or stress. "
                        "Include issues you can actually see in the image (discolored leaves, spots, wilting, etc.) or that are highly likely based on visible conditions. "
                        "If the plant appears completely healthy with no issues, return an empty array []. "
                        "Be specific about what you observe in the image and provide practical advice for home gardening. "
                        "Use simple, common plant names that everyone understands."
                    )
                    
                    # Enhanced user payload with image context
                    user_payload = {
                        "plant_name": result.get('plant_name'),
                        "scientific_name": result.get('scientific_name'),
                        "common_names": result.get('common_names'),
                        "confidence": result.get('confidence'),
                        "wiki_description": wiki_text,
                        "info_url": result.get('info_url'),
                        "current_season": "spring",  # Could be dynamic based on date
                        "analysis_context": "Analyze the actual plant condition visible in the image"
                    }
                    
                    try:
                        # Try SDK path first with vision capabilities
                        from openai import OpenAI
                        # Clear any proxy environment variables that might cause conflicts
                        proxy_vars = ['HTTP_PROXY', 'HTTPS_PROXY', 'http_proxy', 'https_proxy', 'ALL_PROXY', 'all_proxy']
                        original_proxy_values = {}
                        for var in proxy_vars:
                            if var in os.environ:
                                original_proxy_values[var] = os.environ[var]
                                del os.environ[var]
                        
                        # Initialize client with only the api_key to avoid proxy conflicts
                        client = OpenAI(api_key=openai_key)
                        
                        # Restore proxy environment variables
                        for var, value in original_proxy_values.items():
                            os.environ[var] = value
                        
                        # Use vision model for image analysis
                        completion = client.chat.completions.create(
                            model="gpt-4o",  # Use vision-capable model
                            response_format={"type": "json_object"},
                            messages=[
                                {"role": "system", "content": system_prompt},
                                {
                                    "role": "user",
                                    "content": [
                                        {
                                            "type": "text",
                                            "text": f"Analyze this plant image and provide detailed care guidance based on what you see: {user_payload}"
                                        },
                                        {
                                            "type": "image_url",
                                            "image_url": {
                                                "url": f"data:image/jpeg;base64,{image_b64}"
                                            }
                                        }
                                    ]
                                }
                            ],
                            temperature=0.3,  # Lower temperature for more consistent results
                            max_tokens=1000
                        )
                        content = completion.choices[0].message.content
                    except Exception as e:
                        print(f"❌ OpenAI SDK failed: {str(e)}")
                        # Check if it's the proxies error specifically
                        if "proxies" in str(e):
                            print("🔧 Fixing proxies error by using direct HTTP call")
                        # Fallback to plain HTTPS call with vision
                        http_headers = {
                            'Authorization': f'Bearer {openai_key}',
                            'Content-Type': 'application/json'
                        }
                        http_payload = {
                            'model': 'gpt-4o',  # Use vision-capable model
                            'response_format': { 'type': 'json_object' },
                            'messages': [
                                { 'role': 'system', 'content': system_prompt },
                                {
                                    'role': 'user',
                                    'content': [
                                        {
                                            'type': 'text',
                                            'text': f'Analyze this plant image and provide detailed care guidance based on what you see: {user_payload}'
                                        },
                                        {
                                            'type': 'image_url',
                                            'image_url': {
                                                'url': f'data:image/jpeg;base64,{image_b64}'
                                            }
                                        }
                                    ]
                                }
                            ],
                            'temperature': 0.3,
                            'max_tokens': 1000
                        }
                        # Retry with simple backoff on 429
                        backoffs = [0, 2, 5]
                        last_err = None
                        for delay in backoffs:
                            if delay:
                                time.sleep(delay)
                            http_resp = requests.post('https://api.openai.com/v1/chat/completions', json=http_payload, headers=http_headers, timeout=30)
                            if http_resp.status_code == 429:
                                last_err = f"429: {http_resp.text[:200]}"
                                continue
                            http_resp.raise_for_status()
                            content = http_resp.json()['choices'][0]['message']['content']
                            break
                        else:
                            raise Exception(last_err or 'OpenAI HTTP call failed')
                    
                    import json as _json
                    ai = _json.loads(content)
                    
                    # Enhanced merge with more comprehensive data
                    if isinstance(ai, dict):
                        # Update health status with more detailed assessment
                        if ai.get('health_status'):
                            result['health_status'] = ai['health_status']
                        
                        # Update growth stage with more specific information
                        if ai.get('growth_stage'):
                            result['growth_stage'] = ai['growth_stage']
                        
                        # Enhanced care recommendations
                        if isinstance(ai.get('care_recommendations'), dict):
                            care_recs = ai['care_recommendations']
                            result['care_recommendations'].update({
                                'watering': care_recs.get('watering') or result['care_recommendations']['watering'],
                                'sunlight': care_recs.get('sunlight') or result['care_recommendations']['sunlight'],
                                'soil': care_recs.get('soil') or result['care_recommendations']['soil'],
                                'fertilizing': care_recs.get('fertilizing', 'Follow general fertilizing schedule for this plant type'),
                                'pruning': care_recs.get('pruning', 'Prune as needed for health and shape')
                            })
                        
                        # Enhanced common issues with solutions
                        if isinstance(ai.get('common_issues'), list):
                            result['common_issues'] = ai['common_issues']
                        
                        # Enhanced possible_issues (diseases/causes) - AI-detected from image
                        if isinstance(ai.get('possible_issues'), list):
                            result['possible_issues'] = ai['possible_issues']
                        
                        # Enhanced yield estimation
                        if ai.get('estimated_yield'):
                            result['estimated_yield'] = ai['estimated_yield']
                        
                        # Add new AI-enhanced fields
                        if ai.get('seasonal_notes'):
                            result['seasonal_notes'] = ai['seasonal_notes']
                        
                        if ai.get('pest_diseases'):
                            result['pest_diseases'] = ai['pest_diseases']
                        
                        result['ai_enriched'] = True
                        print(f"✅ OpenAI enrichment successful for: {cache_key or 'unknown'}")
                        
                        # Save to cache with enhanced data
                        if cache_key:
                            _AI_CACHE[cache_key] = { 'ts': time.time(), 'data': ai }
                    else:
                        print(f"⚠️ OpenAI returned invalid response format")
                        
        except Exception as ai_err:
            # If AI enrichment fails, fall back to defaults but include server log
            print(f"❌ OpenAI enrichment error: {str(ai_err)}")
            print(f"❌ Error type: {type(ai_err).__name__}")
            import traceback
            print(f"❌ Traceback: {traceback.format_exc()}")
            try:
                import logging
                logging.getLogger(__name__).warning("AI enrichment failed: %s", str(ai_err))
            except Exception:
                ...
            # Optional debug surface in response if enabled
            try:
                if os.getenv('DEBUG_AI') == '1':
                    result['ai_error'] = str(ai_err)
            except Exception:
                ...

        # Filter result based on plan type
        # Basic users: Only basic information (plant name, basic care tips)
        # Premium users: All advanced features (health status, pest identification, AI enrichment, etc.)
        if not is_premium:
            # Basic plan - return simplified result
            basic_result = {
                'plant_name': result.get('plant_name'),
                'plant_type': result.get('plant_type'),
                'scientific_name': result.get('scientific_name'),
                'common_names': result.get('common_names'),
                'confidence': result.get('confidence'),
                'care_recommendations': {
                    'watering': result.get('care_recommendations', {}).get('watering', 'Water as needed'),
                    'sunlight': result.get('care_recommendations', {}).get('sunlight', 'Provide suitable sun exposure'),
                    'soil': result.get('care_recommendations', {}).get('soil', 'Well-draining soil')
                },
                'info_url': result.get('info_url'),
                'plan_type': 'basic'
            }
            if result.get('alternatives'):
                basic_result['alternatives'] = result['alternatives']
            result = basic_result
        else:
            # Premium plan - include all advanced features
            result['plan_type'] = 'premium'
        
        # Track usage after successful analysis (both premium and basic users)
        free_allocation = 10 if is_premium else 3
        usage = _get_or_create_ai_usage(current_user.id)
        is_free = usage.can_use_free(free_allocation)
        old_free_used = usage.free_analyses_used or 0
        old_purchased = usage.purchased_credits or 0
        _track_ai_usage(
            current_user.id,
            'plant_analysis',
            is_free=is_free,
            cost=0.00 if is_free else 20.00,
            analysis_result=str(result.get('plant_name', ''))[:500]
        )
        db.session.commit()
        # Refresh to get updated values
        db.session.refresh(usage)
        new_free_used = usage.free_analyses_used or 0
        new_purchased = usage.purchased_credits or 0
        print(f"✅ Tracked usage for user {current_user.id} (premium={is_premium}, allocation={free_allocation}): free {old_free_used}→{new_free_used}, purchased {old_purchased}→{new_purchased}")
        
        return jsonify(result)
    except Exception as e:
        # Fallback to simple on-device heuristic so the feature still works offline
        try:
            file = request.files.get('image')
            if not file:
                return jsonify({"error": f"Analysis failed: {str(e)}"}), 500
            from PIL import Image
            import numpy as np
            image = Image.open(file.stream).convert('RGB').resize((256, 256))
            arr = np.asarray(image, dtype=np.float32)
            red = arr[:, :, 0]
            green = arr[:, :, 1]
            blue = arr[:, :, 2]
            total = np.maximum(red + green + blue, 1.0)
            green_ratio = float(np.mean(green / total))
            red_ratio = float(np.mean(red / total))
            health = 'Healthy' if green_ratio > 0.33 else 'Possible dryness or nutrient deficiency'
            label = 'Leafy Plant' if green_ratio >= 0.37 else 'Fruit/Flower'
            confidence = round(60 + green_ratio * 40, 1)
            return jsonify({
                'plant_name': label,
                'confidence': confidence,
                'health_status': health,
                'care_recommendations': {
                    'watering': 'Keep soil appropriately moist',
                    'sunlight': 'Provide suitable sun exposure',
                    'soil': 'Well-draining soil'
                },
                'common_issues': [],
                'growth_stage': 'Unknown',
                'estimated_yield': 'N/A'
            })
        except Exception:
            return jsonify({"error": f"Analysis failed: {str(e)}"}), 500

@views.route('/api/soil-analysis', methods=['POST'])
@login_required
def soil_analysis():
    if current_user.is_admin():
        return jsonify({"error": "Admins do not have access to the soil analysis feature."}), 403

    try:
        # Check usage limit (premium users get 10 free tries, basic users get 3)
        is_premium = getattr(current_user, 'subscribed', False)
        can_proceed, remaining, needs_payment = _check_soil_usage_limit(current_user.id, is_premium)
        
        if not can_proceed:
            return jsonify({
                "error": "Free analysis limit reached. Please purchase additional analyses or subscribe to Premium.",
                "limit_reached": True,
                "needs_payment": True,
                "remaining": 0,
                "price_per_analysis": 20.00
            }), 402  # 402 Payment Required
        
        openai_key = os.getenv('OPENAI_API_KEY')
        if not openai_key:
            return jsonify({"error": "Missing OPENAI_API_KEY. Add it to .env and restart backend."}), 200

        file = request.files.get('image')
        if not file:
            return jsonify({"error": "Image file is required (field name: 'image')."}), 200

        # Read and base64 encode image for OpenAI Vision
        image_bytes = file.read()
        image_b64 = base64.b64encode(image_bytes).decode('utf-8')

        # Enhanced system prompt for comprehensive soil analysis with plant recommendations
        system_prompt = (
            "You are an expert soil scientist, agronomist, and horticulturist with extensive experience in soil analysis and plant-soil relationships, "
            "SPECIFICALLY for tropical climates and the Philippines. "
            "Analyze the provided soil image and return detailed, accurate soil assessment with specific plant recommendations. "
            "Consider visual indicators of soil health, texture, moisture, color, structure, and potential issues. "
            "Return your analysis as strict JSON with these exact keys: "
            "moisture_level (detailed assessment with visual indicators), soil_type (clear soil classification: Clay Soil, Sandy Soil, Loam Soil, Silt Soil, or Mixed Soil), "
            "texture (detailed texture description with characteristics), organic_matter (assessment of organic content), "
            "drainage (drainage quality assessment), recommendations (array of specific improvement suggestions), "
            "suitable_plants (detailed object with categories: vegetables, fruits, herbs, flowers, trees - each containing specific plant names with brief explanations), "
            "nutrient_indicators (visual signs of nutrient status), compaction_assessment (soil structure analysis), "
            "soil_health_score (1-10 rating with explanation), seasonal_considerations (best planting times for this soil), "
            "soil_amendments (specific materials to add for improvement), water_retention (how well soil holds water), "
            "root_development (how well roots can grow in this soil). "
            "IMPORTANT: For suitable_plants, prioritize and recommend COMMON PHILIPPINE PLANTS that are widely grown in the Philippines. "
            "For vegetables, recommend: Kangkong (water spinach), Talong (eggplant), Kamatis (tomato), Sitaw (string beans), Okra, Ampalaya (bitter melon), Kalabasa (squash), Pechay (Chinese cabbage), Mustasa (mustard greens), Upo (bottle gourd), Patola (sponge gourd), Sili (chili), Sibuyas (onion), Bawang (garlic), Mais (corn), Labanos (radish), Pipino (cucumber). "
            "For fruits, recommend: Mango, Saging (banana), Papaya, Pinya (pineapple), Bayabas (guava), Niyog (coconut), Calamansi, Atis (sugar apple), Lansones, Rambutan, Durian, Langka (jackfruit), Pakwan (watermelon). "
            "For herbs, recommend: Pandan, Tanglad (lemongrass), Luya (ginger), Luyang Dilaw (turmeric), Balanoy (basil), Oregano, Wansoy (cilantro). "
            "For flowers, recommend: Sampaguita (national flower), Gumamela (hibiscus), Santan, Bougainvillea, Kalachuchi (plumeria), Rosal (rose), Yellow Bell, Adelfa (oleander), Marigold. "
            "Be specific, practical, and accurate for home gardening in the Philippines. Focus on plants that would actually thrive in the specific soil conditions visible in the image and are commonly available in the Philippines."
        )

        try:
            # Use OpenAI Vision API for soil analysis
            from openai import OpenAI
            # Clear any proxy environment variables that might cause conflicts
            proxy_vars = ['HTTP_PROXY', 'HTTPS_PROXY', 'http_proxy', 'https_proxy', 'ALL_PROXY', 'all_proxy']
            original_proxy_values = {}
            for var in proxy_vars:
                if var in os.environ:
                    original_proxy_values[var] = os.environ[var]
                    del os.environ[var]
            
            # Initialize client with only the api_key to avoid proxy conflicts
            client = OpenAI(api_key=openai_key)
            
            # Restore proxy environment variables
            for var, value in original_proxy_values.items():
                os.environ[var] = value
            
            completion = client.chat.completions.create(
                model="gpt-4o",  # Use vision-capable model
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": system_prompt},
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Analyze this soil image and provide detailed soil assessment for home gardening:"
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{image_b64}"
                                }
                            }
                        ]
                    }
                ],
                temperature=0.3,
                max_tokens=1000
            )
            content = completion.choices[0].message.content
        except Exception as e:
            print(f"❌ OpenAI SDK failed: {str(e)}")
            # Check if it's the proxies error specifically
            if "proxies" in str(e):
                print("🔧 Fixing proxies error by using direct HTTP call")
            # Fallback to plain HTTPS call
            http_headers = {
                'Authorization': f'Bearer {openai_key}',
                'Content-Type': 'application/json'
            }
            http_payload = {
                'model': 'gpt-4o',
                'response_format': { 'type': 'json_object' },
                'messages': [
                    { 'role': 'system', 'content': system_prompt },
                    {
                        'role': 'user',
                        'content': [
                            {
                                'type': 'text',
                                'text': 'Analyze this soil image and provide detailed soil assessment for home gardening:'
                            },
                            {
                                'type': 'image_url',
                                'image_url': {
                                    'url': f'data:image/jpeg;base64,{image_b64}'
                                }
                            }
                        ]
                    }
                ],
                'temperature': 0.3,
                'max_tokens': 1000
            }
            
            # Retry with simple backoff on 429
            backoffs = [0, 2, 5]
            last_err = None
            for delay in backoffs:
                if delay:
                    time.sleep(delay)
                http_resp = requests.post('https://api.openai.com/v1/chat/completions', json=http_payload, headers=http_headers, timeout=30)
                if http_resp.status_code == 429:
                    last_err = f"429: {http_resp.text[:200]}"
                    continue
                http_resp.raise_for_status()
                content = http_resp.json()['choices'][0]['message']['content']
                break
            else:
                raise Exception(last_err or 'OpenAI HTTP call failed')

        import json as _json
        ai_result = _json.loads(content)
        
        # Structure the response to match frontend expectations with enhanced data
        # Extract soil type from texture or use texture as soil type
        texture_value = ai_result.get('texture', 'Unable to determine from image')
        soil_type_value = ai_result.get('soil_type', texture_value)
        
        # Normalize soil type into a clear class for recommendations
        soil_type_str = (soil_type_value or '').lower()
        if 'clay' in soil_type_str:
            soil_type_class = 'Clay Soil'
        elif 'loam' in soil_type_str or 'loamy' in soil_type_str:
            soil_type_class = 'Loam Soil'
        elif 'sand' in soil_type_str or 'sandy' in soil_type_str:
            soil_type_class = 'Sandy Soil'
        elif 'silt' in soil_type_str:
            soil_type_class = 'Silt Soil'
        else:
            soil_type_class = 'Mixed Soil'

        # Read user gardening skill level (beginner / experienced)
        user_level = getattr(current_user, 'learning_level', None)
        if user_level not in ('beginner', 'experienced'):
            user_level = None

        # Curated, soil-type + skill-level based plant recommendations (English names)
        soil_based_plants = {
            'Clay Soil': {
                'beginner': {
                    'vegetables': ['Cabbage', 'Kale', 'Spinach'],
                    'fruits': ['Guava', 'Banana'],
                    'herbs': ['Mint', 'Parsley'],
                    'flowers': ['Marigold', 'Zinnia'],
                    'trees': ['Guava Tree']
                },
                'experienced': {
                    'vegetables': ['Brussels Sprout', 'Celery'],
                    'fruits': ['Avocado', 'Orange'],
                    'herbs': ['Rosemary', 'Thyme'],
                    'flowers': ['Hydrangea', 'Roses'],
                    'trees': ['Mango Tree', 'Jackfruit Tree']
                }
            },
            'Loam Soil': {
                'beginner': {
                    'vegetables': ['Tomato', 'Lettuce', 'Carrot', 'Eggplant'],
                    'fruits': ['Mango', 'Banana', 'Papaya'],
                    'herbs': ['Basil', 'Cilantro'],
                    'flowers': ['Sunflower', 'Cosmos'],
                    'trees': ['Calamansi Tree', 'Lemon Tree']
                },
                'experienced': {
                    'vegetables': ['Cauliflower', 'Bell Pepper'],
                    'fruits': ['Grapes', 'Dragon Fruit'],
                    'herbs': ['Rosemary', 'Oregano', 'Thyme'],
                    'flowers': ['Orchids', 'Dahlias'],
                    'trees': ['Avocado Tree', 'Citrus Orchard Mix']
                }
            },
            'Sandy Soil': {
                'beginner': {
                    'vegetables': ['Carrot', 'Sweet Potato', 'Peanut'],
                    'fruits': ['Watermelon', 'Pineapple'],
                    'herbs': ['Lavender', 'Rosemary'],
                    'flowers': ['Portulaca', 'Gazania'],
                    'trees': ['Casuarina Tree', 'Coconut Palm']
                },
                'experienced': {
                    'vegetables': ['Asparagus', 'Artichoke'],
                    'fruits': ['Grapes', 'Fig'],
                    'herbs': ['Sage', 'Thyme'],
                    'flowers': ['Bougainvillea', 'Bird of Paradise'],
                    'trees': ['Olive Tree', 'Date Palm']
                }
            },
            'Silt Soil': {
                'beginner': {
                    'vegetables': ['Onion', 'Garlic', 'Beetroot'],
                    'fruits': ['Strawberry'],
                    'herbs': ['Chives', 'Dill'],
                    'flowers': ['Iris', 'Daylily'],
                    'trees': ['Willow Tree']
                },
                'experienced': {
                    'vegetables': ['Leek', 'Fennel'],
                    'fruits': ['Blueberry', 'Raspberry'],
                    'herbs': ['Tarragon', 'Lovage'],
                    'flowers': ['Peony', 'Astilbe'],
                    'trees': ['Alder Tree', 'Cherry Tree']
                }
            },
            'Mixed Soil': {
                'beginner': {
                    'vegetables': ['Tomato', 'Lettuce', 'Cucumber'],
                    'fruits': ['Mango', 'Banana'],
                    'herbs': ['Basil', 'Mint'],
                    'flowers': ['Marigold', 'Cosmos'],
                    'trees': ['Guava Tree', 'Lemon Tree']
                },
                'experienced': {
                    'vegetables': ['Broccoli', 'Okra'],
                    'fruits': ['Rambutan', 'Lanzones'],
                    'herbs': ['Coriander', 'Thyme'],
                    'flowers': ['Kalachuchi', 'Roses'],
                    'trees': ['Mango Tree', 'Avocado Tree']
                }
            }
        }

        # Prefer curated soil + skill-based plants; fall back to AI output if needed
        curated_by_soil = soil_based_plants.get(soil_type_class, {})
        curated_plants = curated_by_soil.get(user_level) or ai_result.get('suitable_plants', {})

        # Check if user is premium to determine what data to return
        is_premium = getattr(current_user, 'subscribed', False)
        
        if is_premium:
            # Premium plan - return all detailed information
            result = {
                'moisture_level': ai_result.get('moisture_level', 'Unable to determine from image'),
                'soil_type': soil_type_class,
                'texture': texture_value,
                'organic_matter': ai_result.get('organic_matter', 'Unable to assess from image'),
                'drainage': ai_result.get('drainage', 'Unable to assess from image'),
                'recommendations': ai_result.get('recommendations', []),
                'suitable_plants': curated_plants,
                'nutrient_indicators': ai_result.get('nutrient_indicators', 'Unable to assess from image'),
                'compaction_assessment': ai_result.get('compaction_assessment', 'Unable to assess from image'),
                'soil_health_score': ai_result.get('soil_health_score', 'Unable to assess from image'),
                'seasonal_considerations': ai_result.get('seasonal_considerations', 'Unable to assess from image'),
                'soil_amendments': ai_result.get('soil_amendments', 'Unable to assess from image'),
                'water_retention': ai_result.get('water_retention', 'Unable to assess from image'),
                'root_development': ai_result.get('root_development', 'Unable to assess from image'),
                'ai_analyzed': True,
                'plan_type': 'premium'
            }
        else:
            # Basic plan - return only basic information (moisture, soil type, texture)
            result = {
                'moisture_level': ai_result.get('moisture_level', 'Unable to determine from image'),
                'soil_type': soil_type_class,
                'texture': texture_value,
                'ai_analyzed': True,
                'plan_type': 'basic'
            }

        # Track soil usage after successful analysis (both premium and basic users) - separate from plant analysis
        free_allocation = 10 if is_premium else 3
        usage = _get_or_create_soil_usage(current_user.id)
        is_free = usage.can_use_free(free_allocation)
        old_free_used = usage.free_analyses_used or 0
        old_purchased = usage.purchased_credits or 0
        _track_soil_usage(
            current_user.id,
            is_free=is_free,
            cost=0.00 if is_free else 20.00,
            analysis_result=str(result.get('moisture_level', ''))[:500]
        )
        db.session.commit()
        # Refresh to get updated values
        db.session.refresh(usage)
        new_free_used = usage.free_analyses_used or 0
        new_purchased = usage.purchased_credits or 0
        print(f"✅ Tracked usage for user {current_user.id} (premium={is_premium}, allocation={free_allocation}): free {old_free_used}→{new_free_used}, purchased {old_purchased}→{new_purchased}")

        return jsonify(result)
        
    except Exception as e:
        # Fallback to basic heuristic analysis
        try:
            file = request.files.get('image')
            if not file:
                return jsonify({"error": f"Soil analysis failed: {str(e)}"}), 500
            
            from PIL import Image
            import numpy as np
            image = Image.open(file.stream).convert('RGB').resize((256, 256))
            arr = np.asarray(image, dtype=np.float32)
            
            # Basic color analysis for soil characteristics
            red = arr[:, :, 0]
            green = arr[:, :, 1]
            blue = arr[:, :, 2]
            
            # Calculate average colors
            avg_red = float(np.mean(red))
            avg_green = float(np.mean(green))
            avg_blue = float(np.mean(blue))
            
            # Basic soil type estimation based on color
            if avg_red > 120 and avg_green > 100:
                soil_type = "Clay-rich soil (reddish-brown)"
                ph_estimate = "Slightly acidic to neutral (6.0-7.0)"
            elif avg_red > 100 and avg_green > 80:
                soil_type = "Loamy soil (brown)"
                ph_estimate = "Neutral (6.5-7.5)"
            elif avg_red < 80 and avg_green < 80:
                soil_type = "Sandy soil (light colored)"
                ph_estimate = "Variable, may be alkaline (7.0-8.0)"
            else:
                soil_type = "Mixed soil composition"
                ph_estimate = "Neutral range (6.5-7.5)"
            
            # Basic moisture estimation
            brightness = (avg_red + avg_green + avg_blue) / 3
            if brightness < 80:
                moisture = "Appears moist or wet"
            elif brightness > 150:
                moisture = "Appears dry"
            else:
                moisture = "Moderate moisture level"
            
            # Extract soil type classification from soil_type string
            soil_type_class = "Mixed Soil"
            if "clay" in soil_type.lower() or "clay-rich" in soil_type.lower():
                soil_type_class = "Clay Soil"
            elif "loamy" in soil_type.lower() or "loam" in soil_type.lower():
                soil_type_class = "Loam Soil"
            elif "sandy" in soil_type.lower() or "sand" in soil_type.lower():
                soil_type_class = "Sandy Soil"
            elif "silt" in soil_type.lower():
                soil_type_class = "Silt Soil"
            
            # Check if user is premium to determine what data to return
            is_premium = getattr(current_user, 'subscribed', False)
            
            if is_premium:
                # Premium plan - return all detailed information
                return jsonify({
                    'moisture_level': moisture,
                    'soil_type': soil_type_class,
                    'texture': soil_type,
                    'organic_matter': 'Unable to assess from image - consider soil test',
                    'drainage': 'Unable to assess from image - observe after watering',
                    'recommendations': [
                        'Consider professional soil testing for accurate nutrient levels',
                        'Add organic matter like compost to improve soil structure',
                        'Test drainage by observing how quickly water is absorbed'
                    ],
                    'suitable_plants': ['Most common garden plants will grow in this soil type'],
                    'nutrient_indicators': 'Unable to assess from image - consider soil test',
                    'compaction_assessment': 'Unable to assess from image - test with finger or tool',
                    'ai_analyzed': False,
                    'fallback_analysis': True,
                    'plan_type': 'premium'
                })
            else:
                # Basic plan - return only basic information
                return jsonify({
                    'moisture_level': moisture,
                    'soil_type': soil_type_class,
                    'texture': soil_type,
                    'ai_analyzed': False,
                    'fallback_analysis': True,
                    'plan_type': 'basic'
                })
        except Exception:
            return jsonify({"error": f"Soil analysis failed: {str(e)}"}), 500

@views.route('/camera')
@login_required
def camera():
    from flask_login import current_user
    if current_user.is_admin():
        return jsonify({"error": "Admins do not have access to the camera feature."}), 403
    return jsonify({"message": "Camera feature", "user_id": current_user.id})

@views.route('/seasonal-planning')
@login_required
def seasonal_planning():
    # Get user's location from their garden data
    user_gardens = Garden.query.filter_by(user_id=current_user.id).all()
    location = "Unknown"
    if user_gardens:
        # Use the first garden's location as user's location
        garden = user_gardens[0]
        location = f"{garden.location_city}, {garden.location_country}" if garden.location_city else "Unknown"
    
    # Get current season
    month = datetime.now().month
    if month >= 3 and month <= 5:
        season = 'spring'
    elif month >= 6 and month <= 8:
        season = 'summer'
    elif month >= 9 and month <= 11:
        season = 'fall'
    else:
        season = 'winter'
    
    # Enhanced seasonal data with weather integration
    seasonal_data = {
        'location': location,
        'current_season': season,
        'planting_calendar': {
            'spring': ['Tomatoes', 'Peppers', 'Lettuce', 'Herbs'],
            'summer': ['Cucumbers', 'Zucchini', 'Beans', 'Corn'],
            'fall': ['Kale', 'Spinach', 'Garlic', 'Onions'],
            'winter': ['Indoor herbs', 'Microgreens', 'Sprouts']
        },
        'tips': {
            'spring': ['Test soil pH', 'Start seeds indoors', 'Prepare garden beds'],
            'summer': ['Mulch to retain moisture', 'Water deeply', 'Monitor for pests'],
            'fall': ['Plant cool-season crops', 'Clean up garden', 'Add compost'],
            'winter': ['Plan next year', 'Maintain tools', 'Grow indoors']
        },
        'weather_integration': {
            'optimal_planting_conditions': {
                'temperature_range': '50-75°F',
                'humidity_range': '40-70%',
                'soil_temperature': '50°F+ for most seeds'
            },
            'weather_alerts': {
                'frost_warning': 'Protect tender plants',
                'heat_stress': 'Water early morning/evening',
                'high_humidity': 'Watch for fungal diseases'
            }
        }
    }
    
    return jsonify(seasonal_data)

@views.route('/api/plant-weather-tolerance')
def get_plant_weather_tolerance():
    """Get plant-specific weather tolerance data and warnings"""
    city = request.args.get('city', 'Cebu')
    plant_name = request.args.get('plant', '')
    
    try:
        # Get current weather data
        api_key = os.getenv('OPENWEATHER_API_KEY')
        if not api_key:
            return jsonify({
                "error": "Weather API key not configured",
                "success": False
            }), 500
        
        # Get coordinates first
        geocode_url = f"http://api.openweathermap.org/geo/1.0/direct?q={city}&limit=1&appid={api_key}"
        geocode_response = requests.get(geocode_url)
        
        if geocode_response.status_code != 200:
            return jsonify({
                "error": "City not found",
                "success": False
            }), 404
        
        geocode_data = geocode_response.json()
        if not geocode_data:
            return jsonify({
                "error": "City not found",
                "success": False
            }), 404
        
        lat = geocode_data[0]['lat']
        lon = geocode_data[0]['lon']
        
        # Get current weather
        weather_url = f"https://api.openweathermap.org/data/2.5/weather?lat={lat}&lon={lon}&appid={api_key}&units=metric"
        weather_response = requests.get(weather_url)
        
        if weather_response.status_code == 200:
            weather_data = weather_response.json()
            current_temp = weather_data['main']['temp']
            current_humidity = weather_data['main']['humidity']
            current_wind = weather_data['wind']['speed']
            current_conditions = weather_data['weather'][0]['description']
            
            # Comprehensive plant database with weather tolerances
            plant_database = {
                'tomatoes': {
                    'name': 'Tomatoes',
                    'optimal_temp': {'min': 18, 'max': 29},
                    'tolerance_temp': {'min': 10, 'max': 35},
                    'optimal_humidity': {'min': 50, 'max': 70},
                    'tolerance_humidity': {'min': 30, 'max': 85},
                    'wind_tolerance': 15,  # km/h
                    'heat_stress_temp': 32,
                    'cold_damage_temp': 5,
                    'humidity_disease_risk': 80,
                    'varieties': {
                        'cherry': {'cold_tolerance': 2, 'heat_tolerance': 2},
                        'beefsteak': {'cold_tolerance': 0, 'heat_tolerance': 1},
                        'roma': {'cold_tolerance': 1, 'heat_tolerance': 2}
                    }
                },
                'peppers': {
                    'name': 'Peppers',
                    'optimal_temp': {'min': 21, 'max': 29},
                    'tolerance_temp': {'min': 15, 'max': 32},
                    'optimal_humidity': {'min': 50, 'max': 70},
                    'tolerance_humidity': {'min': 40, 'max': 80},
                    'wind_tolerance': 12,
                    'heat_stress_temp': 30,
                    'cold_damage_temp': 10,
                    'humidity_disease_risk': 75,
                    'varieties': {
                        'bell': {'cold_tolerance': 0, 'heat_tolerance': 1},
                        'jalapeno': {'cold_tolerance': 1, 'heat_tolerance': 2},
                        'habanero': {'cold_tolerance': 2, 'heat_tolerance': 3}
                    }
                },
                'lettuce': {
                    'name': 'Lettuce',
                    'optimal_temp': {'min': 7, 'max': 18},
                    'tolerance_temp': {'min': 2, 'max': 24},
                    'optimal_humidity': {'min': 60, 'max': 80},
                    'tolerance_humidity': {'min': 40, 'max': 90},
                    'wind_tolerance': 20,
                    'heat_stress_temp': 25,
                    'cold_damage_temp': -2,
                    'humidity_disease_risk': 85,
                    'varieties': {
                        'romaine': {'cold_tolerance': 2, 'heat_tolerance': 0},
                        'butterhead': {'cold_tolerance': 1, 'heat_tolerance': 0},
                        'iceberg': {'cold_tolerance': 1, 'heat_tolerance': 1}
                    }
                },
                'cucumbers': {
                    'name': 'Cucumbers',
                    'optimal_temp': {'min': 18, 'max': 26},
                    'tolerance_temp': {'min': 10, 'max': 32},
                    'optimal_humidity': {'min': 60, 'max': 80},
                    'tolerance_humidity': {'min': 40, 'max': 90},
                    'wind_tolerance': 10,
                    'heat_stress_temp': 30,
                    'cold_damage_temp': 5,
                    'humidity_disease_risk': 80,
                    'varieties': {
                        'english': {'cold_tolerance': 0, 'heat_tolerance': 1},
                        'pickling': {'cold_tolerance': 1, 'heat_tolerance': 2},
                        'lemon': {'cold_tolerance': 1, 'heat_tolerance': 1}
                    }
                },
                'herbs': {
                    'name': 'Herbs (Basil, Parsley, Cilantro)',
                    'optimal_temp': {'min': 15, 'max': 24},
                    'tolerance_temp': {'min': 5, 'max': 30},
                    'optimal_humidity': {'min': 40, 'max': 60},
                    'tolerance_humidity': {'min': 30, 'max': 70},
                    'wind_tolerance': 15,
                    'heat_stress_temp': 28,
                    'cold_damage_temp': 2,
                    'humidity_disease_risk': 70,
                    'varieties': {
                        'basil': {'cold_tolerance': 0, 'heat_tolerance': 2},
                        'parsley': {'cold_tolerance': 2, 'heat_tolerance': 0},
                        'cilantro': {'cold_tolerance': 1, 'heat_tolerance': 0}
                    }
                },
                'garlic': {
                    'name': 'Garlic',
                    'optimal_temp': {'min': 0, 'max': 15},
                    'tolerance_temp': {'min': -10, 'max': 25},
                    'optimal_humidity': {'min': 50, 'max': 70},
                    'tolerance_humidity': {'min': 30, 'max': 80},
                    'wind_tolerance': 25,
                    'heat_stress_temp': 25,
                    'cold_damage_temp': -15,
                    'humidity_disease_risk': 75,
                    'varieties': {
                        'hardneck': {'cold_tolerance': 3, 'heat_tolerance': 0},
                        'softneck': {'cold_tolerance': 2, 'heat_tolerance': 1}
                    }
                }
            }
            
            # Analyze current conditions for all plants or specific plant
            analysis_results = []
            
            if plant_name and plant_name.lower() in plant_database:
                plants_to_analyze = [plant_name.lower()]
            else:
                plants_to_analyze = list(plant_database.keys())
            
            for plant_key in plants_to_analyze:
                plant = plant_database[plant_key]
                warnings = []
                recommendations = []
                status = "Good"
                status_color = "green"
                
                # Temperature analysis
                if current_temp < plant['optimal_temp']['min']:
                    warnings.append(f"Temperature too low ({current_temp}°C). Optimal: {plant['optimal_temp']['min']}-{plant['optimal_temp']['max']}°C")
                    if current_temp < plant['cold_damage_temp']:
                        warnings.append(f"⚠️ COLD DAMAGE RISK! Temperature below {plant['cold_damage_temp']}°C")
                        status = "Critical"
                        status_color = "red"
                    else:
                        status = "Poor"
                        status_color = "yellow"
                    recommendations.append("Protect with row covers or bring indoors")
                elif current_temp > plant['optimal_temp']['max']:
                    warnings.append(f"Temperature too high ({current_temp}°C). Optimal: {plant['optimal_temp']['min']}-{plant['optimal_temp']['max']}°C")
                    if current_temp > plant['heat_stress_temp']:
                        warnings.append(f"🌡️ HEAT STRESS RISK! Temperature above {plant['heat_stress_temp']}°C")
                        status = "Critical"
                        status_color = "red"
                    else:
                        status = "Poor"
                        status_color = "yellow"
                    recommendations.append("Provide shade and increase watering frequency")
                
                # Humidity analysis
                if current_humidity < plant['optimal_humidity']['min']:
                    warnings.append(f"Humidity too low ({current_humidity}%). Optimal: {plant['optimal_humidity']['min']}-{plant['optimal_humidity']['max']}%")
                    recommendations.append("Increase watering frequency and consider misting")
                elif current_humidity > plant['optimal_humidity']['max']:
                    warnings.append(f"Humidity too high ({current_humidity}%). Optimal: {plant['optimal_humidity']['min']}-{plant['optimal_humidity']['max']}%")
                    if current_humidity > plant['humidity_disease_risk']:
                        warnings.append(f"🦠 DISEASE RISK! High humidity promotes fungal diseases")
                        if status != "Critical":
                            status = "Poor"
                            status_color = "yellow"
                    recommendations.append("Improve air circulation and avoid overhead watering")
                
                # Wind analysis
                if current_wind > plant['wind_tolerance']:
                    warnings.append(f"Wind too strong ({current_wind} km/h). Tolerance: {plant['wind_tolerance']} km/h")
                    recommendations.append("Provide wind protection or stake plants")
                
                # Condition-specific warnings
                if 'rain' in current_conditions.lower():
                    recommendations.append("Rain provides natural watering - reduce irrigation")
                if 'storm' in current_conditions.lower():
                    warnings.append("🌪️ STORM CONDITIONS - Protect plants from damage")
                    if status != "Critical":
                        status = "Poor"
                        status_color = "yellow"
                
                analysis_results.append({
                    'plant_name': plant['name'],
                    'current_conditions': {
                        'temperature': current_temp,
                        'humidity': current_humidity,
                        'wind_speed': current_wind,
                        'conditions': current_conditions
                    },
                    'optimal_ranges': {
                        'temperature': f"{plant['optimal_temp']['min']}-{plant['optimal_temp']['max']}°C",
                        'humidity': f"{plant['optimal_humidity']['min']}-{plant['optimal_humidity']['max']}%"
                    },
                    'warnings': warnings,
                    'recommendations': recommendations,
                    'status': status,
                    'status_color': status_color,
                    'varieties': plant['varieties']
                })
            
            return jsonify({
                "plant_analysis": analysis_results,
                "current_weather": {
                    'temperature': current_temp,
                    'humidity': current_humidity,
                    'wind_speed': current_wind,
                    'conditions': current_conditions
                },
                "city": city,
                "success": True
            })
        else:
            return jsonify({
                "error": f"Weather service error: {weather_response.status_code}",
                "success": False
            }), weather_response.status_code
            
    except Exception as e:
        return jsonify({
            "error": f"Error fetching plant weather tolerance: {str(e)}",
            "success": False
        }), 500

@views.route('/api/microclimate-zones')
def get_microclimate_zones():
    """Get microclimate zone recommendations for garden areas"""
    city = request.args.get('city', 'Cebu')
    
    try:
        # Get current weather data
        api_key = os.getenv('OPENWEATHER_API_KEY')
        if not api_key:
            return jsonify({
                "error": "Weather API key not configured",
                "success": False
            }), 500
        
        # Get coordinates first
        geocode_url = f"http://api.openweathermap.org/geo/1.0/direct?q={city}&limit=1&appid={api_key}"
        geocode_response = requests.get(geocode_url)
        
        if geocode_response.status_code != 200:
            return jsonify({
                "error": "City not found",
                "success": False
            }), 404
        
        geocode_data = geocode_response.json()
        if not geocode_data:
            return jsonify({
                "error": "City not found",
                "success": False
            }), 404
        
        lat = geocode_data[0]['lat']
        lon = geocode_data[0]['lon']
        
        # Get current weather
        weather_url = f"https://api.openweathermap.org/data/2.5/weather?lat={lat}&lon={lon}&appid={api_key}&units=metric"
        weather_response = requests.get(weather_url)
        
        if weather_response.status_code == 200:
            weather_data = weather_response.json()
            current_temp = weather_data['main']['temp']
            current_humidity = weather_data['main']['humidity']
            current_wind = weather_data['wind']['speed']
            
            # Define microclimate zones based on typical garden layouts
            microclimate_zones = [
                {
                    'name': 'Full Sun Zone',
                    'description': 'Areas receiving 6+ hours of direct sunlight daily',
                    'temperature_modifier': 3,  # +3°C warmer than ambient
                    'humidity_modifier': -10,   # -10% humidity
                    'wind_modifier': 1.2,       # 20% more wind exposure
                    'suitable_plants': ['Tomatoes', 'Peppers', 'Basil', 'Rosemary', 'Sunflowers'],
                    'care_tips': [
                        'Water deeply and frequently during hot weather',
                        'Use mulch to retain soil moisture',
                        'Consider shade cloth during extreme heat',
                        'Monitor for sunburn on leaves'
                    ],
                    'seasonal_considerations': {
                        'spring': 'Perfect for starting warm-season crops',
                        'summer': 'May need extra watering and heat protection',
                        'fall': 'Excellent for extending growing season',
                        'winter': 'Good for cool-season crops in mild climates'
                    }
                },
                {
                    'name': 'Partial Shade Zone',
                    'description': 'Areas receiving 3-6 hours of direct sunlight daily',
                    'temperature_modifier': 0,  # Same as ambient
                    'humidity_modifier': 5,     # +5% humidity
                    'wind_modifier': 0.8,       # 20% less wind exposure
                    'suitable_plants': ['Lettuce', 'Spinach', 'Kale', 'Cilantro', 'Mint'],
                    'care_tips': [
                        'Monitor soil moisture - may dry slower',
                        'Good air circulation to prevent fungal issues',
                        'Ideal for succession planting',
                        'Protect from late afternoon sun'
                    ],
                    'seasonal_considerations': {
                        'spring': 'Excellent for cool-season crops',
                        'summer': 'Provides relief from intense heat',
                        'fall': 'Perfect for extending harvest season',
                        'winter': 'Good protection for tender plants'
                    }
                },
                {
                    'name': 'Full Shade Zone',
                    'description': 'Areas receiving less than 3 hours of direct sunlight',
                    'temperature_modifier': -2, # -2°C cooler than ambient
                    'humidity_modifier': 15,    # +15% humidity
                    'wind_modifier': 0.5,       # 50% less wind exposure
                    'suitable_plants': ['Hostas', 'Ferns', 'Mint', 'Parsley', 'Chives'],
                    'care_tips': [
                        'Focus on moisture-loving plants',
                        'Improve drainage to prevent waterlogging',
                        'Use reflective surfaces to increase light',
                        'Monitor for fungal diseases due to high humidity'
                    ],
                    'seasonal_considerations': {
                        'spring': 'Slow to warm up - plant later',
                        'summer': 'Cool refuge for heat-sensitive plants',
                        'fall': 'Good for overwintering tender perennials',
                        'winter': 'Protection from frost damage'
                    }
                },
                {
                    'name': 'Wind-Protected Zone',
                    'description': 'Areas sheltered by buildings, fences, or trees',
                    'temperature_modifier': 1,  # +1°C warmer than ambient
                    'humidity_modifier': 10,    # +10% humidity
                    'wind_modifier': 0.3,       # 70% less wind exposure
                    'suitable_plants': ['Delicate herbs', 'Young seedlings', 'Climbing plants'],
                    'care_tips': [
                        'Excellent for starting seeds and seedlings',
                        'Monitor humidity levels',
                        'Good for vertical growing structures',
                        'Protect from late frosts'
                    ],
                    'seasonal_considerations': {
                        'spring': 'Perfect for early season planting',
                        'summer': 'Reduced water loss from wind',
                        'fall': 'Extended growing season',
                        'winter': 'Excellent frost protection'
                    }
                },
                {
                    'name': 'Heat Trap Zone',
                    'description': 'Areas near walls, patios, or other heat-absorbing surfaces',
                    'temperature_modifier': 5,  # +5°C warmer than ambient
                    'humidity_modifier': -15,   # -15% humidity
                    'wind_modifier': 0.6,       # 40% less wind exposure
                    'suitable_plants': ['Heat-loving peppers', 'Eggplant', 'Okra', 'Sweet potatoes'],
                    'care_tips': [
                        'Water frequently - soil dries quickly',
                        'Use heat-tolerant varieties',
                        'Consider container gardening for mobility',
                        'Monitor for heat stress'
                    ],
                    'seasonal_considerations': {
                        'spring': 'Warms up quickly - early planting possible',
                        'summer': 'May be too hot for most plants',
                        'fall': 'Excellent for extending warm-season crops',
                        'winter': 'Good for overwintering tender plants'
                    }
                }
            ]
            
            # Calculate zone-specific conditions
            zone_analysis = []
            for zone in microclimate_zones:
                zone_temp = round(current_temp + zone['temperature_modifier'], 1)
                zone_humidity = max(0, min(100, current_humidity + zone['humidity_modifier']))
                zone_wind = round(current_wind * zone['wind_modifier'], 1)
                
                # Determine zone suitability for current conditions
                suitability_score = 0
                suitability_notes = []
                
                if 15 <= zone_temp <= 30:
                    suitability_score += 3
                    suitability_notes.append("Optimal temperature range")
                elif 10 <= zone_temp < 15 or 30 < zone_temp <= 35:
                    suitability_score += 2
                    suitability_notes.append("Good temperature range")
                elif zone_temp < 10:
                    suitability_score += 1
                    suitability_notes.append("Cool conditions - good for cool-season crops")
                else:
                    suitability_score += 0
                    suitability_notes.append("Extreme temperatures - monitor closely")
                
                if 40 <= zone_humidity <= 80:
                    suitability_score += 2
                    suitability_notes.append("Good humidity levels")
                elif zone_humidity > 80:
                    suitability_score += 1
                    suitability_notes.append("High humidity - watch for diseases")
                else:
                    suitability_score += 1
                    suitability_notes.append("Low humidity - increase watering")
                
                if zone_wind < 15:
                    suitability_score += 2
                    suitability_notes.append("Calm conditions")
                elif zone_wind < 25:
                    suitability_score += 1
                    suitability_notes.append("Moderate wind")
                else:
                    suitability_score += 0
                    suitability_notes.append("High winds - provide protection")
                
                # Determine overall suitability
                if suitability_score >= 6:
                    suitability = "Excellent"
                    suitability_color = "green"
                elif suitability_score >= 4:
                    suitability = "Good"
                    suitability_color = "blue"
                elif suitability_score >= 2:
                    suitability = "Fair"
                    suitability_color = "yellow"
                else:
                    suitability = "Poor"
                    suitability_color = "red"
                
                zone_analysis.append({
                    'zone_info': zone,
                    'current_conditions': {
                        'temperature': zone_temp,
                        'humidity': zone_humidity,
                        'wind_speed': zone_wind
                    },
                    'suitability': {
                        'score': suitability_score,
                        'rating': suitability,
                        'color': suitability_color,
                        'notes': suitability_notes
                    }
                })
            
            return jsonify({
                "microclimate_zones": zone_analysis,
                "base_conditions": {
                    'temperature': current_temp,
                    'humidity': current_humidity,
                    'wind_speed': current_wind
                },
                "city": city,
                "success": True
            })
        else:
            return jsonify({
                "error": f"Weather service error: {weather_response.status_code}",
                "success": False
            }), weather_response.status_code
            
    except Exception as e:
        return jsonify({
            "error": f"Error fetching microclimate zones: {str(e)}",
            "success": False
        }), 500

@views.route('/api/soil-temperature')
def get_soil_temperature():
    """Get soil temperature recommendations based on weather data"""
    city = request.args.get('city', 'Cebu')
    
    try:
        # Get current weather data
        api_key = os.getenv('OPENWEATHER_API_KEY')
        if not api_key:
            return jsonify({
                "error": "Weather API key not configured",
                "success": False
            }), 500
        
        # Get coordinates first
        geocode_url = f"http://api.openweathermap.org/geo/1.0/direct?q={city}&limit=1&appid={api_key}"
        geocode_response = requests.get(geocode_url)
        
        if geocode_response.status_code != 200:
            return jsonify({
                "error": "City not found",
                "success": False
            }), 404
        
        geocode_data = geocode_response.json()
        if not geocode_data:
            return jsonify({
                "error": "City not found",
                "success": False
            }), 404
        
        lat = geocode_data[0]['lat']
        lon = geocode_data[0]['lon']
        
        # Get current weather
        weather_url = f"https://api.openweathermap.org/data/2.5/weather?lat={lat}&lon={lon}&appid={api_key}&units=metric"
        weather_response = requests.get(weather_url)
        
        if weather_response.status_code == 200:
            weather_data = weather_response.json()
            air_temp = weather_data['main']['temp']
            
            # Estimate soil temperature using improved formula (without IoT sensors)
            # Based on agricultural research and meteorological data:
            # - Soil temp varies by depth, time of day, season, and soil type
            # - Shallow soil (2-4") follows air temp more closely
            # - Deep soil (6-8") is more stable and cooler
            # - Daytime: soil is 2-4°C cooler than air
            # - Nighttime: soil retains heat, 1-2°C cooler than air
            # - Summer: soil stays warmer
            # - Winter: soil stays cooler
            
            from datetime import datetime
            now = datetime.now()
            hour = now.hour
            month = now.month
            is_daytime = 6 <= hour < 18
            is_summer = 4 <= month <= 9  # April-September in Philippines
            
            if is_daytime:
                # Daytime: soil is cooler than air
                soil_temp_shallow = round(air_temp - (2 if is_summer else 3), 1)  # 2-4 inches deep
                soil_temp_deep = round(air_temp - (4 if is_summer else 5), 1)      # 6-8 inches deep
            else:
                # Nighttime: soil retains heat better
                soil_temp_shallow = round(air_temp - (1 if is_summer else 2), 1)   # 2-4 inches deep
                soil_temp_deep = round(air_temp - (3 if is_summer else 4), 1)     # 6-8 inches deep
            
            # Note: This is an estimation. For accurate readings, use a soil thermometer.
            
            # Generate planting recommendations based on soil temperature
            recommendations = []
            
            if soil_temp_shallow >= 10:
                recommendations.append({
                    "plant_type": "Warm-season vegetables",
                    "plants": ["Tomatoes", "Peppers", "Eggplant", "Cucumbers", "Squash"],
                    "soil_temp_range": "10-15°C+",
                    "current_temp": f"{soil_temp_shallow}°C",
                    "status": "Ready to plant" if soil_temp_shallow >= 15 else "Almost ready",
                    "color": "green" if soil_temp_shallow >= 15 else "yellow"
                })
            
            if soil_temp_shallow >= 5:
                recommendations.append({
                    "plant_type": "Cool-season vegetables",
                    "plants": ["Lettuce", "Spinach", "Kale", "Peas", "Radishes"],
                    "soil_temp_range": "5-10°C",
                    "current_temp": f"{soil_temp_shallow}°C",
                    "status": "Ready to plant" if 5 <= soil_temp_shallow <= 15 else "Check conditions",
                    "color": "green" if 5 <= soil_temp_shallow <= 15 else "blue"
                })
            
            if soil_temp_shallow < 5:
                recommendations.append({
                    "plant_type": "Cold-tolerant crops",
                    "plants": ["Garlic", "Onions", "Carrots", "Beets"],
                    "soil_temp_range": "2-5°C",
                    "current_temp": f"{soil_temp_shallow}°C",
                    "status": "Ready to plant" if soil_temp_shallow >= 2 else "Too cold",
                    "color": "green" if soil_temp_shallow >= 2 else "red"
                })
            
            # Add general soil temperature tips
            tips = []
            if soil_temp_shallow < 10:
                tips.append("Soil is still cool. Consider using row covers or black plastic to warm the soil.")
            elif soil_temp_shallow > 20:
                tips.append("Soil is warm. Perfect for heat-loving plants like tomatoes and peppers.")
            else:
                tips.append("Soil temperature is moderate. Good for most cool and warm-season crops.")
            
            return jsonify({
                "soil_temperature": {
                    "shallow": soil_temp_shallow,
                    "deep": soil_temp_deep,
                    "air_temperature": air_temp
                },
                "recommendations": recommendations,
                "tips": tips,
                "city": city,
                "success": True
            })
        else:
            return jsonify({
                "error": f"Weather service error: {weather_response.status_code}",
                "success": False
            }), weather_response.status_code
            
    except Exception as e:
        return jsonify({
            "error": f"Error fetching soil temperature data: {str(e)}",
            "success": False
        }), 500

@views.route('/api/weather-forecast')
def get_weather_forecast():
    """Get 7-day weather forecast for better planting timing"""
    city = request.args.get('city', 'Cebu')
    
    try:
        # Get 7-day forecast from OpenWeatherMap
        api_key = os.getenv('OPENWEATHER_API_KEY')
        if not api_key:
            return jsonify({
                "error": "Weather API key not configured",
                "success": False
            }), 500
        
        # Get coordinates first
        geocode_url = f"http://api.openweathermap.org/geo/1.0/direct?q={city}&limit=1&appid={api_key}"
        geocode_response = requests.get(geocode_url)
        
        if geocode_response.status_code != 200:
            return jsonify({
                "error": "City not found",
                "success": False
            }), 404
        
        geocode_data = geocode_response.json()
        if not geocode_data:
            return jsonify({
                "error": "City not found",
                "success": False
            }), 404
        
        lat = geocode_data[0]['lat']
        lon = geocode_data[0]['lon']
        
        # Get current weather for today's temperature
        current_weather_url = f"http://api.openweathermap.org/data/2.5/weather?lat={lat}&lon={lon}&appid={api_key}&units=metric"
        current_weather_response = requests.get(current_weather_url, timeout=10)
        current_temp = None
        if current_weather_response.status_code == 200:
            current_weather_data = current_weather_response.json()
            current_temp = round(current_weather_data['main']['temp'])
        
        # Get 7-day forecast
        forecast_url = f"https://api.openweathermap.org/data/2.5/forecast?lat={lat}&lon={lon}&appid={api_key}&units=metric"
        forecast_response = requests.get(forecast_url)
        
        if forecast_response.status_code == 200:
            forecast_data = forecast_response.json()
            
            # Get today's date in YYYY-MM-DD format for comparison
            today = datetime.now().strftime('%Y-%m-%d')
            
            # Process forecast data for planting recommendations
            daily_forecasts = []
            current_date = None
            daily_data = {}
            
            for item in forecast_data['list']:
                date = item['dt_txt'].split(' ')[0]
                if date != current_date:
                    if current_date and daily_data:
                        daily_forecasts.append(daily_data)
                    current_date = date
                    daily_data = {
                        'date': date,
                        'temperatures': [],
                        'humidity': [],
                        'conditions': [],
                        'wind_speed': [],
                        'rain_probability': []
                    }
                
                daily_data['temperatures'].append(item['main']['temp'])
                daily_data['humidity'].append(item['main']['humidity'])
                daily_data['conditions'].append(item['weather'][0]['description'])
                daily_data['wind_speed'].append(item['wind']['speed'])
                if 'rain' in item:
                    daily_data['rain_probability'].append(item['rain'].get('3h', 0))
                else:
                    daily_data['rain_probability'].append(0)
            
            if daily_data:
                daily_forecasts.append(daily_data)
            
            # Process daily data
            processed_forecast = []
            for day in daily_forecasts[:7]:  # Limit to 7 days
                # Use current temperature for today if available, otherwise use average
                is_today = day['date'] == today
                if is_today and current_temp is not None:
                    # For today, use current temperature as the average
                    avg_temp = current_temp
                else:
                    avg_temp = round(sum(day['temperatures']) / len(day['temperatures']))
                min_temp = round(min(day['temperatures']))
                max_temp = round(max(day['temperatures']))
                avg_humidity = round(sum(day['humidity']) / len(day['humidity']))
                avg_wind = round(sum(day['wind_speed']) / len(day['wind_speed']))
                total_rain = sum(day['rain_probability'])
                
                # Determine planting conditions
                planting_score = 0
                conditions = []
                
                if 50 <= avg_temp <= 75:
                    planting_score += 3
                    conditions.append("Optimal temperature")
                elif 45 <= avg_temp < 50 or 75 < avg_temp <= 80:
                    planting_score += 2
                    conditions.append("Good temperature")
                elif avg_temp < 45:
                    planting_score += 0
                    conditions.append("Too cold for most plants")
                else:
                    planting_score += 1
                    conditions.append("Hot weather - water frequently")
                
                if 40 <= avg_humidity <= 70:
                    planting_score += 2
                    conditions.append("Good humidity")
                elif avg_humidity > 70:
                    planting_score += 1
                    conditions.append("High humidity - watch for diseases")
                else:
                    planting_score += 1
                    conditions.append("Low humidity - water more")
                
                if avg_wind < 15:
                    planting_score += 2
                    conditions.append("Calm conditions")
                elif avg_wind < 25:
                    planting_score += 1
                    conditions.append("Moderate wind")
                else:
                    planting_score += 0
                    conditions.append("High winds - avoid planting")
                
                if total_rain > 5:
                    planting_score += 1
                    conditions.append("Rain expected - natural watering")
                
                # Determine overall recommendation
                if planting_score >= 6:
                    recommendation = "Excellent planting day"
                    recommendation_color = "green"
                elif planting_score >= 4:
                    recommendation = "Good planting day"
                    recommendation_color = "blue"
                elif planting_score >= 2:
                    recommendation = "Fair planting day"
                    recommendation_color = "yellow"
                else:
                    recommendation = "Poor planting day"
                    recommendation_color = "red"
                
                processed_forecast.append({
                    'date': day['date'],
                    'temperature': {
                        'average': avg_temp,
                        'min': min_temp,
                        'max': max_temp
                    },
                    'humidity': avg_humidity,
                    'wind_speed': avg_wind,
                    'rain_probability': round(total_rain, 1),
                    'conditions': conditions,
                    'planting_score': planting_score,
                    'recommendation': recommendation,
                    'recommendation_color': recommendation_color
                })
            
            # Fill in missing days if API only returned 5 days (OpenWeatherMap free tier limit)
            if len(processed_forecast) < 7:
                # Get the last date from processed forecast
                last_date_str = processed_forecast[-1]['date'] if processed_forecast else today
                last_date = datetime.strptime(last_date_str, '%Y-%m-%d')
                
                # Get average values from last few days for extrapolation
                if len(processed_forecast) >= 3:
                    recent_temps = [d['temperature']['average'] for d in processed_forecast[-3:]]
                    recent_humidity = [d['humidity'] for d in processed_forecast[-3:]]
                    recent_wind = [d['wind_speed'] for d in processed_forecast[-3:]]
                    recent_scores = [d['planting_score'] for d in processed_forecast[-3:]]
                    
                    avg_recent_temp = round(sum(recent_temps) / len(recent_temps))
                    avg_recent_humidity = round(sum(recent_humidity) / len(recent_humidity))
                    avg_recent_wind = round(sum(recent_wind) / len(recent_wind))
                    avg_recent_score = round(sum(recent_scores) / len(recent_scores))
                else:
                    # Fallback to last day's values
                    last_day = processed_forecast[-1] if processed_forecast else None
                    if last_day:
                        avg_recent_temp = last_day['temperature']['average']
                        avg_recent_humidity = last_day['humidity']
                        avg_recent_wind = last_day['wind_speed']
                        avg_recent_score = last_day['planting_score']
                    else:
                        avg_recent_temp = 28
                        avg_recent_humidity = 70
                        avg_recent_wind = 10
                        avg_recent_score = 6
                
                # Generate missing days (6th and 7th day)
                for i in range(len(processed_forecast), 7):
                    next_date = last_date + timedelta(days=i - len(processed_forecast) + 1)
                    next_date_str = next_date.strftime('%Y-%m-%d')
                    
                    # Slight variation in temperature (±2°C)
                    temp_variation = (i % 2) * 2 - 1  # Alternate between -1 and +1
                    extrapolated_temp = avg_recent_temp + temp_variation
                    extrapolated_min = extrapolated_temp - 3
                    extrapolated_max = extrapolated_temp + 3
                    
                    # Determine planting score based on extrapolated temperature
                    planting_score = 6
                    conditions = []
                    
                    if 50 <= extrapolated_temp <= 75:
                        planting_score = 6
                        conditions.append("Optimal temperature")
                    elif 45 <= extrapolated_temp < 50 or 75 < extrapolated_temp <= 80:
                        planting_score = 5
                        conditions.append("Good temperature")
                    elif extrapolated_temp < 45:
                        planting_score = 3
                        conditions.append("Too cold for most plants")
                    else:
                        planting_score = 4
                        conditions.append("Hot weather - water frequently")
                    
                    if 40 <= avg_recent_humidity <= 70:
                        planting_score += 2
                        conditions.append("Good humidity")
                    elif avg_recent_humidity > 70:
                        planting_score += 1
                        conditions.append("High humidity - watch for diseases")
                    else:
                        planting_score += 1
                        conditions.append("Low humidity - water more")
                    
                    if avg_recent_wind < 15:
                        planting_score += 2
                        conditions.append("Calm conditions")
                    elif avg_recent_wind < 25:
                        planting_score += 1
                        conditions.append("Moderate wind")
                    else:
                        planting_score += 0
                        conditions.append("High winds - avoid planting")
                    
                    # Cap score at 8
                    planting_score = min(planting_score, 8)
                    
                    # Determine recommendation
                    if planting_score >= 6:
                        recommendation = "Excellent planting day"
                        recommendation_color = "green"
                    elif planting_score >= 4:
                        recommendation = "Good planting day"
                        recommendation_color = "blue"
                    elif planting_score >= 2:
                        recommendation = "Fair planting day"
                        recommendation_color = "yellow"
                    else:
                        recommendation = "Poor planting day"
                        recommendation_color = "red"
                    
                    processed_forecast.append({
                        'date': next_date_str,
                        'temperature': {
                            'average': extrapolated_temp,
                            'min': extrapolated_min,
                            'max': extrapolated_max
                        },
                        'humidity': avg_recent_humidity,
                        'wind_speed': avg_recent_wind,
                        'rain_probability': 0.0,
                        'conditions': conditions,
                        'planting_score': planting_score,
                        'recommendation': recommendation,
                        'recommendation_color': recommendation_color
                    })
            
            return jsonify({
                "forecast": processed_forecast,
                "city": city,
                "success": True
            })
        else:
            return jsonify({
                "error": f"Weather service error: {forecast_response.status_code}",
                "success": False
            }), forecast_response.status_code
            
    except Exception as e:
        return jsonify({
            "error": f"Error fetching forecast: {str(e)}",
            "success": False
        }), 500

@views.route('/learning-paths')
def learning_paths():
    return jsonify({"message": "Learning paths feature"})

@views.route('/track-plants')
def track_plants():
    return jsonify({"message": "Plant tracking feature"})

@views.route('/gardening-tips')
def gardening_tips():
    return jsonify({"message": "Gardening tips feature"})

@views.route('/notifications')
@login_required
def notifications():
    """Get active notifications for the current user"""
    try:
        # Get all active notifications that haven't expired
        now = datetime.now(timezone.utc)
        active_notifications = Notification.query.filter(
            Notification.is_active == True,
            db.or_(
                Notification.expires_at.is_(None),
                Notification.expires_at > now
            )
        ).order_by(
            # Sort by priority (High > Medium > Low), then by created_at (newest first)
            db.case(
                (Notification.priority == 'High', 3),
                (Notification.priority == 'Medium', 2),
                (Notification.priority == 'Low', 1),
                else_=0
            ).desc(),
            Notification.created_at.desc()
        ).all()
        
        # Convert to dictionary format
        notifications_data = [notification.to_dict() for notification in active_notifications]
        
        return jsonify(notifications_data)
    except Exception as e:
        print(f"Error fetching notifications: {str(e)}")
        return jsonify([])  # Return empty array on error

@views.route('/garden', methods=['GET'])
@login_required
def garden():
    # Only regular users can manage gardens
    if hasattr(current_user, 'is_admin') and current_user.is_admin():
        return jsonify({"error": "Only user accounts can access gardens. Please sign in with a user account."}), 403
    # List gardens and plants for the current user
    gardens = Garden.query.filter_by(user_id=current_user.id).all()
    # For each garden, get plants via PlantTracking
    plant_trackings = PlantTracking.query.join(Garden).filter(Garden.user_id == current_user.id).all()
    plants = []
    for pt in plant_trackings:
        plant = Plant.query.get(pt.plant_id)
        if plant:
            # Get the latest image for this plant from grid spaces
            latest_grid_space = GridSpace.query.filter_by(plant_id=plant.id).filter(GridSpace.image_path.isnot(None)).order_by(GridSpace.last_updated.desc()).first()
            latest_image = latest_grid_space.image_path if latest_grid_space else None
            print(f"🌱 Plant {plant.name} (ID: {plant.id}) - Latest image: {latest_image}")
            
            plants.append({
                'tracking': {
                    'id': pt.id,
                    'planting_date': pt.planting_date.isoformat() if pt.planting_date else None
                },
                'plant': {
                    'id': plant.id,
                    'name': plant.name,
                    'type': plant.type,
                    'environment': plant.environment,
                    'care_guide': plant.care_guide,
                    'ideal_soil_type': plant.ideal_soil_type,
                    'watering_frequency': plant.watering_frequency,
                    'fertilizing_frequency': plant.fertilizing_frequency,
                    'pruning_frequency': plant.pruning_frequency,
                    'latest_image': latest_image or plant.image_path
                },
                'garden': {
                    'id': pt.garden.id,
                    'name': pt.garden.name,
                    'garden_type': pt.garden.garden_type,
                    'location_city': pt.garden.location_city,
                    'location_country': pt.garden.location_country
                }
            })
    
    print(f"🌱 Returning {len(plants)} plants to frontend")
    for plant_data in plants:
        print(f"  - {plant_data['plant']['name']} (ID: {plant_data['plant']['id']}) - Image: {plant_data['plant']['latest_image']}")
    
    return jsonify({
        "gardens": [{
            'id': g.id,
            'name': g.name,
            'garden_type': g.garden_type,
            'location_city': g.location_city,
            'location_country': g.location_country,
            'grid_size': getattr(g, 'grid_size', '3x3'),
            'base_grid_spaces': getattr(g, 'base_grid_spaces', 9),
            'additional_spaces_purchased': getattr(g, 'additional_spaces_purchased', 0),
            'total_grid_spaces': getattr(g, 'base_grid_spaces', 9) + getattr(g, 'additional_spaces_purchased', 0),
            'used_grid_spaces': getattr(g, 'used_grid_spaces', 0)
        } for g in gardens],
        "plants": plants
    })

@views.route('/garden/add', methods=['POST'])
@login_required
def add_garden():
    # Only regular users can create gardens
    if hasattr(current_user, 'is_admin') and current_user.is_admin():
        return jsonify({"error": "Admins cannot create gardens. Please use a user account."}), 403
    try:
        data = request.get_json(force=True, silent=True) or {}
        name = data.get('name')
        garden_type = data.get('garden_type')
        location_city = data.get('location_city')
        location_country = data.get('location_country')

        if not name or not garden_type:
            return jsonify({"error": "Name and type are required."}), 400

        # Determine grid size based on subscription plan
        is_premium = getattr(current_user, 'subscribed', False) or getattr(current_user, 'subscription_plan', 'basic') == 'premium'
        grid_size = '6x6' if is_premium else '3x3'
        base_grid_spaces = 36 if is_premium else 9

        garden = Garden(
            user_id=current_user.id,
            name=name,
            garden_type=garden_type,
            location_city=location_city,
            location_country=location_country,
            grid_size=grid_size,
            base_grid_spaces=base_grid_spaces,
            additional_spaces_purchased=0,
            used_grid_spaces=0
        )
        db.session.add(garden)
        db.session.flush()  # Get the garden ID
        
        # Create grid spaces for the garden
        create_grid_spaces_for_garden(garden.id, grid_size)
        
        db.session.commit()

        return jsonify({"success": True, "message": "Garden added successfully!", "garden_id": garden.id})
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@views.route('/garden/edit/<int:garden_id>', methods=['POST'])
@login_required
def edit_garden(garden_id):
    garden = Garden.query.get_or_404(garden_id)
    if garden.user_id != current_user.id:
        return jsonify({"error": "Unauthorized."}), 403
    
    data = request.get_json()
    garden.name = data.get('name')
    garden.garden_type = data.get('garden_type')
    garden.location_city = data.get('location_city')
    garden.location_country = data.get('location_country')
    db.session.commit()
    
    return jsonify({"message": "Garden updated successfully!"})

@views.route('/garden/delete/<int:garden_id>', methods=['POST'])
@login_required
def delete_garden(garden_id):
    try:
        garden = Garden.query.get_or_404(garden_id)
        if garden.user_id != current_user.id:
            return jsonify({"error": "Unauthorized."}), 403
        
        # Clean up related data before deleting garden
        # Remove all plants from grid spaces in this garden
        grid_spaces = GridSpace.query.filter_by(garden_id=garden_id).all()
        for space in grid_spaces:
            if space.plant_id:
                # Remove plant from this space
                space.plant_id = None
                space.planting_date = None
                space.last_watered = None
                space.last_fertilized = None
                space.last_pruned = None
                space.notes = None
                space.image_path = None
                space.care_suggestions = None
                space.last_updated = None
        
        # Delete all grid spaces for this garden
        GridSpace.query.filter_by(garden_id=garden_id).delete()
        
        # Delete all plant trackings for this garden
        PlantTracking.query.filter_by(garden_id=garden_id).delete()
        
        # Delete the garden
        db.session.delete(garden)
        db.session.commit()
        
        return jsonify({"message": "Garden deleted successfully!"})
        
    except Exception as e:
        db.session.rollback()
        print(f"Error deleting garden {garden_id}: {str(e)}")
        return jsonify({"error": f"Failed to delete garden: {str(e)}"}), 500

@views.route('/plant/add', methods=['POST'])
@login_required
def add_plant():
    try:
        # Only regular users can add plants to gardens
        if hasattr(current_user, 'is_admin') and current_user.is_admin():
            return jsonify({"error": "Admins cannot add plants. Please use a user account."}), 403
        
        # Check if this is a form data request (with image) or JSON request
        if request.content_type and 'multipart/form-data' in request.content_type:
            # Handle form data with image
            name = request.form.get('name')
            type_ = request.form.get('type')
            environment = request.form.get('environment')
            care_guide = request.form.get('care_guide')
            ideal_soil_type = request.form.get('ideal_soil_type')
            watering_frequency = request.form.get('watering_frequency')
            fertilizing_frequency = request.form.get('fertilizing_frequency')
            pruning_frequency = request.form.get('pruning_frequency')
            garden_id = request.form.get('garden_id')
            planting_date = request.form.get('planting_date')
            
            # Handle image upload
            image_path = None
            if 'image' in request.files:
                file = request.files['image']
                if file and file.filename:
                    # Create uploads directory if it doesn't exist
                    upload_dir = os.path.join(os.getcwd(), 'uploads', 'plants')
                    os.makedirs(upload_dir, exist_ok=True)
                    
                    # Generate unique filename with proper extension
                    import time
                    # Get file extension from content type or default to jpg
                    content_type = file.content_type
                    if content_type and 'image/' in content_type:
                        ext = content_type.split('/')[-1]
                        if ext not in ['jpg', 'jpeg', 'png', 'webp', 'gif']:
                            ext = 'jpg'
                    else:
                        ext = 'jpg'
                    
                    # Clean filename to remove blob references
                    original_filename = file.filename or 'image'
                    if 'blob' in original_filename.lower():
                        original_filename = 'camera_image'
                    
                    filename = f"plant_{int(time.time())}_{original_filename}.{ext}"
                    file_path = os.path.join(upload_dir, filename)
                    file.save(file_path)
                    image_path = f"uploads/plants/{filename}"
        else:
            # Handle JSON data
            data = request.get_json()
            name = data.get('name')
            type_ = data.get('type')
            environment = data.get('environment')
            care_guide = data.get('care_guide')
            ideal_soil_type = data.get('ideal_soil_type')
            watering_frequency = data.get('watering_frequency')
            fertilizing_frequency = data.get('fertilizing_frequency')
            pruning_frequency = data.get('pruning_frequency')
            garden_id = data.get('garden_id')
            planting_date = data.get('planting_date')
            image_path = None
        
        # Validate required fields
        if not type_ or not environment or not care_guide or not garden_id or not planting_date:
            return jsonify({"error": "Missing required fields: type, environment, care_guide, garden_id, and planting_date are required."}), 400

        # Handle static/demo gardens vs real gardens
        if str(garden_id).startswith('static-'):
            return jsonify({"error": "Cannot add plants to demo gardens. Please create a real garden first."}), 400
        
        # Convert garden_id to integer for real gardens
        try:
            garden_id = int(garden_id)
        except (ValueError, TypeError):
            return jsonify({"error": "Invalid garden_id. Must be a valid number."}), 400

        # Verify garden exists and belongs to current user
        garden = Garden.query.get(garden_id)
        if not garden:
            return jsonify({"error": "Garden not found."}), 404
        if garden.user_id != current_user.id:
            return jsonify({"error": "You don't have permission to add plants to this garden."}), 403

        # Convert date string to date object
        try:
            from datetime import datetime
            planting_date = datetime.strptime(planting_date, '%Y-%m-%d').date()
        except ValueError:
            return jsonify({"error": "Invalid planting_date format. Use YYYY-MM-DD."}), 400

        # Convert frequency fields to integers if provided
        if watering_frequency:
            try:
                watering_frequency = int(watering_frequency)
            except (ValueError, TypeError):
                watering_frequency = None
        
        if fertilizing_frequency:
            try:
                fertilizing_frequency = int(fertilizing_frequency)
            except (ValueError, TypeError):
                fertilizing_frequency = None
        
        if pruning_frequency:
            try:
                pruning_frequency = int(pruning_frequency)
            except (ValueError, TypeError):
                pruning_frequency = None

        # Default name to type if not provided
        if not name:
            name = str(type_).strip().title()
        
        # Create plant record
        plant = Plant(
            name=name, 
            type=type_, 
            environment=environment, 
            care_guide=care_guide, 
            ideal_soil_type=ideal_soil_type, 
            watering_frequency=watering_frequency, 
            fertilizing_frequency=fertilizing_frequency, 
            pruning_frequency=pruning_frequency,
            image_path=image_path
        )
        db.session.add(plant)
        db.session.flush()  # Flush to get the plant ID without committing
        
        # Create plant tracking record
        tracking = PlantTracking(
            garden_id=garden_id, 
            plant_id=plant.id, 
            planting_date=planting_date
        )
        db.session.add(tracking)
        db.session.commit()
        
        return jsonify({"message": "Plant added successfully!", "plant_id": plant.id})
        
    except Exception as e:
        db.session.rollback()  # Rollback any partial changes
        print(f"Error adding plant: {str(e)}")
        return jsonify({"error": f"Failed to add plant: {str(e)}"}), 500

@views.route('/plant/edit/<int:tracking_id>', methods=['POST'])
@login_required
def edit_plant(tracking_id):
    tracking = PlantTracking.query.get_or_404(tracking_id)
    garden = Garden.query.get(tracking.garden_id)
    if garden.user_id != current_user.id:
        return jsonify({"error": "Unauthorized."}), 403
    
    data = request.get_json() or {}
    plant = Plant.query.get(tracking.plant_id)
    
    # Enforce update limit: 3 free total; can purchase additional credits
    usage = _get_or_create_update_usage(current_user.id)
    free_used = int(usage.free_updates_used or 0)
    paid_credits = int(usage.purchased_credits or 0)
    free_remaining = max(0, 3 - free_used)
    total_remaining = free_remaining + paid_credits
    
    if total_remaining <= 0:
        return jsonify({
            "error": "Update limit reached. Purchase more updates to continue.",
            "purchase_required": True,
            "price_per_update": 20
        }), 402
    
    # Track if any changes are actually being made
    any_changes = False
    if 'name' in data and data.get('name'):
        plant.name = data.get('name')
        any_changes = True
    if 'type' in data and data.get('type'):
        plant.type = data.get('type')
        any_changes = True
    if 'environment' in data and data.get('environment'):
        plant.environment = data.get('environment')
        any_changes = True
    if 'care_guide' in data and data.get('care_guide'):
        plant.care_guide = data.get('care_guide')
        any_changes = True
    if 'ideal_soil_type' in data:
        plant.ideal_soil_type = data.get('ideal_soil_type')
        any_changes = True
    if 'watering_frequency' in data:
        plant.watering_frequency = data.get('watering_frequency')
        any_changes = True
    if 'fertilizing_frequency' in data:
        plant.fertilizing_frequency = data.get('fertilizing_frequency')
        any_changes = True
    if 'pruning_frequency' in data:
        plant.pruning_frequency = data.get('pruning_frequency')
        any_changes = True
    if 'planting_date' in data and data.get('planting_date'):
        tracking.planting_date = data.get('planting_date')
        any_changes = True
    
    if not any_changes:
        return jsonify({"message": "No changes detected. Nothing to update."})
    
    # Deduct from free quota first, then from paid credits
    if free_remaining > 0:
        usage.free_updates_used = free_used + 1
    else:
        usage.purchased_credits = max(0, paid_credits - 1)
    
    db.session.commit()
    
    remaining_after = usage.total_remaining()
    return jsonify({
        "message": "Plant updated successfully!",
        "remaining_updates": remaining_after,
        "price_per_update": 20
    })

@views.route('/plant/delete/<int:tracking_id>', methods=['POST'])
@login_required
def delete_plant(tracking_id):
    tracking = PlantTracking.query.get_or_404(tracking_id)
    garden = Garden.query.get(tracking.garden_id)
    if garden.user_id != current_user.id:
        return jsonify({"error": "Unauthorized."}), 403
    
    plant_id = tracking.plant_id
    
    # Remove plant from all grid spaces
    grid_spaces = GridSpace.query.filter_by(plant_id=plant_id).all()
    for space in grid_spaces:
        space.plant_id = None
        space.planting_date = None
        space.last_watered = None
        space.last_fertilized = None
        space.last_pruned = None
        space.notes = None
        space.image_path = None
        space.care_suggestions = None
        space.last_updated = None
    
    # Delete the plant record
    plant = Plant.query.get(plant_id)
    if plant:
        db.session.delete(plant)
    
    # Delete the tracking record
    db.session.delete(tracking)
    
    db.session.commit()
    
    return jsonify({"message": "Plant deleted successfully!"})

# Grid Spaces API endpoints
@views.route('/garden/<int:garden_id>/grid-spaces', methods=['GET'])
@login_required
def get_grid_spaces(garden_id):
    """Get all grid spaces for a specific garden"""
    garden = Garden.query.get_or_404(garden_id)
    if garden.user_id != current_user.id:
        return jsonify({"error": "Unauthorized."}), 403
    
    try:
        grid_spaces = GridSpace.query.filter_by(garden_id=garden_id, is_active=True).all()
        spaces_data = []
        
        for space in grid_spaces:
            # Parse care suggestions if available
            care_suggestions = None
            if space.care_suggestions:
                try:
                    import json
                    care_suggestions = json.loads(space.care_suggestions)
                except:
                    care_suggestions = None
            
            spaces_data.append({
                'id': space.id,
                'garden_id': space.garden_id,
                'grid_position': space.grid_position,
                'plant_id': space.plant_id,
                'planting_date': space.planting_date.isoformat() if space.planting_date else None,
                'last_watered': space.last_watered.isoformat() if space.last_watered else None,
                'last_fertilized': space.last_fertilized.isoformat() if space.last_fertilized else None,
                'last_pruned': space.last_pruned.isoformat() if space.last_pruned else None,
                'notes': space.notes,
                'image_path': space.image_path,
                'care_suggestions': care_suggestions,
                'last_updated': space.last_updated.isoformat() if space.last_updated else None,
                'is_active': space.is_active
            })
        
        return jsonify({"grid_spaces": spaces_data})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@views.route('/api/ai-analysis/usage', methods=['GET'])
@login_required
def get_ai_usage_status():
    """Get current AI analysis usage status for the user (plant analyses only)"""
    try:
        is_premium = getattr(current_user, 'subscribed', False)
        usage = _get_or_create_ai_usage(current_user.id)
        # Premium users get 10 free analyses, basic users get 3
        free_allocation = 10 if is_premium else 3
        
        free_used = usage.free_analyses_used or 0
        free_remaining = max(0, free_allocation - free_used)
        purchased_credits = usage.purchased_credits or 0
        total_remaining = usage.total_remaining(free_allocation)
        
        return jsonify({
            "success": True,
            "is_premium": is_premium,
            "free_allocation": free_allocation,
            "free_used": free_used,
            "free_remaining": free_remaining,
            "purchased_credits": purchased_credits,
            "total_remaining": total_remaining,
            "price_per_analysis": 20.00
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@views.route('/api/soil-analysis/usage', methods=['GET'])
@login_required
def get_soil_usage_status():
    """Get current soil analysis usage status for the user (separate from plant analysis)"""
    try:
        is_premium = getattr(current_user, 'subscribed', False)
        usage = _get_or_create_soil_usage(current_user.id)
        # Premium users get 10 free analyses, basic users get 3
        free_allocation = 10 if is_premium else 3
        
        free_used = usage.free_analyses_used or 0
        free_remaining = max(0, free_allocation - free_used)
        purchased_credits = usage.purchased_credits or 0
        total_remaining = usage.total_remaining(free_allocation)
        
        return jsonify({
            "success": True,
            "is_premium": is_premium,
            "free_allocation": free_allocation,
            "free_used": free_used,
            "free_remaining": free_remaining,
            "purchased_credits": purchased_credits,
            "total_remaining": total_remaining,
            "price_per_analysis": 20.00
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@views.route('/api/ai-analysis/purchase', methods=['POST'])
@login_required
def purchase_ai_analysis():
    """Purchase a single AI analysis for ₱20"""
    try:
        data = request.get_json() or {}
        quantity = data.get('quantity', 1)
        payment_method = data.get('payment_method', 'demo')
        gcash_number = data.get('gcash_number')
        
        if quantity <= 0:
            return jsonify({"success": False, "error": "Invalid quantity"}), 400
        
        # Premium users can also purchase additional credits if they use up their 10 free tries
        # No need to block them from purchasing
        
        # Validate GCash number if GCash payment method is selected
        if payment_method == 'gcash' and not gcash_number:
            return jsonify({"success": False, "error": "GCash mobile number is required"}), 400
        
        # For demo purposes, simulate successful payment
        payment_info = f"via {payment_method}"
        if payment_method == 'gcash' and gcash_number:
            payment_info += f" (GCash: {gcash_number})"
        print(f"💰 AI ANALYSIS PURCHASE: User {current_user.id} purchasing {quantity} analysis(es) {payment_info}")
        
        # Update usage with purchased credits
        usage = _get_or_create_ai_usage(current_user.id)
        usage.purchased_credits = (usage.purchased_credits or 0) + quantity
        db.session.commit()
        
        # Calculate total remaining with correct free allocation
        is_premium = getattr(current_user, 'subscribed', False)
        free_allocation = 10 if is_premium else 3
        total_remaining = usage.total_remaining(free_allocation)
        
        return jsonify({
            "success": True,
            "message": f"Successfully purchased {quantity} analysis(es).",
            "credits_added": quantity,
            "total_remaining": total_remaining,
            "price_per_analysis": 20.00,
            "total_paid": quantity * 20.00
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "error": f"Purchase failed: {str(e)}"}), 500

@views.route('/api/soil-analysis/purchase', methods=['POST'])
@login_required
def purchase_soil_analysis():
    """Purchase a single soil analysis for ₱20"""
    try:
        data = request.get_json() or {}
        quantity = data.get('quantity', 1)
        payment_method = data.get('payment_method', 'demo')
        gcash_number = data.get('gcash_number')
        
        if quantity <= 0:
            return jsonify({"success": False, "error": "Invalid quantity"}), 400
        
        # Premium users can also purchase additional credits if they use up their 10 free tries
        # No need to block them from purchasing
        
        # Validate GCash number if GCash payment method is selected
        if payment_method == 'gcash' and not gcash_number:
            return jsonify({"success": False, "error": "GCash mobile number is required"}), 400
        
        # For demo purposes, simulate successful payment
        payment_info = f"via {payment_method}"
        if payment_method == 'gcash' and gcash_number:
            payment_info += f" (GCash: {gcash_number})"
        print(f"💰 SOIL ANALYSIS PURCHASE: User {current_user.id} purchasing {quantity} analysis(es) {payment_info}")
        
        # Update usage with purchased credits
        usage = _get_or_create_soil_usage(current_user.id)
        usage.purchased_credits = (usage.purchased_credits or 0) + quantity
        db.session.commit()
        
        # Calculate total remaining with correct free allocation
        is_premium = getattr(current_user, 'subscribed', False)
        free_allocation = 10 if is_premium else 3
        total_remaining = usage.total_remaining(free_allocation)
        
        return jsonify({
            "success": True,
            "message": f"Successfully purchased {quantity} soil analysis(es).",
            "credits_added": quantity,
            "total_remaining": total_remaining,
            "price_per_analysis": 20.00,
            "total_paid": quantity * 20.00
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "error": f"Purchase failed: {str(e)}"}), 500

@views.route('/plant/purchase-updates', methods=['POST'])
@login_required
def purchase_plant_updates():
    """Simulate purchase of additional plant update credits at ₱20 per update."""
    try:
        data = request.get_json(force=True, silent=True) or {}
        amount = data.get('amount')  # total amount in PHP
        quantity = data.get('quantity')  # number of update credits to purchase
        transaction_id = data.get('transaction_id')
        
        # Determine quantity: prefer explicit quantity; else derive from amount
        if quantity is None:
            try:
                quantity = int(float(amount) / 20) if amount is not None else 0
            except Exception:
                quantity = 0
        try:
            quantity = int(quantity)
        except Exception:
            quantity = 0
        
        if quantity <= 0:
            return jsonify({"success": False, "error": "Invalid purchase quantity"}), 400
        
        # Simulate successful payment
        print(f"💰 Plant update credits purchase by User {current_user.id}: {quantity} credits (Txn: {transaction_id})")
        
        usage = _get_or_create_update_usage(current_user.id)
        usage.purchased_credits = int(usage.purchased_credits or 0) + quantity
        db.session.commit()
        
        total_remaining = usage.total_remaining()
        
        return jsonify({
            "success": True,
            "message": f"Purchased {quantity} update credit(s) successfully.",
            "credits_added": quantity,
            "remaining_updates": total_remaining,
            "price_per_update": 20
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "error": f"Purchase failed: {str(e)}"}), 500

@views.route('/garden/verify-payment', methods=['POST'])
@login_required
def verify_payment():
    """Verify GCash payment and add purchased spaces to garden"""
    try:
        data = request.get_json(force=True, silent=True) or {}
        garden_id = data.get('garden_id')
        spaces_purchased = data.get('spaces_purchased', 0)
        transaction_id = data.get('transaction_id')
        
        if not garden_id or spaces_purchased <= 0:
            return jsonify({"success": False, "error": "Invalid payment data"}), 400
        
        # Get the garden
        garden = Garden.query.get_or_404(garden_id)
        if garden.user_id != current_user.id:
            return jsonify({"success": False, "error": "Unauthorized"}), 403
        
        # For demo purposes, we'll simulate successful payment verification
        # In production, this would verify with GCash API
        print(f"💰 Payment verification for Garden {garden_id}: {spaces_purchased} spaces")
        print(f"💰 Transaction ID: {transaction_id}")
        
        # Update garden with additional spaces
        current_additional = getattr(garden, 'additional_spaces_purchased', 0)
        garden.additional_spaces_purchased = current_additional + spaces_purchased
        
        # Create additional grid spaces for the purchased spaces
        create_additional_grid_spaces(garden.id, spaces_purchased, garden.grid_size)
        
        db.session.commit()
        
        return jsonify({
            "success": True, 
            "message": f"Payment verified! {spaces_purchased} additional spaces added to your garden.",
            "total_spaces": getattr(garden, 'base_grid_spaces', 9) + garden.additional_spaces_purchased
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"Error verifying payment: {str(e)}")
        return jsonify({"success": False, "error": f"Payment verification failed: {str(e)}"}), 500

@views.route('/garden/place-plant', methods=['POST'])
@login_required
def place_plant():
    """Place a plant in a specific grid space"""
    try:
        data = request.get_json()
        space_id = data.get('space_id')
        plant_id = data.get('plant_id')
        planting_date = data.get('planting_date')
        notes = data.get('notes', '')
        
        if not space_id or not plant_id or not planting_date:
            return jsonify({"error": "Missing required fields."}), 400
        
        # Get the grid space
        grid_space = GridSpace.query.get_or_404(space_id)
        garden = Garden.query.get(grid_space.garden_id)
        
        # Check authorization
        if garden.user_id != current_user.id:
            return jsonify({"error": "Unauthorized."}), 403
        
        # Check if space is already occupied
        if grid_space.plant_id:
            return jsonify({"error": "This grid space is already occupied."}), 400
        
        # Update the grid space with plant information
        grid_space.plant_id = plant_id
        grid_space.planting_date = planting_date
        grid_space.notes = notes
        
        db.session.commit()
        
        return jsonify({"message": "Plant placed successfully!"})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@views.route('/garden/remove-plant/<int:space_id>', methods=['POST'])
@login_required
def remove_plant(space_id):
    """Remove a plant from a specific grid space"""
    try:
        grid_space = GridSpace.query.get_or_404(space_id)
        garden = Garden.query.get(grid_space.garden_id)
        
        # Check authorization
        if garden.user_id != current_user.id:
            return jsonify({"error": "Unauthorized."}), 403
        
        # Remove plant from grid space
        grid_space.plant_id = None
        grid_space.planting_date = None
        grid_space.notes = ''
        
        db.session.commit()
        
        return jsonify({"message": "Plant removed successfully!"})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@views.route('/garden/upload-plant-image', methods=['POST'])
@login_required
def upload_plant_image():
    """Upload an image for a planted space"""
    try:
        print("🔄 Starting image upload process...")
        
        space_id = request.form.get('space_id')
        print(f"📋 Space ID: {space_id}")
        if not space_id:
            return jsonify({"error": "Space ID is required."}), 400
        
        grid_space = GridSpace.query.get_or_404(space_id)
        garden = Garden.query.get(grid_space.garden_id)
        if garden.user_id != current_user.id:
            return jsonify({"error": "Unauthorized."}), 403
        
        file = request.files.get('image')
        print(f"📁 File received: {file}")
        if not file:
            return jsonify({"error": "Image file is required."}), 400
        
        # Create uploads directory if it doesn't exist
        upload_dir = os.path.join(os.getcwd(), 'uploads', 'plants')
        print(f"📁 Upload directory: {upload_dir}")
        os.makedirs(upload_dir, exist_ok=True)
        
        # Generate unique filename with proper extension
        import time
        # Get file extension from content type or default to jpg
        content_type = file.content_type
        if content_type and 'image/' in content_type:
            ext = content_type.split('/')[-1]
            if ext not in ['jpg', 'jpeg', 'png', 'webp', 'gif']:
                ext = 'jpg'
        else:
            ext = 'jpg'
        
        filename = f"space_{space_id}_{int(time.time())}.{ext}"
        file_path = os.path.join(upload_dir, filename)
        print(f"💾 File path: {file_path}")
        
        # Save the file
        print("💾 Saving file...")
        file.save(file_path)
        print("✅ File saved successfully!")
        
        # Update grid space with image path
        grid_space.image_path = f"uploads/plants/{filename}"
        print(f"🗄️ Updated image path: {grid_space.image_path}")
        
        # AI Analysis for care suggestions
        import json  # Import json at the beginning of the function
        care_suggestions = {
            "needs_water": False,
            "needs_fertilize": False,
            "needs_prune": False,
            "confidence": 0.0,
            "reasoning": "AI analysis not available"
        }
        
        try:
            openai_key = os.getenv('OPENAI_API_KEY')
            if openai_key:
                # Check usage limit before performing AI analysis (premium users get 10 free tries, basic users get 3)
                is_premium = getattr(current_user, 'subscribed', False)
                can_proceed, remaining, needs_payment = _check_ai_usage_limit(current_user.id, is_premium)
                
                if not can_proceed:
                    # Still save the image but without AI analysis
                    grid_space.care_suggestions = json.dumps(care_suggestions)
                    db.session.commit()
                    return jsonify({
                        "error": "Free analysis limit reached. Please purchase additional analyses or subscribe to Premium.",
                        "limit_reached": True,
                        "needs_payment": True,
                        "remaining": 0,
                        "price_per_analysis": 20.00,
                        "image_path": grid_space.image_path,
                        "care_suggestions": care_suggestions
                    }), 402  # 402 Payment Required
                
                print("🤖 Starting AI analysis...")
                
                # Read and base64 encode image for OpenAI Vision
                with open(file_path, 'rb') as img_file:
                    image_bytes = img_file.read()
                    image_b64 = base64.b64encode(image_bytes).decode('utf-8')
                
                # Get plant name for context
                plant_name = "plant"
                if grid_space.plant_id:
                    plant = Plant.query.get(grid_space.plant_id)
                    if plant:
                        plant_name = plant.name
                
                print(f"🌱 Analyzing {plant_name} plant...")
                
                # OpenAI Vision analysis
                try:
                    from openai import OpenAI
                    # Clear any proxy environment variables that might cause conflicts
                    proxy_vars = ['HTTP_PROXY', 'HTTPS_PROXY', 'http_proxy', 'https_proxy', 'ALL_PROXY', 'all_proxy']
                    original_proxy_values = {}
                    for var in proxy_vars:
                        if var in os.environ:
                            original_proxy_values[var] = os.environ[var]
                            del os.environ[var]
                    
                    # Initialize client with only the api_key to avoid proxy conflicts
                    client = OpenAI(api_key=openai_key)
                    
                    # Restore proxy environment variables
                    for var, value in original_proxy_values.items():
                        os.environ[var] = value
                    print(f"✅ OpenAI client initialized successfully")
                    
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": f"Analyze this {plant_name} plant image for care needs. You MUST respond with ONLY valid JSON in this exact format: {{\"needs_water\": true, \"needs_fertilize\": false, \"needs_prune\": false, \"confidence\": 0.8, \"reasoning\": \"description of what you see\"}}. \n\nSet needs_water=true if you see: wilted leaves, drooping, dry soil, brown leaf edges, dehydration signs, or underwatering symptoms.\n\nSet needs_fertilize=true if you see: yellowing leaves, stunted growth, pale foliage, or nutrient deficiency signs.\n\nSet needs_prune=true if you see: dead branches, damaged leaves, overgrowth, brown/shriveled leaves, or shape issues.\n\nYou can set MULTIPLE needs to true if the plant needs multiple types of care. Be very specific about what you see in the image."
                                    },
                                    {
                                        "type": "image_url",
                                        "image_url": {
                                            "url": f"data:image/jpeg;base64,{image_b64}"
                                        }
                                    }
                                ]
                            }
                        ],
                        max_tokens=300
                    )
                    
                    # Parse AI response
                    ai_response = response.choices[0].message.content
                    print(f"🤖 AI Response: {ai_response}")
                    print(f"🌱 Plant being analyzed: {plant_name}")
                    print(f"🔍 Looking for care needs in AI response...")
                    
                    try:
                        # Clean the response before parsing
                        cleaned_response = ai_response.strip()
                        if cleaned_response.startswith('```json'):
                            cleaned_response = cleaned_response.replace('```json', '').replace('```', '').strip()
                        elif cleaned_response.startswith('```'):
                            cleaned_response = cleaned_response.replace('```', '').strip()
                        
                        suggestions = json.loads(cleaned_response)
                        care_suggestions = suggestions
                        print(f"✅ Parsed AI analysis: {care_suggestions}")
                        print(f"💧 Needs water: {care_suggestions.get('needs_water', False)}")
                        print(f"🌱 Needs fertilize: {care_suggestions.get('needs_fertilize', False)}")
                        print(f"✂️ Needs prune: {care_suggestions.get('needs_prune', False)}")
                        # Clean up the reasoning text if it contains JSON formatting
                        if 'reasoning' in care_suggestions and isinstance(care_suggestions['reasoning'], str):
                            reasoning = care_suggestions['reasoning']
                            if '```json' in reasoning or '```' in reasoning:
                                # Extract clean text from reasoning
                                clean_reasoning = reasoning.replace('```json', '').replace('```', '').strip()
                                if clean_reasoning.startswith('{'):
                                    try:
                                        json_data = json.loads(clean_reasoning)
                                        care_suggestions['reasoning'] = json_data.get('reasoning', reasoning)
                                    except:
                                        care_suggestions['reasoning'] = reasoning
                                else:
                                    care_suggestions['reasoning'] = clean_reasoning
                        print(f"✅ AI Analysis completed: {care_suggestions}")
                        
                        # Track AI usage after successful analysis
                        free_allocation = 10 if is_premium else 3
                        usage = _get_or_create_ai_usage(current_user.id)
                        is_free = usage.can_use_free(free_allocation)
                        old_free_used = usage.free_analyses_used or 0
                        old_purchased = usage.purchased_credits or 0
                        _track_ai_usage(
                            current_user.id,
                            'plant_analysis',
                            is_free=is_free,
                            cost=0.00 if is_free else 20.00,
                            image_path=grid_space.image_path,
                            analysis_result=str(care_suggestions.get('reasoning', ''))[:500]
                        )
                        db.session.flush()  # Ensure usage is saved
                        db.session.refresh(usage)
                        new_free_used = usage.free_analyses_used or 0
                        new_purchased = usage.purchased_credits or 0
                        print(f"✅ Tracked usage for user {current_user.id} (premium={is_premium}, allocation={free_allocation}): free {old_free_used}→{new_free_used}, purchased {old_purchased}→{new_purchased}")
                        
                        # Enhanced health issue detection - check for negative statements first
                        negative_indicators = ['no signs of', 'no visible', 'no evidence of', 'appears healthy', 'looks healthy', 'no mold', 'no rot', 'no spoilage', 'no disease', 'no problems', 'no issues']
                        has_negative_indicators = any(indicator in ai_response.lower() for indicator in negative_indicators)
                        
                        health_issue_phrases = [
                            'mold detected', 'rot detected', 'disease detected', 'spoilage detected',
                            'fungal infection', 'powdery mildew', 'white coating', 'gray coating', 
                            'green coating', 'shriveled fruit', 'wrinkled fruit', 'damaged fruit',
                            'discolored fruit', 'brown spots', 'black spots', 'soft spots',
                            'mushy fruit', 'decay present', 'health issues', 'plant problems',
                            'needs treatment', 'urgent care', 'immediate attention'
                        ]
                        detected_issues = [phrase for phrase in health_issue_phrases if phrase in ai_response.lower()]
                        
                        # Only override if health issues detected AND no negative indicators
                        if detected_issues and not has_negative_indicators:
                            print(f"🚨 AI detected health issues: {detected_issues}")
                            print(f"🚨 Full AI response: {ai_response}")
                            if not care_suggestions.get('needs_water', False):
                                print(f"⚠️ WARNING: AI detected health issues but didn't set needs_water to true!")
                                care_suggestions['needs_water'] = True
                                care_suggestions['reasoning'] = f"URGENT: Health issues detected ({', '.join(detected_issues)}) - {care_suggestions.get('reasoning', '')}"
                                care_suggestions['confidence'] = max(care_suggestions.get('confidence', 0.5), 0.8)  # Increase confidence for health issues
                        elif has_negative_indicators:
                            print(f"✅ AI indicates healthy plant: {ai_response}")
                            # Override to healthy if AI explicitly says no problems
                            care_suggestions['needs_water'] = False
                            # Extract clean reasoning from AI response
                            clean_reasoning = ai_response.replace('```json', '').replace('```', '').strip()
                            if clean_reasoning.startswith('{'):
                                # If it's JSON, extract just the reasoning field
                                try:
                                    json_data = json.loads(clean_reasoning)
                                    clean_reasoning = json_data.get('reasoning', 'Plant appears healthy')
                                except:
                                    clean_reasoning = 'Plant appears healthy'
                            care_suggestions['reasoning'] = clean_reasoning[:200] + ('...' if len(clean_reasoning) > 200 else '')
                            care_suggestions['confidence'] = max(care_suggestions.get('confidence', 0.5), 0.8)
                    except Exception as parse_error:
                        print(f"⚠️ JSON parsing failed: {parse_error}")
                        print(f"🔍 Raw AI response: {ai_response}")
                        
                        # Enhanced fallback with negative statement detection
                        negative_indicators = ['no signs of', 'no visible', 'no evidence of', 'appears healthy', 'looks healthy', 'no mold', 'no rot', 'no spoilage', 'no disease', 'no problems', 'no issues']
                        has_negative_indicators = any(indicator in ai_response.lower() for indicator in negative_indicators)
                        
                        health_issue_phrases = [
                            'mold detected', 'rot detected', 'disease detected', 'spoilage detected',
                            'fungal infection', 'powdery mildew', 'white coating', 'gray coating', 
                            'green coating', 'shriveled fruit', 'wrinkled fruit', 'damaged fruit',
                            'discolored fruit', 'brown spots', 'black spots', 'soft spots',
                            'mushy fruit', 'decay present', 'health issues', 'plant problems',
                            'needs treatment', 'urgent care', 'immediate attention'
                        ]
                        detected_issues = [phrase for phrase in health_issue_phrases if phrase in ai_response.lower()]
                        
                        # Also check for positive health indicators
                        healthy_indicators = ['healthy', 'fresh', 'ripe', 'good condition', 'no problems', 'no issues', 'looks good', 'appears healthy']
                        healthy_detected = [indicator for indicator in healthy_indicators if indicator in ai_response.lower()]
                        
                        # Priority: negative indicators override everything
                        if has_negative_indicators:
                            print(f"✅ Fallback: Detected negative indicators (healthy): {[ind for ind in negative_indicators if ind in ai_response.lower()]}")
                            # Extract clean reasoning from AI response
                            clean_reasoning = ai_response.replace('```json', '').replace('```', '').strip()
                            if clean_reasoning.startswith('{'):
                                # If it's JSON, extract just the reasoning field
                                try:
                                    json_data = json.loads(clean_reasoning)
                                    clean_reasoning = json_data.get('reasoning', 'Plant appears healthy')
                                except:
                                    clean_reasoning = 'Plant appears healthy'
                            care_suggestions = {
                                "needs_water": False,
                                "needs_fertilize": False, 
                                "needs_prune": False,
                                "confidence": 0.8,
                                "reasoning": clean_reasoning[:200] + ('...' if len(clean_reasoning) > 200 else '')
                            }
                        elif detected_issues and not healthy_detected:
                            print(f"🚨 Fallback: Detected health issues in response: {detected_issues}")
                            care_suggestions = {
                                "needs_water": True,
                                "needs_fertilize": False, 
                                "needs_prune": False,
                                "confidence": 0.8,
                                "reasoning": f"URGENT: Health issues detected ({', '.join(detected_issues)}) - {ai_response[:200]}..."
                            }
                        elif healthy_detected:
                            print(f"✅ Fallback: Detected healthy indicators: {healthy_detected}")
                            # Extract clean reasoning from AI response
                            clean_reasoning = ai_response.replace('```json', '').replace('```', '').strip()
                            if clean_reasoning.startswith('{'):
                                # If it's JSON, extract just the reasoning field
                                try:
                                    json_data = json.loads(clean_reasoning)
                                    clean_reasoning = json_data.get('reasoning', 'Plant appears healthy')
                                except:
                                    clean_reasoning = 'Plant appears healthy'
                            care_suggestions = {
                                "needs_water": False,
                                "needs_fertilize": False, 
                                "needs_prune": False,
                                "confidence": 0.7,
                                "reasoning": clean_reasoning[:200] + ('...' if len(clean_reasoning) > 200 else '')
                            }
                        else:
                            care_suggestions = {
                                "needs_water": False,
                                "needs_fertilize": False, 
                                "needs_prune": False,
                                "confidence": 0.3,
                                "reasoning": f"AI analysis completed but response format unclear: {ai_response[:100]}..."
                            }
                        
                except Exception as openai_error:
                    print(f"❌ OpenAI client error: {openai_error}")
                    print(f"❌ Error type: {type(openai_error).__name__}")
                    
                    # Check if it's the proxies error specifically
                    if "proxies" in str(openai_error):
                        print("🔧 Detected proxies error - trying direct HTTP API call")
                        try:
                            # Try direct HTTP API call as fallback
                            import requests
                            headers = {
                                'Authorization': f'Bearer {openai_key}',
                                'Content-Type': 'application/json'
                            }
                            payload = {
                                "model": "gpt-4o",
                                "messages": [
                                    {
                                        "role": "user",
                                        "content": [
                                            {
                                                "type": "text",
                                                "text": f"Analyze this {plant_name} plant image for health issues. Look for mold, rot, disease, spoilage, or any problems. You MUST respond with ONLY valid JSON in this exact format: {{\"needs_water\": true, \"needs_fertilize\": false, \"needs_prune\": false, \"confidence\": 0.8, \"reasoning\": \"description of what you see\"}}. If you see mold, rot, spoilage, or any health issues, set needs_water to true. If plant looks healthy, set needs_water to false. Be very specific about what you see in the image."
                                            },
                                            {
                                                "type": "image_url",
                                                "image_url": {
                                                    "url": f"data:image/jpeg;base64,{image_b64}"
                                                }
                                            }
                                        ]
                                    }
                                ],
                                "max_tokens": 300
                            }
                            
                            response = requests.post(
                                'https://api.openai.com/v1/chat/completions',
                                headers=headers,
                                json=payload,
                                timeout=30
                            )
                            
                            if response.status_code == 200:
                                ai_response = response.json()['choices'][0]['message']['content']
                                print(f"✅ Direct HTTP API call successful: {ai_response}")
                                
                                # Parse the response
                                try:
                                    cleaned_response = ai_response.strip()
                                    if cleaned_response.startswith('```json'):
                                        cleaned_response = cleaned_response.replace('```json', '').replace('```', '').strip()
                                    elif cleaned_response.startswith('```'):
                                        cleaned_response = cleaned_response.replace('```', '').strip()
                                    
                                    care_suggestions = json.loads(cleaned_response)
                                    
                                    # Track AI usage after successful analysis (fallback HTTP method)
                                    free_allocation = 10 if is_premium else 3
                                    usage = _get_or_create_ai_usage(current_user.id)
                                    is_free = usage.can_use_free(free_allocation)
                                    old_free_used = usage.free_analyses_used or 0
                                    old_purchased = usage.purchased_credits or 0
                                    _track_ai_usage(
                                        current_user.id,
                                        'plant_analysis',
                                        is_free=is_free,
                                        cost=0.00 if is_free else 20.00,
                                        image_path=grid_space.image_path,
                                        analysis_result=str(care_suggestions.get('reasoning', ''))[:500]
                                    )
                                    db.session.flush()
                                    db.session.refresh(usage)
                                    new_free_used = usage.free_analyses_used or 0
                                    new_purchased = usage.purchased_credits or 0
                                    print(f"✅ Tracked usage (HTTP fallback) for user {current_user.id}: free {old_free_used}→{new_free_used}, purchased {old_purchased}→{new_purchased}")
                                except:
                                    care_suggestions = {
                                        "needs_water": False,
                                        "needs_fertilize": False, 
                                        "needs_prune": False,
                                        "confidence": 0.8,
                                        "reasoning": "Plant appears healthy based on visual analysis"
                                    }
                            else:
                                raise Exception(f"HTTP API call failed: {response.status_code}")
                                
                        except Exception as http_error:
                            print(f"❌ Direct HTTP API call also failed: {http_error}")
                            care_suggestions = {
                                "needs_water": False,
                                "needs_fertilize": False, 
                                "needs_prune": False,
                                "confidence": 0.0,
                                "reasoning": "AI analysis failed: Client initialization error (proxies parameter not supported)"
                            }
                    else:
                        care_suggestions = {
                            "needs_water": False,
                            "needs_fertilize": False, 
                            "needs_prune": False,
                            "confidence": 0.0,
                            "reasoning": f"AI analysis failed: {str(openai_error)}"
                        }
            else:
                print("⚠️ OpenAI API key not configured")
                care_suggestions = {
                    "needs_water": False,
                    "needs_fertilize": False,
                    "needs_prune": False, 
                    "confidence": 0.0,
                    "reasoning": "OpenAI API key not configured"
                }
        except Exception as ai_error:
            print(f"❌ AI analysis error: {ai_error}")
            care_suggestions = {
                "needs_water": False,
                "needs_fertilize": False,
                "needs_prune": False,
                "confidence": 0.0,
                "reasoning": f"AI analysis failed: {str(ai_error)}"
            }
        
        # Determine health status from care_suggestions
        def determine_health_status(care_suggestions_dict):
            """Determine health status from care suggestions"""
            if not care_suggestions_dict:
                return 'Unknown'
            
            reasoning = care_suggestions_dict.get('reasoning', '').lower()
            needs_water = care_suggestions_dict.get('needs_water', False)
            needs_fertilize = care_suggestions_dict.get('needs_fertilize', False)
            needs_prune = care_suggestions_dict.get('needs_prune', False)
            
            # Check for explicit healthy indicators FIRST (highest priority)
            # These are strong positive indicators that override everything else
            strong_healthy_indicators = [
                'appears healthy', 'looks healthy', 'healthy and', 'healthy and green',
                'healthy and vibrant', 'no visible signs of', 'no visible signs',
                'no clear signs of', 'no clear signs', 'no signs of dehydration',
                'no signs of nutrient deficiency', 'no signs of disease',
                'ripe and abundant', 'abundant', 'thriving', 'flourishing',
                'in good health', 'good condition', 'excellent condition'
            ]
            
            has_strong_healthy = any(indicator in reasoning for indicator in strong_healthy_indicators)
            
            # Check for negative health indicators (but only if not explicitly healthy)
            urgent_indicators = [
                'urgent', 'health issues detected', 'mold detected', 'rot detected', 
                'disease detected', 'spoilage detected', 'fungal infection', 
                'powdery mildew', 'immediate attention', 'critical'
            ]
            
            # Check for negative context phrases (e.g., "no signs of", "no visible")
            # If these are present, we should NOT count damage indicators
            negative_context_phrases = [
                'no visible signs of', 'no visible signs', 'no clear signs of',
                'no clear signs', 'no signs of', 'no evidence of', 'no indication of',
                'no sign of', 'no symptoms of', 'no trace of'
            ]
            
            has_negative_context = any(phrase in reasoning for phrase in negative_context_phrases)
            
            # Check for damage indicators (but exclude if mentioned in negative context)
            damage_indicators = [
                'shriveled', 'damaged', 'decay', 'dehydration', 'wilted', 
                'brown spots', 'black spots', 'mushy', 'dying', 'dying leaves'
            ]
            
            # Only check for damage if NOT in negative context
            # For example, "no signs of dehydration" should NOT count as damage
            has_damage = False
            if not has_negative_context:
                has_damage = any(indicator in reasoning for indicator in damage_indicators)
            
            has_urgent_issues = any(indicator in reasoning for indicator in urgent_indicators)
            
            # Additional healthy indicators (weaker, but still positive)
            healthy_indicators = [
                'no problems', 'no issues', 'fresh', 'ripe', 'vibrant',
                'no damage', 'no disease', 'no mold', 'no rot', 'green leaves',
                'healthy leaves', 'leaves appear healthy', 'leaves are healthy'
            ]
            
            has_healthy_indicators = any(indicator in reasoning for indicator in healthy_indicators)
            
            # Priority logic:
            # 1. If explicitly healthy (strong indicators), it's Healthy
            # 2. If urgent issues AND not explicitly healthy, it's Unhealthy
            # 3. If damage indicators (and not in negative context) AND not explicitly healthy, it's Unhealthy
            # 4. If healthy indicators (weaker) AND no urgent/damage, it's Healthy
            # 5. If needs care actions, it's Needs Care
            # 6. Otherwise Unknown
            
            if has_strong_healthy:
                # Explicitly healthy - override everything
                return 'Healthy'
            elif has_urgent_issues:
                # Urgent issues always mean unhealthy
                return 'Unhealthy'
            elif has_damage and not has_strong_healthy:
                # Damage indicators (when not in negative context) mean unhealthy
                return 'Unhealthy'
            elif has_healthy_indicators and not has_urgent_issues and not has_damage:
                # Positive indicators without negative ones
                return 'Healthy'
            elif needs_water or needs_fertilize or needs_prune:
                # Needs care but not explicitly unhealthy
                return 'Needs Care'
            else:
                return 'Unknown'
        
        # Get previous health status
        old_care_suggestions = None
        old_health_status = 'Unknown'
        if grid_space.care_suggestions:
            try:
                old_care_suggestions = json.loads(grid_space.care_suggestions)
                old_health_status = determine_health_status(old_care_suggestions)
            except:
                pass
        
        # Determine new health status
        new_health_status = determine_health_status(care_suggestions)
        
        # Store care suggestions in grid space
        grid_space.care_suggestions = json.dumps(care_suggestions)
        grid_space.ai_analyzed = True
        grid_space.ai_analysis_date = datetime.now(timezone.utc)
        grid_space.ai_analysis_result = json.dumps(care_suggestions)
        grid_space.last_updated = datetime.now(timezone.utc)
        print(f"🗄️ Stored care suggestions: {care_suggestions}")
        print(f"🤖 Marked as AI analyzed: {grid_space.ai_analyzed}")
        print(f"📅 Updated last_updated timestamp")
        
        # Track health status update (both changes and same status)
        # Log all health status updates, including the first one
        from website.models import ActivityLog
        from datetime import date
        
        plant_name = grid_space.plant.name if grid_space.plant else 'Unknown Plant'
        
        # Create description based on whether this is the first update or a change
        if old_health_status == 'Unknown':
            # First update - log the initial health status
            health_change_description = f"Initial health status: {new_health_status}"
        elif old_health_status != new_health_status:
            # Status changed
            health_change_description = f"Health status changed from {old_health_status} to {new_health_status}"
        else:
            # Status remained the same
            health_change_description = f"Health status remains {new_health_status}"
        
        # Add reasoning context if available
        if care_suggestions.get('reasoning'):
            health_change_description += f": {care_suggestions['reasoning'][:150]}"
        
        health_log = ActivityLog(
            user_id=current_user.id,
            garden_id=garden.id,
            space_id=grid_space.id,
            plant_id=grid_space.plant_id,
            action='health_change',
            action_date=date.today(),
            notes=health_change_description
        )
        db.session.add(health_log)
        
        if old_health_status == 'Unknown':
            print(f"📊 Initial health status logged: {new_health_status}")
        elif old_health_status != new_health_status:
            print(f"📊 Health status change logged: {old_health_status} → {new_health_status}")
        else:
            print(f"📊 Health status update logged: {new_health_status} (unchanged)")
        
        print("💾 Committing to database...")
        db.session.commit()
        print("✅ Database commit successful!")
        
        print("🎉 Upload process completed successfully!")
        return jsonify({
            "message": "Image uploaded successfully!", 
            "image_path": grid_space.image_path,
            "care_suggestions": care_suggestions
        })
    except Exception as e:
        print(f"❌ Error in upload process: {str(e)}")
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@views.route('/garden/update-plant-care', methods=['POST'])
@login_required
def update_plant_care():
    """Update plant care activities (water, fertilize, prune)"""
    try:
        data = request.get_json()
        space_id = data.get('space_id')
        action = data.get('action')
        date = data.get('date')
        
        if not space_id or not action:
            return jsonify({"error": "Space ID and action are required."}), 400
        
        grid_space = GridSpace.query.get_or_404(space_id)
        garden = Garden.query.get(grid_space.garden_id)
        if garden.user_id != current_user.id:
            return jsonify({"error": "Unauthorized."}), 403
        
        # Update the appropriate field based on action
        if action == 'water':
            grid_space.last_watered = date
        elif action == 'fertilize':
            grid_space.last_fertilized = date
        elif action == 'prune':
            grid_space.last_pruned = date
        else:
            return jsonify({"error": "Invalid action. Must be 'water', 'fertilize', or 'prune'."}), 400
        
        # Update the last_updated timestamp
        grid_space.last_updated = datetime.now(timezone.utc)
        
        db.session.commit()
        
        return jsonify({"message": f"Plant {action}ed successfully!"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@views.route('/uploads/<path:filename>')
def uploaded_file(filename):
    """Serve uploaded files"""
    return send_from_directory(os.path.join(os.getcwd(), 'uploads'), filename)

@views.route('/uploads/learning_paths/<file_type>/<filename>')
def uploaded_learning_path_file(file_type, filename):
    """Serve uploaded learning path files"""
    upload_dir = os.path.join(os.getcwd(), 'uploads', 'learning_paths', file_type)
    return send_from_directory(upload_dir, filename)

@views.route('/api/admin/delete-file', methods=['POST'])
@login_required
def admin_delete_file():
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        data = request.get_json()
        file_url = data.get('fileUrl')
        
        if not file_url:
            return jsonify({"error": "File URL is required"}), 400
        
        # Extract file path from URL
        if '/uploads/learning_paths/' in file_url:
            # Remove the leading slash and get the relative path
            relative_path = file_url.lstrip('/')
            file_path = os.path.join(os.getcwd(), relative_path)
            
            # Check if file exists and delete it
            if os.path.exists(file_path):
                os.remove(file_path)
                return jsonify({"message": "File deleted successfully"})
            else:
                return jsonify({"error": "File not found"}), 404
        else:
            return jsonify({"error": "Invalid file URL"}), 400
            
    except Exception as e:
        print(f"Error deleting file: {str(e)}")
        return jsonify({"error": f"Delete failed: {str(e)}"}), 500

def get_watering_recommendation(plant, days_since_watered):
    """Generate plant-specific watering recommendations"""
    plant_type = getattr(plant, 'type', '').lower()
    environment = getattr(plant, 'environment', '').lower()
    
    recommendations = []
    
    # Base recommendations by plant type
    if 'tomato' in plant.name.lower():
        recommendations.append("Water at the base, avoid wetting leaves to prevent disease")
        recommendations.append("Use 1-2 inches of water per week")
    elif 'basil' in plant.name.lower() or 'herb' in plant_type:
        recommendations.append("Keep soil consistently moist but not waterlogged")
        recommendations.append("Water when top inch of soil feels dry")
    elif 'succulent' in plant.name.lower() or 'cactus' in plant.name.lower():
        recommendations.append("Water deeply but infrequently - let soil dry between waterings")
        recommendations.append("Use well-draining soil")
    else:
        recommendations.append("Water when top 1-2 inches of soil feels dry")
    
    # Environment-specific recommendations
    if environment == 'indoor':
        recommendations.append("Check for proper drainage to prevent root rot")
    elif environment == 'outdoor':
        recommendations.append("Water early morning for best absorption")
    
    # Overdue recommendations
    if days_since_watered > 7:
        recommendations.append("⚠️ Plant may be stressed - water immediately and check for wilting")
    
    return " | ".join(recommendations)

def get_new_plant_recommendation(plant):
    """Generate recommendations for newly planted plants"""
    plant_type = getattr(plant, 'type', '').lower()
    
    recommendations = []
    recommendations.append("First watering: water thoroughly to help establish roots")
    
    if 'tomato' in plant.name.lower():
        recommendations.append("Add support stake now for easier training later")
    elif 'herb' in plant_type:
        recommendations.append("Pinch back any flowers to encourage leaf growth")
    
    recommendations.append("Monitor for signs of transplant shock")
    
    return " | ".join(recommendations)

def get_fertilizing_recommendation(plant, days_since_fertilized):
    """Generate plant-specific fertilizing recommendations"""
    plant_type = getattr(plant, 'type', '').lower()
    
    recommendations = []
    
    if 'tomato' in plant.name.lower():
        recommendations.append("Use balanced fertilizer (10-10-10) or tomato-specific fertilizer")
        recommendations.append("Apply around base, avoid direct contact with stems")
    elif 'herb' in plant_type:
        recommendations.append("Use diluted liquid fertilizer (half strength)")
        recommendations.append("Organic options: compost tea or fish emulsion")
    else:
        recommendations.append("Use balanced fertilizer according to package instructions")
    
    if days_since_fertilized > 30:
        recommendations.append("⚠️ Long overdue - use gentle fertilizer to avoid shock")
    
    return " | ".join(recommendations)

def get_pruning_recommendation(plant, days_since_pruned):
    """Generate plant-specific pruning recommendations"""
    plant_type = getattr(plant, 'type', '').lower()
    
    recommendations = []
    
    if 'tomato' in plant.name.lower():
        recommendations.append("Remove suckers (small shoots between main stem and branches)")
        recommendations.append("Prune lower leaves that touch the ground")
    elif 'herb' in plant_type:
        recommendations.append("Pinch back tips to encourage bushier growth")
        recommendations.append("Remove any yellow or dead leaves")
    else:
        recommendations.append("Remove dead, damaged, or diseased branches")
        recommendations.append("Prune for shape and air circulation")
    
    return " | ".join(recommendations)

@views.route('/api/smart-alerts')
@login_required
def smart_alerts():
    """Generate smart alerts for user's plants based on care schedules"""
    try:
        alerts = []
        today = datetime.now().date()
        
        print(f"🔍 Generating alerts for user {current_user.id}")
        
        # Get user's gardens
        user_gardens = Garden.query.filter_by(user_id=current_user.id).all()
        print(f"🏡 Found {len(user_gardens)} gardens for user")
        
        for garden in user_gardens:
            # Check GridSpace plants (from GridPlanner) - only plants with care suggestions
            grid_spaces = GridSpace.query.filter_by(garden_id=garden.id).filter(
                GridSpace.plant_id.isnot(None),
                GridSpace.care_suggestions.isnot(None)
            ).all()
            
            for space in grid_spaces:
                if space.plant_id:
                    plant = Plant.query.get(space.plant_id)
                    if not plant:
                        continue
                    
                    # Check AI analysis results first
                    ai_analysis = None
                    if space.care_suggestions:
                        try:
                            ai_analysis = json.loads(space.care_suggestions)
                            print(f"🔍 AI Analysis for {plant.name}: {ai_analysis}")
                        except:
                            ai_analysis = None
                            print(f"❌ Failed to parse AI analysis for {plant.name}")
                    else:
                        print(f"ℹ️ No AI analysis found for {plant.name}")
                    
                    # Generate alerts based on AI analysis - check each care type separately
                    # Only process if AI analysis exists and indicates care needs
                    if ai_analysis:
                        # Check if plant actually needs care
                        needs_water = ai_analysis.get('needs_water', False)
                        needs_fertilize = ai_analysis.get('needs_fertilize', False)
                        needs_prune = ai_analysis.get('needs_prune', False)
                        has_care_needs = needs_water or needs_fertilize or needs_prune
                        
                        # Only proceed if plant needs care
                        if has_care_needs:
                            # Check if this is a newly added plant (to skip initial alerts)
                            # Only skip if: planted today AND no care history AND analyzed within last 10 seconds
                            is_newly_added = False
                            has_care_history = space.last_watered or space.last_fertilized or space.last_pruned
                            
                            # If plant has care history, it's definitely an existing plant
                            if has_care_history:
                                print(f"✅ Showing alerts for existing plant {plant.name} (has care history)")
                                is_newly_added = False
                            # Check if plant was placed today and might be newly added
                            elif space.planting_date:
                                planting_date = space.planting_date
                                if hasattr(planting_date, 'date'):
                                    planting_date = planting_date.date()
                                
                                # Only check if planted today
                                if isinstance(planting_date, type(today)) and (today - planting_date).days == 0:
                                    # For plants placed today, only skip if analyzed within last 10 seconds
                                    if space.last_updated:
                                        try:
                                            last_updated = space.last_updated
                                            now = datetime.now(timezone.utc)
                                            
                                            # Handle timezone-aware and naive datetimes
                                            if isinstance(last_updated, datetime):
                                                if last_updated.tzinfo is None:
                                                    last_updated_dt = last_updated.replace(tzinfo=timezone.utc)
                                                else:
                                                    last_updated_dt = last_updated
                                            else:
                                                # If it's a date, convert to datetime
                                                last_updated_dt = datetime.combine(last_updated, datetime.min.time()).replace(tzinfo=timezone.utc)
                                            
                                            seconds_since_update = (now - last_updated_dt).total_seconds()
                                            
                                            # Only skip if analyzed within last 10 seconds (very strict - immediate first-time only)
                                            if seconds_since_update <= 10:
                                                is_newly_added = True
                                                print(f"⏭️ Skipping alerts for newly added plant {plant.name} (planted today, analyzed {seconds_since_update:.0f} seconds ago)")
                                            else:
                                                # Analyzed more than 10 seconds ago - treat as existing (being updated)
                                                print(f"✅ Showing alerts for existing plant {plant.name} (planted today, analyzed {seconds_since_update:.0f} seconds ago - being updated)")
                                                is_newly_added = False
                                        except Exception as e:
                                            print(f"⚠️ Error checking last_updated for {plant.name}: {e}")
                                            # On error, show alerts to be safe
                                            is_newly_added = False
                                    else:
                                        # No last_updated - show alerts
                                        print(f"✅ Showing alerts for plant {plant.name} (planted today, no last_updated)")
                                        is_newly_added = False
                                else:
                                    # Planted before today - definitely existing
                                    print(f"✅ Showing alerts for existing plant {plant.name} (planted before today)")
                                    is_newly_added = False
                            else:
                                # No planting_date - show alerts to be safe
                                print(f"✅ Showing alerts for plant {plant.name} (no planting_date, treating as existing)")
                                is_newly_added = False
                            
                            # Only create alerts if this is not a newly added plant
                            if not is_newly_added:
                                confidence = ai_analysis.get('confidence', 0.5)
                                reasoning = ai_analysis.get('reasoning', 'AI analysis suggests care needed')
                                
                                print(f"🔍 AI Analysis for {plant.name}: {ai_analysis}")
                                print(f"🔍 Needs water: {needs_water}")
                                print(f"🔍 Needs fertilize: {needs_fertilize}")
                                print(f"🔍 Needs prune: {needs_prune}")
                                
                                # Generate AI suggested actions tags
                                suggested_actions = []
                                if needs_water:
                                    suggested_actions.append({'type': 'water', 'label': 'Needs Water', 'icon': '💧'})
                                if needs_fertilize:
                                    suggested_actions.append({'type': 'fertilize', 'label': 'Needs Fertilizer', 'icon': '🌱'})
                                if needs_prune:
                                    suggested_actions.append({'type': 'prune', 'label': 'Needs Pruning', 'icon': '✂️'})
                                
                                # Check for watering needs
                                if needs_water:
                                    # Skip if plant was watered today (already completed)
                                    was_watered_today = False
                                    if space.last_watered:
                                        last_watered_date = space.last_watered
                                        if hasattr(last_watered_date, 'date'):
                                            last_watered_date = last_watered_date.date()
                                        if isinstance(last_watered_date, type(today)) and (today - last_watered_date).days == 0:
                                            was_watered_today = True
                                            print(f"⏭️ Skipping watering alert for {plant.name} - already watered today")
                                    
                                    if not was_watered_today:
                                        print(f"💧 AI detected watering need for {plant.name}")
                                        print(f"🚨 GENERATING AI WATERING ALERT for {plant.name}: {reasoning}")
                                        
                                        alerts.append({
                                            'id': f"ai_water_{space.id}",
                                            'type': 'watering',
                                            'plant_name': plant.name,
                                            'garden_name': garden.name,
                                            'message': f'Your {plant.name} needs watering based on AI analysis',
                                            'due_date': today.isoformat(),
                                            'priority': 'high' if confidence > 0.7 else 'medium',
                                            'status': 'pending',
                                            'space_id': space.id,
                                            'garden_id': garden.id,
                                            'recommendation': f"AI Analysis: {reasoning}",
                                            'grid_position': space.grid_position,
                                            'ai_confidence': confidence,
                                            'ai_suggested_actions': suggested_actions
                                        })
                                        print(f"✅ WATERING ALERT ADDED: {plant.name} - Total alerts now: {len(alerts)}")
                                
                                # Check for fertilizing needs
                                if needs_fertilize:
                                    # Skip if plant was fertilized today (already completed)
                                    was_fertilized_today = False
                                    if space.last_fertilized:
                                        last_fertilized_date = space.last_fertilized
                                        if hasattr(last_fertilized_date, 'date'):
                                            last_fertilized_date = last_fertilized_date.date()
                                        if isinstance(last_fertilized_date, type(today)) and (today - last_fertilized_date).days == 0:
                                            was_fertilized_today = True
                                            print(f"⏭️ Skipping fertilizing alert for {plant.name} - already fertilized today")
                                    
                                    if not was_fertilized_today:
                                        print(f"🌱 AI detected fertilizing need for {plant.name}")
                                        print(f"🚨 GENERATING FERTILIZING ALERT for {plant.name}: {reasoning}")
                                        
                                        alerts.append({
                                            'id': f"ai_fertilize_{space.id}",
                                            'type': 'fertilizing',
                                            'plant_name': plant.name,
                                            'garden_name': garden.name,
                                            'message': f'Your {plant.name} needs fertilizer based on AI analysis',
                                            'due_date': today.isoformat(),
                                            'priority': 'high' if confidence > 0.7 else 'medium',
                                            'status': 'pending',
                                            'space_id': space.id,
                                            'garden_id': garden.id,
                                            'recommendation': f"AI Analysis: {reasoning}",
                                            'grid_position': space.grid_position,
                                            'ai_confidence': confidence,
                                            'ai_suggested_actions': suggested_actions
                                        })
                                        print(f"✅ FERTILIZING ALERT ADDED: {plant.name} - Total alerts now: {len(alerts)}")
                                
                                # Check for pruning needs
                                if needs_prune:
                                    # Skip if plant was pruned today (already completed)
                                    was_pruned_today = False
                                    if space.last_pruned:
                                        last_pruned_date = space.last_pruned
                                        if hasattr(last_pruned_date, 'date'):
                                            last_pruned_date = last_pruned_date.date()
                                        if isinstance(last_pruned_date, type(today)) and (today - last_pruned_date).days == 0:
                                            was_pruned_today = True
                                            print(f"⏭️ Skipping pruning alert for {plant.name} - already pruned today")
                                    
                                    if not was_pruned_today:
                                        print(f"✂️ AI detected pruning need for {plant.name}")
                                        print(f"🚨 GENERATING PRUNING ALERT for {plant.name}: {reasoning}")
                                        
                                        alerts.append({
                                            'id': f"ai_prune_{space.id}",
                                            'type': 'pruning',
                                            'plant_name': plant.name,
                                            'garden_name': garden.name,
                                            'message': f'Your {plant.name} needs pruning based on AI analysis',
                                            'due_date': today.isoformat(),
                                            'priority': 'high' if confidence > 0.7 else 'medium',
                                            'status': 'pending',
                                            'space_id': space.id,
                                            'garden_id': garden.id,
                                            'recommendation': f"AI Analysis: {reasoning}",
                                            'grid_position': space.grid_position,
                                            'ai_confidence': confidence,
                                            'ai_suggested_actions': suggested_actions
                                        })
                                        print(f"✅ PRUNING ALERT ADDED: {plant.name} - Total alerts now: {len(alerts)}")
                            else:
                                print(f"⏭️ Skipping alerts for {plant.name} - plant was just added via AI recognition")
                    
                    # Only check traditional schedule if AI analysis doesn't exist
                    # If AI analysis exists, we rely on it exclusively to avoid duplicates
                    if not ai_analysis:
                        if space.last_watered:
                            # Ensure last_watered is datetime.date for proper subtraction
                            last_watered_date = space.last_watered
                            if hasattr(last_watered_date, 'date'):
                                last_watered_date = last_watered_date.date()
                            days_since_watered = (today - last_watered_date).days
                            # Default watering frequency if not set in plant
                            watering_freq = getattr(plant, 'watering_frequency', 3) or 3
                            
                            if days_since_watered >= watering_freq:
                                # Generate plant-specific watering recommendations
                                watering_recommendation = get_watering_recommendation(plant, days_since_watered)
                                
                                alerts.append({
                                    'id': f"grid_water_{space.id}",
                                    'type': 'watering',
                                    'plant_name': plant.name,
                                    'garden_name': garden.name,
                                    'message': f'Time to water your {plant.name} in {garden.name}',
                                    'due_date': (space.last_watered + timedelta(days=watering_freq)).isoformat(),
                                    'priority': 'high' if days_since_watered > watering_freq + 2 else 'medium',
                                    'status': 'overdue' if days_since_watered > watering_freq else 'pending',
                                    'space_id': space.id,
                                    'garden_id': garden.id,
                                    'recommendation': watering_recommendation,
                                    'grid_position': space.grid_position
                                })
                        else:
                            # New plant that hasn't been watered yet
                            new_plant_recommendation = get_new_plant_recommendation(plant)
                            
                            alerts.append({
                                'id': f"grid_water_new_{space.id}",
                                'type': 'watering',
                                'plant_name': plant.name,
                                'garden_name': garden.name,
                                'message': f'Your newly planted {plant.name} needs its first watering',
                                'due_date': today.isoformat(),
                                'priority': 'high',
                                'status': 'pending',
                                'space_id': space.id,
                                'garden_id': garden.id,
                                'recommendation': new_plant_recommendation,
                                'grid_position': f"Row {space.row}, Col {space.column}" if hasattr(space, 'row') and hasattr(space, 'column') else None
                            })
                
                # Only check traditional schedule if AI analysis doesn't exist
                # If AI analysis exists, we rely on it exclusively to avoid duplicates
                if not ai_analysis:
                    if space.last_fertilized:
                        # Ensure last_fertilized is datetime.date for proper subtraction
                        last_fertilized_date = space.last_fertilized
                        if hasattr(last_fertilized_date, 'date'):
                            last_fertilized_date = last_fertilized_date.date()
                        days_since_fertilized = (today - last_fertilized_date).days
                        fertilizing_freq = getattr(plant, 'fertilizing_frequency', 14) or 14
                        
                        if days_since_fertilized >= fertilizing_freq:
                            fertilizing_recommendation = get_fertilizing_recommendation(plant, days_since_fertilized)
                            
                            alerts.append({
                                'id': f"grid_fertilize_{space.id}",
                                'type': 'fertilizing',
                                'plant_name': plant.name,
                                'garden_name': garden.name,
                                'message': f'Your {plant.name} needs fertilizer',
                                'due_date': (space.last_fertilized + timedelta(days=fertilizing_freq)).isoformat(),
                                'priority': 'medium',
                                'status': 'overdue' if days_since_fertilized > fertilizing_freq else 'pending',
                                'space_id': space.id,
                                'garden_id': garden.id,
                                'recommendation': fertilizing_recommendation,
                                'grid_position': space.grid_position
                            })
                
                # Only check traditional schedule if AI analysis doesn't exist
                # If AI analysis exists, we rely on it exclusively to avoid duplicates
                if not ai_analysis:
                    if space.last_pruned:
                        # Ensure last_pruned is datetime.date for proper subtraction
                        last_pruned_date = space.last_pruned
                        if hasattr(last_pruned_date, 'date'):
                            last_pruned_date = last_pruned_date.date()
                        days_since_pruned = (today - last_pruned_date).days
                        pruning_freq = getattr(plant, 'pruning_frequency', 30) or 30
                        
                        if days_since_pruned >= pruning_freq:
                            pruning_recommendation = get_pruning_recommendation(plant, days_since_pruned)
                            
                            alerts.append({
                                'id': f"grid_prune_{space.id}",
                                'type': 'pruning',
                                'plant_name': plant.name,
                                'garden_name': garden.name,
                                'message': f'Time to prune your {plant.name}',
                                'due_date': (space.last_pruned + timedelta(days=pruning_freq)).isoformat(),
                                'priority': 'low',
                                'status': 'overdue' if days_since_pruned > pruning_freq else 'pending',
                                'space_id': space.id,
                                'garden_id': garden.id,
                                'recommendation': pruning_recommendation,
                                'grid_position': space.grid_position
                            })
        
        # Also check traditional PlantTracking (legacy system)
        for garden in user_gardens:
            plant_trackings = PlantTracking.query.filter_by(garden_id=garden.id).all()
            
            for pt in plant_trackings:
                plant = Plant.query.get(pt.plant_id)
                if not plant:
                    continue
                
                # Skip alerts for newly added plants (added within last 48 hours via AI recognition)
                # Check if this is a newly added plant with no care history
                has_care_history = pt.last_watered or pt.last_fertilized or pt.last_pruned
                
                # Check if plant was added recently (within 48 hours)
                is_newly_added = False
                if pt.planting_date and not has_care_history:
                    planting_date = pt.planting_date
                    if hasattr(planting_date, 'date'):
                        planting_date = planting_date.date()
                    
                    # Check if plant was added within last 48 hours
                    if isinstance(planting_date, type(today)):
                        days_since_planted = (today - planting_date).days
                        if days_since_planted < 2:  # Less than 2 days old
                            is_newly_added = True
                            print(f"⏭️ Skipping PlantTracking alerts for newly added plant {plant.name} (planted {days_since_planted} days ago, no care history)")
                
                # Skip all alerts for newly added plants
                if is_newly_added:
                    continue
                    
                # Check watering schedule
                if plant.watering_frequency:
                    # Ensure last_watered is datetime.date for proper subtraction
                    if pt.last_watered:
                        last_watered_date = pt.last_watered
                        if hasattr(last_watered_date, 'date'):
                            last_watered_date = last_watered_date.date()
                        days_since_watered = (today - last_watered_date).days
                    else:
                        days_since_watered = 999
                    if days_since_watered >= plant.watering_frequency:
                        alerts.append({
                            'id': f"track_water_{pt.id}",
                            'type': 'watering',
                            'plant_name': plant.name,
                            'garden_name': garden.name,
                            'message': f'Time to water your {plant.name}',
                            'due_date': (pt.last_watered + timedelta(days=plant.watering_frequency)).isoformat() if pt.last_watered else today.isoformat(),
                            'priority': 'high' if days_since_watered > plant.watering_frequency + 2 else 'medium',
                            'status': 'overdue' if days_since_watered > plant.watering_frequency else 'pending'
                        })
                
                # Check fertilizing schedule
                if plant.fertilizing_frequency:
                    # Ensure last_fertilized is datetime.date for proper subtraction
                    if pt.last_fertilized:
                        last_fertilized_date = pt.last_fertilized
                        if hasattr(last_fertilized_date, 'date'):
                            last_fertilized_date = last_fertilized_date.date()
                        days_since_fertilized = (today - last_fertilized_date).days
                    else:
                        days_since_fertilized = 999
                    if days_since_fertilized >= plant.fertilizing_frequency:
                        alerts.append({
                            'id': f"track_fertilize_{pt.id}",
                            'type': 'fertilizing',
                            'plant_name': plant.name,
                            'garden_name': garden.name,
                            'message': f'Your {plant.name} needs fertilizer',
                            'due_date': (pt.last_fertilized + timedelta(days=plant.fertilizing_frequency)).isoformat() if pt.last_fertilized else today.isoformat(),
                            'priority': 'medium',
                            'status': 'overdue' if days_since_fertilized > plant.fertilizing_frequency else 'pending'
                        })
                
                # Check pruning schedule
                if plant.pruning_frequency:
                    # Ensure last_pruned is datetime.date for proper subtraction
                    if pt.last_pruned:
                        last_pruned_date = pt.last_pruned
                        if hasattr(last_pruned_date, 'date'):
                            last_pruned_date = last_pruned_date.date()
                        days_since_pruned = (today - last_pruned_date).days
                    else:
                        days_since_pruned = 999
                    if days_since_pruned >= plant.pruning_frequency:
                        alerts.append({
                            'id': f"track_prune_{pt.id}",
                            'type': 'pruning',
                            'plant_name': plant.name,
                            'garden_name': garden.name,
                            'message': f'Time to prune your {plant.name}',
                            'due_date': (pt.last_pruned + timedelta(days=plant.pruning_frequency)).isoformat() if pt.last_pruned else today.isoformat(),
                            'priority': 'low',
                            'status': 'overdue' if days_since_pruned > plant.pruning_frequency else 'pending'
                        })
    
        # Sort alerts by priority and due date
        priority_order = {'high': 3, 'medium': 2, 'low': 1}
        alerts.sort(key=lambda x: (priority_order.get(x['priority'], 0), x['due_date']), reverse=True)
        
        print(f"📊 Generated {len(alerts)} total alerts for user {current_user.id}")
        for alert in alerts:
            print(f"  - {alert['type']}: {alert['plant_name']} in {alert['garden_name']} (Priority: {alert['priority']})")
        
        # Get completed actions for the same time period (no limit - show all completed actions)
        completed_actions = ActivityLog.query.filter_by(
            user_id=current_user.id
        ).order_by(ActivityLog.action_date.desc()).all()
        
        # Filter out alerts that have been completed recently (within last 24 hours)
        print(f"🔍 Filtering out recently completed actions...")
        print(f"🔍 Found {len(completed_actions)} completed actions to check against {len(alerts)} alerts")
        filtered_alerts = []
        for alert in alerts:
            print(f"🔍 Checking alert: {alert['id']} ({alert['type']}) for space {alert.get('space_id')}")
            # Check if this alert has been completed recently
            recently_completed = False
            for action in completed_actions:
                print(f"🔍 Checking against completed action: space_id={action.space_id}, action={action.action}, date={action.action_date}")
                print(f"🔍 Alert type: {alert.get('type')}, Alert space_id: {alert.get('space_id')}")
                print(f"🔍 Action type: {action.action}, Action space_id: {action.space_id}")
                print(f"🔍 Type conversion: {alert.get('type')} -> {alert.get('type').replace('ing', '')}")
                
                # Check if this is the same space and same action type
                space_match = alert.get('space_id') and action.space_id == int(alert.get('space_id'))
                action_match = action.action == alert.get('type').replace('ing', '')
                print(f"🔍 Space match: {space_match} (alert: {alert.get('space_id')}, action: {action.space_id})")
                print(f"🔍 Action match: {action_match} (alert: {alert.get('type')}, action: {action.action})")
                
                if (alert.get('space_id') and 
                    action.space_id == int(alert.get('space_id')) and
                    action.action == alert.get('type').replace('ing', '')):  # 'watering' -> 'water'
                    
                    # Check if completed within last 24 hours
                    # Ensure action_date is datetime.datetime, not datetime.date
                    action_datetime = action.action_date
                    if hasattr(action_datetime, 'date') and not hasattr(action_datetime, 'hour'):
                        # If it's a date object, convert to datetime at midnight
                        action_datetime = datetime.combine(action_datetime, datetime.min.time())
                    
                    # Use timezone-aware datetime for proper calculation
                    now = datetime.now(timezone.utc)
                    if action_datetime.tzinfo is None:
                        action_datetime = action_datetime.replace(tzinfo=timezone.utc)
                    
                    hours_since_completed = (now - action_datetime).total_seconds() / 3600
                    print(f"🔍 Match found! Hours since completed: {hours_since_completed:.1f}")
                    if hours_since_completed < 24:
                        print(f"🚫 Filtering out recently completed alert: {alert['type']} for {alert['plant_name']} (completed {hours_since_completed:.1f} hours ago)")
                        recently_completed = True
                        break
                    else:
                        print(f"✅ Alert not filtered - completed {hours_since_completed:.1f} hours ago (beyond 24h threshold)")
            
            if not recently_completed:
                filtered_alerts.append(alert)
                print(f"✅ Alert kept: {alert['id']} ({alert['type']})")
            else:
                print(f"🚫 Alert filtered out: {alert['id']} ({alert['type']})")
        
        alerts = filtered_alerts
        print(f"📊 After filtering completed actions: {len(alerts)} alerts remaining")
        
        completed_data = []
        for action in completed_actions:
            plant = Plant.query.get(action.plant_id) if action.plant_id else None
            garden = Garden.query.get(action.garden_id) if action.garden_id else None
            
            completed_data.append({
                'id': f"completed_{action.id}",
                'type': action.action,
                'plant_name': plant.name if plant else 'Unknown Plant',
                'garden_name': garden.name if garden else 'Unknown Garden',
                'action_date': action.action_date.isoformat(),
                'created_at': action.created_at.isoformat() if action.created_at else action.action_date.isoformat(),
                'status': 'completed',
                'notes': action.notes
            })
        
        return jsonify({
            "alerts": alerts,
            "completed_actions": completed_data
        })
    
    except Exception as e:
        print(f"❌ Error in smart_alerts: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "error": f"Failed to load alerts: {str(e)}",
            "alerts": [],
            "completed_actions": []
        }), 500

@views.route('/api/alerts/mark-completed', methods=['POST'])
@login_required
def mark_alert_completed():
    """Mark an alert as completed when user performs the care action"""
    try:
        data = request.get_json()
        alert_id = data.get('alert_id')
        action = data.get('action')  # 'water', 'fertilize', 'prune'
        
        print(f"🎯 ALERT COMPLETION: User {current_user.id} marking alert {alert_id} as {action}")
        
        if not alert_id or not action:
            return jsonify({"error": "Alert ID and action are required"}), 400
        
        # Parse alert ID to determine if it's from grid or tracking
        if alert_id.startswith('grid_') or alert_id.startswith('ai_'):
            # Handle grid space alerts (both traditional and AI-based)
            space_id = alert_id.split('_')[-1]
            print(f"🎯 ALERT COMPLETION: Processing grid space {space_id}")
            
            space = GridSpace.query.get(space_id)
            
            if not space:
                print(f"🎯 ALERT COMPLETION: Grid space {space_id} not found")
                return jsonify({"error": "Grid space not found"}), 404
            
            garden = Garden.query.get(space.garden_id)
            if garden.user_id != current_user.id:
                print(f"🎯 ALERT COMPLETION: Unauthorized access to garden {garden.id}")
                return jsonify({"error": "Unauthorized"}), 403
            
            # Update the appropriate care field
            today = datetime.now().date()
            if action == 'water':
                space.last_watered = today
                print(f"🎯 ALERT COMPLETION: Updated last_watered to {today}")
            elif action == 'fertilize':
                space.last_fertilized = today
                print(f"🎯 ALERT COMPLETION: Updated last_fertilized to {today}")
            elif action == 'prune':
                space.last_pruned = today
                print(f"🎯 ALERT COMPLETION: Updated last_pruned to {today}")
            else:
                print(f"🎯 ALERT COMPLETION: Invalid action {action}")
                return jsonify({"error": f"Invalid action: {action}"}), 400
            
            # Don't clear AI analysis - keep it so other care needs can still show alerts
            # The filtering logic will prevent duplicate alerts for the completed action
            # Only the specific action type that was completed will be filtered out
            
            # Create activity log entry for completed action
            activity_log = ActivityLog(
                user_id=current_user.id,
                garden_id=space.garden_id,
                space_id=space.id,
                plant_id=space.plant_id,
                action=action,
                action_date=today,
                notes=f"Completed {action}ing for {space.plant.name if space.plant else 'plant'}"
            )
            db.session.add(activity_log)
            
            space.last_updated = datetime.now(timezone.utc)
            db.session.commit()
            
            print(f"🎯 ALERT COMPLETION: Successfully updated grid space {space_id}")
            return jsonify({
                "success": True,
                "message": f"Plant {action}ed successfully! Alert marked as completed."
            })
            
        elif alert_id.startswith('track_'):
            # Handle plant tracking alerts
            tracking_id = alert_id.split('_')[-1]
            tracking = PlantTracking.query.get(tracking_id)
            
            if not tracking:
                return jsonify({"error": "Plant tracking not found"}), 404
            
            garden = Garden.query.get(tracking.garden_id)
            if garden.user_id != current_user.id:
                return jsonify({"error": "Unauthorized"}), 403
            
            # Update the appropriate care field
            today = datetime.now().date()
            if action == 'water':
                tracking.last_watered = today
            elif action == 'fertilize':
                tracking.last_fertilized = today
            elif action == 'prune':
                tracking.last_pruned = today
            
            db.session.commit()
            
            return jsonify({
                "success": True,
                "message": f"Plant {action}ed successfully! Alert marked as completed."
            })
        
        else:
            return jsonify({"error": "Invalid alert ID format"}), 400
            
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@views.route('/api/completed-actions')
@login_required
def get_completed_actions():
    """Get completed plant care actions for reports"""
    try:
        # Get user's completed actions from activity logs
        completed_actions = ActivityLog.query.filter_by(
            user_id=current_user.id
        ).order_by(ActivityLog.action_date.desc()).all()
        
        actions_data = []
        for action in completed_actions:
            # Get plant and garden info
            plant = Plant.query.get(action.plant_id) if action.plant_id else None
            garden = Garden.query.get(action.garden_id) if action.garden_id else None
            
            actions_data.append({
                'id': action.id,
                'action': action.action,
                'action_date': action.action_date.isoformat(),
                'plant_name': plant.name if plant else 'Unknown Plant',
                'garden_name': garden.name if garden else 'Unknown Garden',
                'notes': action.notes,
                'space_id': action.space_id
            })
        
        return jsonify({
            "success": True,
            "completed_actions": actions_data
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@views.route('/api/plant-activity-logs', methods=['GET'])
@login_required
def get_plant_activity_logs():
    """Get activity logs for a specific grid space/plant"""
    try:
        from website.models import ActivityLog, GridSpace, Plant
        
        space_id = request.args.get('space_id', type=int)
        plant_id = request.args.get('plant_id', type=int)
        
        if not space_id and not plant_id:
            return jsonify({"error": "space_id or plant_id is required"}), 400
        
        # Build query
        query = ActivityLog.query.filter_by(user_id=current_user.id)
        
        if space_id:
            query = query.filter_by(space_id=space_id)
        if plant_id:
            query = query.filter_by(plant_id=plant_id)
        
        # Order by most recent first
        activity_logs = query.order_by(ActivityLog.created_at.desc()).all()
        
        # Enhance with plant and space info
        logs_data = []
        for log in activity_logs:
            log_dict = log.to_dict()
            if log.space:
                log_dict['grid_position'] = log.space.grid_position
            if log.plant:
                log_dict['plant_name'] = log.plant.name
            logs_data.append(log_dict)
        
        return jsonify({
            "success": True,
            "activity_logs": logs_data
        })
    except Exception as e:
        print(f"Error fetching plant activity logs: {str(e)}")
        return jsonify({"error": str(e)}), 500

@views.route('/features')
@login_required
def features():
    return jsonify({"message": "Features feature"})

@views.route('/about')
def about():
    return jsonify({"message": "About feature"})

@views.route('/admin/reports')
@login_required
def admin_reports():
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    # Get statistics for reports
    total_users = User.query.count()
    active_users = User.query.filter_by(is_active=True).count()
    subscribed_users = User.query.filter_by(subscribed=True).count()
    total_admins = Admin.query.count()
    
    # Get recent user registrations
    recent_users = User.query.order_by(User.created_at.desc()).limit(10).all()
    
    # Get user activity data (placeholder for now)
    user_activity = {
        'daily': 45,
        'weekly': 320,
        'monthly': 1200
    }
    
    return jsonify({
        "message": "Admin reports",
        "stats": {
            "total_users": total_users,
            "active_users": active_users,
            "subscribed_users": subscribed_users,
            "total_admins": total_admins,
            "recent_users": [{"id": u.id, "email": u.email, "full_name": u.full_name} for u in recent_users],
            "user_activity": user_activity
        }
    })

@views.route('/admin/feedback')
@login_required
def admin_feedback():
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    try:
        # Get all feedback from database
        feedbacks = Feedback.query.order_by(Feedback.created_at.desc()).all()
        feedback_list = [feedback.to_dict() for feedback in feedbacks]
        
        return jsonify({"feedback_list": feedback_list})
    except Exception as e:
        return jsonify({"error": "Failed to fetch feedback. Please try again."}), 500

@views.route('/admin/feedback/<int:feedback_id>/update', methods=['POST'])
@login_required
def update_feedback_status(feedback_id):
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    data = request.get_json()
    status = data.get('status')
    admin_response = data.get('admin_response', '')
    
    if not status:
        return jsonify({"error": "Status is required."}), 400
    
    try:
        # Find the feedback
        feedback = Feedback.query.get_or_404(feedback_id)
        
        # Update status and admin response
        feedback.status = status
        if admin_response:
            feedback.admin_response = admin_response
        
        db.session.commit()
        
        return jsonify({"message": f"Feedback status updated to {status}."})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": "Failed to update feedback status. Please try again."}), 500

@views.route('/admin/system-content')
@login_required
def admin_system_content():
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    # Placeholder data for system content management
    ai_data = {
        'plant_recognition_model': 'v2.1',
        'last_updated': '2024-01-15',
        'accuracy': '94.5%',
        'total_plants': 1500
    }
    
    learning_paths = [
        {'id': 1, 'name': 'Beginner Gardening', 'status': 'active', 'users': 45},
        {'id': 2, 'name': 'Intermediate Techniques', 'status': 'active', 'users': 23},
        {'id': 3, 'name': 'Advanced Horticulture', 'status': 'draft', 'users': 0}
    ]
    
    seasonal_content = [
        {'season': 'Spring', 'status': 'published', 'last_updated': '2024-03-01'},
        {'season': 'Summer', 'status': 'published', 'last_updated': '2024-06-01'},
        {'season': 'Fall', 'status': 'draft', 'last_updated': '2024-09-01'},
        {'season': 'Winter', 'status': 'draft', 'last_updated': '2024-12-01'}
    ]
    
    notification_templates = [
        {'id': 1, 'name': 'Watering Reminder', 'status': 'active'},
        {'id': 2, 'name': 'Fertilizing Schedule', 'status': 'active'},
        {'id': 3, 'name': 'Seasonal Tips', 'status': 'active'},
        {'id': 4, 'name': 'New Feature Alert', 'status': 'draft'}
    ]
    
    return jsonify({
        "message": "Admin system content",
        "ai_data": ai_data,
        "learning_paths": learning_paths,
        "seasonal_content": seasonal_content,
        "notification_templates": notification_templates
    })

@views.route('/admin/ai-data', methods=['GET', 'POST'])
@login_required
def admin_ai_data():
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    if request.method == 'POST':
        # Handle AI data updates
        return jsonify({"message": "AI data updated successfully!"})
    
    return jsonify({"message": "Admin AI data"})

@views.route('/admin/learning-paths', methods=['GET', 'POST'])
@login_required
def admin_learning_paths():
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    if request.method == 'POST':
        # Handle learning path updates
        return jsonify({"message": "Learning paths updated successfully!"})
    
    return jsonify({"message": "Admin learning paths"})

@views.route('/admin/notifications', methods=['GET', 'POST'])
@login_required
def admin_notifications():
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    if request.method == 'POST':
        # Handle notification template updates
        return jsonify({"message": "Notification templates updated successfully!"})
    
    return jsonify({"message": "Admin notifications"})

@views.route('/admin/seasonal-planning', methods=['GET', 'POST'])
@login_required
def admin_seasonal_planning():
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    if request.method == 'POST':
        # Handle seasonal content updates
        return jsonify({"message": "Seasonal content updated successfully!"})
    
    return jsonify({"message": "Admin seasonal planning"})

@views.route('/admin/rollback-content', methods=['GET', 'POST'])
@login_required
def admin_rollback_content():
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    if request.method == 'POST':
        # Handle content rollback
        return jsonify({"message": "Content rollback completed successfully!"})
    
    return jsonify({"message": "Admin rollback content"})

@views.route('/admin/users')
@login_required
def admin_users():
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    users = User.query.all()
    return jsonify({
        "users": [{
            "id": u.id,
            "email": u.email,
            "full_name": u.full_name,
            "contact": u.contact,
            "role": u.role,
            "is_active": u.is_active,
            "created_at": u.created_at.isoformat() if u.created_at else None
        } for u in users]
    })

@views.route('/admin/users/<int:user_id>/status', methods=['PUT'])
@login_required
def admin_user_status(user_id):
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    data = request.get_json()
    is_active = data.get('is_active')
    
    user = User.query.get_or_404(user_id)
    user.is_active = is_active
    db.session.commit()
    
    return jsonify({"message": f"User status updated successfully"})

@views.route('/admin/users/<int:user_id>', methods=['DELETE'])
@login_required
def admin_delete_user(user_id):
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    try:
        user = User.query.get_or_404(user_id)
        if user.role == 'admin' and user.id != current_user.id:
            return jsonify({"error": "Cannot delete other admin users"}), 403
        
        # Manually delete related records to avoid foreign key constraint issues
        UserSubscription.query.filter_by(user_id=user_id).delete()
        ActivityLog.query.filter_by(user_id=user_id).delete()
        db.session.commit()
        
        db.session.delete(user)
        db.session.commit()
        
        return jsonify({"message": "User deleted successfully"})
    except Exception as e:
        db.session.rollback()
        import traceback
        error_msg = str(e)
        print(f"Error deleting user {user_id}: {error_msg}")
        print(traceback.format_exc())
        return jsonify({"error": f"Failed to delete user: {error_msg}"}), 500

@views.route('/admin/feedbacks')
@login_required
def admin_feedbacks():
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    try:
        # Get all feedback from database with user relationship loaded
        from sqlalchemy.orm import joinedload
        feedbacks = Feedback.query.options(joinedload(Feedback.user)).order_by(Feedback.created_at.desc()).all()
        feedback_list = [feedback.to_dict() for feedback in feedbacks]
        
        return jsonify({"feedbacks": feedback_list})
    except Exception as e:
        import traceback
        print(f"Error fetching feedback: {str(e)}")
        print(traceback.format_exc())
        return jsonify({"error": f"Failed to fetch feedback: {str(e)}"}), 500

@views.route('/admin/feedbacks/<int:feedback_id>/status', methods=['PUT'])
@login_required
def admin_feedback_status(feedback_id):
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    data = request.get_json()
    status = data.get('status')
    admin_response = data.get('admin_response', '')
    
    if not status:
        return jsonify({"error": "Status is required."}), 400
    
    try:
        # Find the feedback
        feedback = Feedback.query.get_or_404(feedback_id)
        
        # Update status and admin response
        feedback.status = status
        if admin_response:
            feedback.admin_response = admin_response
        
        db.session.commit()
        
        return jsonify({"message": "Feedback status updated successfully"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": "Failed to update feedback status. Please try again."}), 500

@views.route('/admin/feedback/<int:feedback_id>', methods=['DELETE'])
@login_required
def delete_feedback(feedback_id):
    if not current_user.is_admin():
        return jsonify({"error": "Access denied. Admin privileges required."}), 403
    
    try:
        # Find the feedback
        feedback = Feedback.query.get_or_404(feedback_id)
        
        # Delete the feedback
        db.session.delete(feedback)
        db.session.commit()
        
        return jsonify({"message": "Feedback deleted successfully"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": "Failed to delete feedback. Please try again."}), 500

@views.route('/feedback/submit', methods=['POST'])
@login_required
def submit_feedback():
    data = request.get_json()
    subject = data.get('subject')
    message = data.get('message')
    rating = data.get('rating', 5)
    category = data.get('category', 'general')
    
    if not subject or not message:
        return jsonify({"error": "Subject and message are required."}), 400
    
    try:
        # Create new feedback entry
        feedback = Feedback(
            user_id=current_user.id,
            subject=subject,
            message=message,
            rating=rating,
            category=category,
            status='pending'
        )
        
        db.session.add(feedback)
        db.session.commit()
        
        return jsonify({"success": True, "message": "Feedback submitted successfully!"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": "Failed to submit feedback. Please try again."}), 500

@views.route('/feedback/user')
@login_required
def user_feedbacks():
    try:
        # Fetch user's feedback from database
        feedbacks = Feedback.query.filter_by(user_id=current_user.id).order_by(Feedback.created_at.desc()).all()
        
        feedback_list = [feedback.to_dict() for feedback in feedbacks]
        
        return jsonify({"feedbacks": feedback_list})
    except Exception as e:
        return jsonify({"error": "Failed to fetch feedback. Please try again."}), 500


@views.route('/concepts', methods=['GET'])
@login_required
def get_concepts():
    """Fetch community-shared concepts and the current user's own submissions."""
    try:
        community_query = UserSharedConcept.query.filter(
            (UserSharedConcept.is_public.is_(True)) | (UserSharedConcept.user_id == current_user.id)
        ).order_by(UserSharedConcept.created_at.desc())
        my_query = UserSharedConcept.query.filter_by(user_id=current_user.id).order_by(UserSharedConcept.created_at.desc())
        
        community = [concept.to_dict(include_owner=True) for concept in community_query]
        mine = [concept.to_dict(include_owner=True) for concept in my_query]
        
        return jsonify({
            "concepts": {
                "community": community,
                "mine": mine
            }
        })
    except Exception as e:
        return jsonify({"error": f"Failed to load concepts: {str(e)}"}), 500


@views.route('/concepts', methods=['POST'])
@login_required
def create_concept():
    """Create a new concept/technique shared by the user."""
    data = request.get_json() or {}
    title = (data.get('title') or '').strip()
    if not title:
        return jsonify({"error": "Title is required."}), 400
    
    summary = (data.get('summary') or '').strip() or None
    technique_steps = data.get('technique_steps')
    tips = data.get('tips')
    tags = _normalize_tags(data.get('tags'))
    is_public = bool(data.get('is_public', True))
    
    try:
        concept = UserSharedConcept(
            user_id=current_user.id,
            title=title,
            summary=summary,
            technique_steps=technique_steps,
            tips=tips,
            tags=tags,
            is_public=is_public,
            source='manual'
        )
        db.session.add(concept)
        db.session.commit()
        
        return jsonify({"success": True, "concept": concept.to_dict(include_owner=True)})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": f"Failed to save concept: {str(e)}"}), 500


@views.route('/concepts/<int:concept_id>', methods=['PUT'])
@login_required
def update_concept(concept_id):
    """Update an existing concept owned by the user."""
    concept = UserSharedConcept.query.get_or_404(concept_id)
    if concept.user_id != current_user.id:
        return jsonify({"error": "Unauthorized to update this concept."}), 403
    
    data = request.get_json() or {}
    title = data.get('title')
    if title is not None:
        title = title.strip()
        if not title:
            return jsonify({"error": "Title cannot be empty."}), 400
        concept.title = title
    
    if 'summary' in data:
        concept.summary = (data.get('summary') or '').strip() or None
    if 'technique_steps' in data:
        concept.technique_steps = data.get('technique_steps')
    if 'tips' in data:
        concept.tips = data.get('tips')
    if 'tags' in data:
        concept.tags = _normalize_tags(data.get('tags'))
    if 'is_public' in data:
        concept.is_public = bool(data.get('is_public'))
    
    try:
        db.session.commit()
        return jsonify({"success": True, "concept": concept.to_dict(include_owner=True)})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": f"Failed to update concept: {str(e)}"}), 500


@views.route('/concepts/<int:concept_id>', methods=['DELETE'])
@login_required
def delete_concept(concept_id):
    """Delete a concept owned by the user."""
    concept = UserSharedConcept.query.get_or_404(concept_id)
    if concept.user_id != current_user.id:
        return jsonify({"error": "Unauthorized to delete this concept."}), 403
    
    try:
        db.session.delete(concept)
        db.session.commit()
        return jsonify({"success": True})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": f"Failed to delete concept: {str(e)}"}), 500


@views.route('/concepts/<int:concept_id>/export', methods=['GET'])
@login_required
def export_concept(concept_id):
    """Export a concept as a downloadable Word document (.docx)."""
    try:
        # Load concept
        concept = UserSharedConcept.query.get_or_404(concept_id)
        if not concept.is_public and concept.user_id != current_user.id and not current_user.is_admin():
            return jsonify({"error": "Unauthorized to export this concept."}), 403
        
        # Create a new Document
        doc = Document()
        
        # Helper function to ensure black text
        def set_black_text(run):
            """Set text color to black for a run."""
            if run:
                run.font.color.rgb = RGBColor(0, 0, 0)
            return run
        
        # Title
        title_text = concept.title or 'Untitled Concept'
        title = doc.add_heading(title_text, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title.runs:
            set_black_text(run)
        
        # Summary
        if concept.summary:
            summary_heading = doc.add_heading('Summary', level=2)
            for run in summary_heading.runs:
                set_black_text(run)
            summary_para = doc.add_paragraph()
            summary_run = summary_para.add_run(concept.summary)
            set_black_text(summary_run)
        
        # Technique / Steps
        if concept.technique_steps:
            technique_heading = doc.add_heading('Technique / Steps', level=2)
            for run in technique_heading.runs:
                set_black_text(run)
            # Split by newlines to preserve formatting
            steps = concept.technique_steps.split('\n')
            for step in steps:
                if step.strip():
                    step_para = doc.add_paragraph()
                    step_run = step_para.add_run(step.strip())
                    set_black_text(step_run)
                    # Add bullet if it starts with certain characters
                    if step.strip().startswith(('-', '*', '•')):
                        step_para.style = 'List Bullet'
        
        # Tips
        if concept.tips:
            tips_heading = doc.add_heading('Tips', level=2)
            for run in tips_heading.runs:
                set_black_text(run)
            tips = concept.tips.split('\n')
            for tip in tips:
                if tip.strip():
                    tip_para = doc.add_paragraph()
                    tip_run = tip_para.add_run(tip.strip())
                    set_black_text(tip_run)
                    # Add bullet if it starts with certain characters
                    if tip.strip().startswith(('-', '*', '•')):
                        tip_para.style = 'List Bullet'
        
        # Tags - tags are stored as comma-separated string
        if concept.tags:
            tags_list = [tag.strip() for tag in concept.tags.split(',') if tag.strip()]
            if tags_list:
                tags_heading = doc.add_heading('Tags', level=2)
                for run in tags_heading.runs:
                    set_black_text(run)
                tags_text = ', '.join([f'#{tag}' for tag in tags_list])
                tags_para = doc.add_paragraph()
                tags_run = tags_para.add_run(tags_text)
                set_black_text(tags_run)
        
        # Metadata
        doc.add_paragraph()  # Empty line
        metadata_para = doc.add_paragraph()
        shared_by_run = metadata_para.add_run('Shared by: ')
        set_black_text(shared_by_run)
        
        # Get owner from user relationship
        try:
            if concept.user:
                if hasattr(concept.user, 'full_name') and concept.user.full_name:
                    owner_name = concept.user.full_name
                elif hasattr(concept.user, 'email'):
                    owner_name = concept.user.email
                else:
                    owner_name = 'Unknown'
            else:
                owner_name = 'Unknown'
        except Exception:
            owner_name = 'Unknown'
        
        owner_run = metadata_para.add_run(owner_name)
        set_black_text(owner_run)
        
        if concept.created_at:
            created_date = concept.created_at.strftime('%Y-%m-%d %H:%M:%S') if hasattr(concept.created_at, 'strftime') else str(concept.created_at)
            created_run = metadata_para.add_run(f' • Created: {created_date}')
            set_black_text(created_run)
        
        # Ensure all text in document is black
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                set_black_text(run)
        
        # Save to BytesIO
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Clean filename
        safe_title = re.sub(r'[^\w\s-]', '', concept.title or 'concept').strip().replace(' ', '_')
        filename = f"{safe_title}_{concept_id}.docx"
        
        return Response(
            doc_io.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"'
            }
        )
    except Exception as e:
        import traceback
        print(f"Export error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({"error": f"Failed to export concept: {str(e)}"}), 500


@views.route('/concepts/import', methods=['POST'])
@login_required
def import_concept():
    """Import a concept from uploaded JSON or direct payload."""
    concept_data = None
    source_name = 'import'
    
    if 'file' in request.files:
        file = request.files['file']
        if not file:
            return jsonify({"error": "No file provided."}), 400
        try:
            payload = json.load(file)
            concept_data = payload.get('concept', payload)
            source_name = file.filename or 'imported_file'
        except Exception as e:
            return jsonify({"error": f"Invalid file format: {str(e)}"}), 400
    else:
        payload = request.get_json(force=True, silent=True) or {}
        concept_data = payload.get('concept', payload)
        source_name = payload.get('imported_from', 'import')
    
    if not concept_data:
        return jsonify({"error": "No concept data provided."}), 400
    
    title = (concept_data.get('title') or '').strip()
    if not title:
        return jsonify({"error": "Concept title is required."}), 400
    
    try:
        concept = UserSharedConcept(
            user_id=current_user.id,
            title=title,
            summary=(concept_data.get('summary') or '').strip() or None,
            technique_steps=concept_data.get('technique_steps'),
            tips=concept_data.get('tips'),
            tags=_normalize_tags(concept_data.get('tags')),
            is_public=bool(concept_data.get('is_public', True)),
            source='import',
            imported_from=source_name
        )
        db.session.add(concept)
        db.session.commit()
        
        return jsonify({"success": True, "concept": concept.to_dict(include_owner=True)})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": f"Failed to import concept: {str(e)}"}), 500

@views.route('/api/weather')
@login_required
def get_weather():
    # Simple rate limiting - prevent too many requests
    if not hasattr(get_weather, '_last_request_time'):
        get_weather._last_request_time = 0
    
    rate_limit_time = time.time()
    if rate_limit_time - get_weather._last_request_time < 2:  # 2 seconds minimum between requests
        print(f"🌤️ Weather API rate limited - too soon since last request")
        return jsonify({
            "error": "Rate limited - please wait before making another request",
            "success": False
        }), 429
    
    get_weather._last_request_time = rate_limit_time
    
    """Get weather data for a specific city using OpenWeatherMap API"""
    # Initialize variables outside try block to ensure they're available in except
    city = request.args.get('city', 'Cebu')
    cache_key = f"weather_{city.lower()}"
    current_time = time.time()
    
    try:
        if cache_key in _WEATHER_CACHE:
            cached_data, cache_time = _WEATHER_CACHE[cache_key]
            if current_time - cache_time < _WEATHER_CACHE_TTL_SECONDS:
                print(f"🌤️ Weather cache hit for {city} (cached {int(current_time - cache_time)}s ago)")
                return jsonify(cached_data)
            else:
                # Cache expired, remove it
                del _WEATHER_CACHE[cache_key]
        
        
        api_key = os.getenv('OPENWEATHER_API_KEY')
        
        if not api_key:
            # Return mock data if API key is not available
            mock_data = {
                "temperature": 28,
                "humidity": 75,
                "description": "Partly cloudy",
                "windSpeed": 12,
                "visibility": 10,
                "city": city,
                "mock": True
            }
            # Cache mock data too
            _WEATHER_CACHE[cache_key] = (mock_data, current_time)
            return jsonify(mock_data)
        
        # Make request to OpenWeatherMap API
        url = f"http://api.openweathermap.org/data/2.5/weather"
        params = {
            'q': city,
            'appid': api_key,
            'units': 'metric'
        }
        
        response = requests.get(url, params=params, timeout=10)
        
        if response.status_code == 200:
            data = response.json()
            weather_data = {
                "temperature": round(data['main']['temp']),
                "humidity": data['main']['humidity'],
                "description": data['weather'][0]['description'],
                "windSpeed": round(data['wind']['speed'] * 3.6),  # Convert m/s to km/h
                "visibility": round(data.get('visibility', 10000) / 1000),  # Convert m to km
                "city": data['name'],
                "country": data['sys']['country'],
                "mock": False,
                "success": True
            }
            # Cache the successful response
            _WEATHER_CACHE[cache_key] = (weather_data, current_time)
            print(f"🌤️ Weather data cached for {city}")
            return jsonify(weather_data)
        elif response.status_code == 404:
            # City not found
            return jsonify({
                "error": "City not found. Please check the spelling and try again.",
                "success": False,
                "city": city
            }), 404
        else:
            # Other API errors
            return jsonify({
                "error": f"Weather service temporarily unavailable (Error {response.status_code})",
                "success": False,
                "city": city
            }), response.status_code
            
    except Exception as e:
        # Return mock data if any error occurs
        mock_data = {
            "temperature": 28,
            "humidity": 75,
            "description": "Partly cloudy",
            "windSpeed": 12,
            "visibility": 10,
            "city": city,
            "mock": True,
            "error": str(e)
        }
        # Cache mock data too
        _WEATHER_CACHE[cache_key] = (mock_data, current_time)
        return jsonify(mock_data)

# Additional Admin API endpoints for the new admin panel
@views.route('/api/admin/stats')
@login_required
def admin_api_stats():
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        # Get basic statistics
        total_users = User.query.count()
        active_users = User.query.filter_by(is_active=True).count()
        
        # Mock data for now - in a real app, you'd query actual feedback and system health
        total_feedback = 0  # Would query feedback table
        system_health = "Good"
        
        return jsonify({
            "totalUsers": total_users,
            "activeUsers": active_users,
            "subscribedUsers": User.query.filter_by(subscribed=True).count()
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/users')
@login_required
def admin_api_users():
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        # Get all users
        users = User.query.all()
        users_data = []
        for user in users:
            users_data.append({
                "id": user.id,
                "email": user.email,
                "firstname": user.firstname,
                "lastname": user.lastname,
                "full_name": user.full_name,
                "contact": user.contact,
                "role": user.role,
                "is_active": user.is_active,
                "email_verified": user.email_verified,
                "subscribed": user.subscribed,
                "learning_level": user.learning_level,
                "created_at": user.created_at.isoformat() if user.created_at else None,
                "updated_at": user.updated_at.isoformat() if user.updated_at else None
            })
        
        # Get all admins and format them to match user structure with role="admin"
        admins = Admin.query.all()
        for admin in admins:
            users_data.append({
                "id": admin.id,
                "email": admin.email,
                "username": admin.username,
                "firstname": None,  # Admins don't have firstname/lastname, they have full_name
                "lastname": None,
                "full_name": admin.full_name,
                "contact": None,  # Admins don't have contact
                "role": "admin",  # Mark as admin role
                "is_active": admin.is_active,
                "email_verified": True,  # Admins are always verified
                "subscribed": False,  # Admins don't have subscriptions
                "learning_level": None,  # Admins don't have learning level
                "created_at": admin.created_at.isoformat() if admin.created_at else None,
                "updated_at": None,
                "is_super_admin": admin.is_super_admin  # Include super admin status
            })
        
        return jsonify(users_data)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/users/<int:user_id>', methods=['DELETE'])
@login_required
def admin_api_delete_user(user_id):
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        # Prevent deleting yourself
        if user_id == current_user.id:
            return jsonify({"error": "Cannot delete your own account"}), 403
        
        # Check if it's an Admin account
        admin = Admin.query.get(user_id)
        if admin:
            # It's an Admin account
            db.session.delete(admin)
            db.session.commit()
            return jsonify({"message": "Admin account deleted successfully"})
        
        # Check if it's a User account
        user = User.query.get(user_id)
        if user:
            if user.role == 'admin' and user.id != current_user.id:
                return jsonify({"error": "Cannot delete other admin users"}), 403
            
            # Manually delete related records to avoid foreign key constraint issues
            # Delete UserSubscription records first
            UserSubscription.query.filter_by(user_id=user_id).delete()
            
            # Delete ActivityLog records (they don't have CASCADE in the model)
            ActivityLog.query.filter_by(user_id=user_id).delete()
            
            # Commit the deletions before deleting the user
            db.session.commit()
            
            # Now delete the user - cascades will handle other related data (gardens, etc.)
            db.session.delete(user)
            db.session.commit()
            
            return jsonify({"message": "User deleted successfully"})
        
        # If neither found
        return jsonify({"error": "User or admin not found"}), 404
        
    except Exception as e:
        db.session.rollback()
        import traceback
        error_msg = str(e)
        print(f"Error deleting user/admin {user_id}: {error_msg}")
        print(traceback.format_exc())
        return jsonify({"error": f"Failed to delete: {error_msg}"}), 500

@views.route('/api/admin/users/<int:user_id>/status', methods=['PATCH'])
@login_required
def admin_api_toggle_user_status(user_id):
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        # Prevent deactivating yourself
        if user_id == current_user.id:
            return jsonify({"error": "Cannot modify your own account status"}), 403
        
        # Check if it's an Admin account
        admin = Admin.query.get(user_id)
        if admin:
            # It's an Admin account
            data = request.get_json()
            admin.is_active = data.get('is_active', not admin.is_active)
            db.session.commit()
            return jsonify({"message": "Admin status updated successfully"})
        
        # Check if it's a User account
        user = User.query.get(user_id)
        if user:
            if user.role == 'admin' and user.id != current_user.id:
                return jsonify({"error": "Cannot modify other admin users"}), 403
            
            data = request.get_json()
            user.is_active = data.get('is_active', not user.is_active)
            db.session.commit()
            return jsonify({"message": "User status updated successfully"})
        
        # If neither found
        return jsonify({"error": "User or admin not found"}), 404
        
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/users/<int:user_id>/subscription', methods=['PATCH'])
@login_required
def admin_api_toggle_user_subscription(user_id):
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        user = User.query.get_or_404(user_id)
        if user.role == 'admin':
            return jsonify({"error": "Cannot modify admin user subscriptions"}), 403
        
        data = request.get_json()
        new_subscribed_status = data.get('subscribed', not user.subscribed)
        was_subscribed = user.subscribed
        
        user.subscribed = new_subscribed_status
        
        # Update subscription plan based on subscription status
        if user.subscribed:
            user.subscription_plan = 'premium'
            # Reset free analysis usage when upgrading to premium
            # This gives premium users a fresh start with 10 free analyses
            usage = _get_or_create_ai_usage(user_id)
            old_free_used = usage.free_analyses_used or 0
            usage.free_analyses_used = 0
            print(f"🔄 RESET AI USAGE: Admin upgraded user {user_id} to premium - reset free_analyses_used from {old_free_used} to 0")
        else:
            user.subscription_plan = 'basic'
        
        # If downgrading from premium to basic, revert garden features
        if was_subscribed and not new_subscribed_status:
            from website.models import UserSubscription, Garden, GridSpace
            # Update active subscription status to cancelled
            active_subscription = UserSubscription.query.filter_by(
                user_id=user_id,
                status='active'
            ).first()
            
            if active_subscription:
                active_subscription.status = 'cancelled'
                active_subscription.updated_at = datetime.now(timezone.utc)
            
            # Revert garden features to basic
            user_gardens = Garden.query.filter_by(user_id=user_id).all()
            
            for garden in user_gardens:
                # Log the garden changes
                log_history_change(
                    table_name='garden',
                    record_id=garden.id,
                    action='UPDATE',
                    old_values={'grid_size': garden.grid_size, 'base_grid_spaces': garden.base_grid_spaces},
                    new_values={'grid_size': '3x3', 'base_grid_spaces': 9},
                    changed_by=f'admin_{current_user.id}'
                )
                
                # Revert to 3x3 grid for basic plan
                garden.grid_size = '3x3'
                garden.base_grid_spaces = 9
                garden.additional_spaces_purchased = 0
                
                # Remove any plants beyond the basic 9 spaces
                grid_spaces = GridSpace.query.filter_by(garden_id=garden.id).all()
                
                # Keep only the first 9 spaces (3x3 grid)
                spaces_to_remove = grid_spaces[9:]  # Remove spaces beyond 9
                for space in spaces_to_remove:
                    if space.plant_id:  # If there's a plant in this space
                        # Remove the plant from the space
                        space.plant_id = None
                        space.planting_date = None
                        space.notes = None
                    db.session.delete(space)
                
                # Update used grid spaces count
                remaining_spaces = GridSpace.query.filter_by(garden_id=garden.id).all()
                garden.used_grid_spaces = len([s for s in remaining_spaces if s.plant_id is not None])
            
            # Log the subscription cancellation
            log_subscription_activity(
                user_id=user_id,
                user_name=user.full_name,
                action='subscription_cancelled',
                plan_name='Premium Plan',
                amount=150.00,
                currency='PHP',
                payment_method='admin_action',
                status='cancelled',
                subscription_id=active_subscription.id if active_subscription else None
            )
            
            # Log the user activity
            log_activity(
                user_id=user_id,
                user_name=user.full_name,
                action='subscription_cancelled',
                description=f'Admin {current_user.full_name} cancelled Premium subscription and reverted to basic plan'
            )
            
        db.session.commit()
        return jsonify({"message": "User subscription updated successfully"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/learning-paths')
@login_required
def admin_api_learning_paths():
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    # Real learning paths data based on actual learning path files
    learning_paths = [
                    {
                        "id": 1,
                        "title": "Beginner Gardener",
                        "description": "Perfect for those just starting their gardening journey. Learn plant basics, soil compatibility, and essential gardening skills.",
                        "modules_count": 11,  # Based on actual modules in BeginnerLearningPath.jsx
                        "difficulty": "Beginner",
                        "duration": "4 weeks",
                        "is_active": True,
                        "created_at": "2024-01-15T10:00:00Z"
                    },
        {
            "id": 2,
            "title": "Intermediate Gardener", 
            "description": "For gardeners ready to take their skills to the next level. Advanced plant care, nutrition, and soil science.",
            "modules_count": 2,  # Based on actual modules in IntermediateLearningPath.jsx
            "difficulty": "Intermediate",
            "duration": "6 weeks",
            "is_active": True,
            "created_at": "2024-01-20T14:30:00Z"
        },
        {
            "id": 3,
            "title": "Expert Gardener",
            "description": "Advanced techniques for experienced gardeners. Professional pruning, soil analysis, and master-level skills.",
            "modules_count": 2,  # Based on actual modules in ExpertLearningPath.jsx
            "difficulty": "Expert",
            "duration": "8 weeks",
            "is_active": True,
            "created_at": "2024-01-25T09:15:00Z"
        }
    ]
    return jsonify(learning_paths)

@views.route('/api/learning-paths/<difficulty>')
def get_learning_path_content(difficulty):
    """Serve learning path content - prioritize saved modules from database, fallback to static modules"""
    try:
        # Check if the learning path is active by checking if any modules are active
        active_modules_count = LearningPathContent.query.filter_by(
            path_difficulty=difficulty,
            content_type='module',
            is_active=True
        ).count()
        
        # If no active modules exist, the path is deactivated
        if active_modules_count == 0:
            # Check if there are any modules at all (to distinguish between deactivated and empty)
            total_modules_count = LearningPathContent.query.filter_by(
                path_difficulty=difficulty,
                content_type='module'
            ).count()
            
            if total_modules_count > 0:
                # Path has been deactivated by admin
                return jsonify({
                    "error": "This learning path has been deactivated",
                    "is_active": False,
                    "message": "This learning path is currently unavailable. Please check back later."
                }), 403
        
        # First, get complete modules saved by admin (content_type='module')
        saved_modules = LearningPathContent.query.filter_by(
            path_difficulty=difficulty,
            content_type='module',
            is_active=True
        ).all()
        
        module_list = []
        
        # Parse saved modules from database
        for saved_module in saved_modules:
            try:
                module_data = json.loads(saved_module.content) if saved_module.content else {}
                if module_data:
                    module_list.append(module_data)
            except json.JSONDecodeError:
                print(f"Error parsing module data for {saved_module.module_id}")
                continue
        
        # If we have saved modules, ensure we have the 3 static modules (basic-information, lessons, quizzes)
        # Check if static modules exist, if not, they'll be added by the admin panel
        static_module_ids = ['basic-information', 'lessons', 'quizzes']
        existing_ids = [m.get('id') for m in module_list]
        
        # If we have saved modules, return them (they should include the static ones)
        if module_list:
            # Ensure all 3 static modules are present
            for static_id in static_module_ids:
                if static_id not in existing_ids:
                    # This shouldn't happen if admin auto-saved them, but handle gracefully
                    print(f"Warning: Static module {static_id} not found for {difficulty}")
            return jsonify(module_list)
        
        # Fallback: Get individual content items (lessons, quiz questions) for backward compatibility
        content_items = LearningPathContent.query.filter_by(
            path_difficulty=difficulty,
            is_active=True
        ).all()
        
        # Organize content by module
        modules = {}
        for item in content_items:
            module_id = item.module_id
            if module_id not in modules:
                modules[module_id] = {
                    'id': module_id,
                    'lessons': [],
                    'quiz': {'title': '', 'questions': []}
                }
            
            if item.content_type == 'lesson':
                modules[module_id]['lessons'].append({
                    'id': item.content_id,
                    'title': item.title,
                    'content': item.content,
                    'images': [{'url': item.media_url, 'description': item.media_description}] if item.media_url else [],
                    'videos': []
                })
            elif item.content_type == 'quiz_question':
                # Find or create the quiz question
                question_id = item.content_id
                question = next((q for q in modules[module_id]['quiz']['questions'] if q['id'] == question_id), None)
                
                if not question:
                    question = {
                        'id': question_id,
                        'question': item.title,
                        'options': [],
                        'correct': 0,
                        'explanation': '',
                        'image': item.media_url if item.media_type == 'image' else None,
                        'imageDescription': item.media_description
                    }
                    modules[module_id]['quiz']['questions'].append(question)
        
        # Convert to list format expected by frontend
        module_list = []
        for module_id, module_data in modules.items():
            module_list.append({
                'id': module_id,
                'title': f"{module_id.replace('-', ' ').title()} Module",
                'difficulty': difficulty,
                'estimatedTime': '30 min',
                'description': f'Learn about {module_id.replace("-", " ")}',
                'lessons': module_data['lessons'],
                'quiz': module_data['quiz']
            })
        
        # If no content found in database, return empty list to trigger fallback to static data in frontend
        # The frontend will use getModuleData() or the admin will auto-save static modules
        if not module_list:
            return jsonify([])
        
        return jsonify(module_list)
        
    except Exception as e:
        print(f"Error serving learning path content: {str(e)}")
        return jsonify({"error": "Failed to load learning path content"}), 500

@views.route('/api/admin/ai-data')
@login_required
def admin_api_ai_data():
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    # Mock data for AI recognition data
    ai_data = [
        {
            "id": 1,
            "name": "Plant Recognition Model",
            "description": "AI model for identifying plant species",
            "type": "Plant Recognition",
            "accuracy": 94.5,
            "updated_at": datetime.now().isoformat()
        },
        {
            "id": 2,
            "name": "Disease Detection Model",
            "description": "AI model for detecting plant diseases",
            "type": "Disease Detection", 
            "accuracy": 89.2,
            "updated_at": datetime.now().isoformat()
        }
    ]
    return jsonify(ai_data)

@views.route('/api/admin/notifications', methods=['GET'])
@login_required
def admin_api_notifications():
    """Get all notifications for admin management"""
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        notifications = Notification.query.order_by(Notification.created_at.desc()).all()
        return jsonify([n.to_dict() for n in notifications])
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/notifications', methods=['POST'])
@login_required
def admin_api_create_notification():
    """Create a new notification"""
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        data = request.get_json()
        
        # Validate required fields
        if not data.get('title') or not data.get('message'):
            return jsonify({"error": "Title and message are required"}), 400
        
        # Parse expiration date if provided
        expires_at = None
        if data.get('expires_at'):
            try:
                # Frontend sends ISO string with timezone (from toISOString())
                # Parse it directly - it's already in UTC format
                date_str = data['expires_at']
                if date_str.endswith('Z'):
                    expires_at = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
                elif '+' in date_str or date_str.count('-') >= 3:
                    # It's an ISO string with timezone
                    expires_at = datetime.fromisoformat(date_str)
                else:
                    # Fallback: treat as naive datetime and assume UTC
                    expires_at = datetime.fromisoformat(date_str).replace(tzinfo=timezone.utc)
            except Exception as e:
                print(f"Error parsing expires_at: {e}")
                pass
        
        # Handle created_by: if current_user is Admin (from admins table), 
        # we need to check if they have a corresponding User record
        # Otherwise, set to NULL since the foreign key references users.id
        created_by = None
        
        # Check if current_user is a User (from users table)
        if isinstance(current_user, User):
            # Verify the user exists in the database
            user_exists = User.query.get(current_user.id)
            if user_exists:
                created_by = current_user.id
        # If current_user is Admin (from admins table), try to find corresponding User by email
        else:
            # Try to find a User with the same email
            try:
                corresponding_user = User.query.filter_by(email=current_user.email).first()
                if corresponding_user:
                    created_by = corresponding_user.id
            except:
                pass
            # If no corresponding user found, set to NULL (foreign key allows NULL)
            # This is safe because created_by is nullable
        
        notification = Notification(
            title=data['title'],
            message=data['message'],
            type=data.get('type', 'Update'),
            priority=data.get('priority', 'Medium'),
            is_active=data.get('is_active', True),
            created_by=created_by,  # Will be NULL if admin has no corresponding user
            expires_at=expires_at
        )
        
        db.session.add(notification)
        db.session.commit()
        
        return jsonify(notification.to_dict()), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/notifications/<int:notification_id>', methods=['PUT'])
@login_required
def admin_api_update_notification(notification_id):
    """Update an existing notification"""
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        notification = Notification.query.get_or_404(notification_id)
        data = request.get_json()
        
        # Update fields
        if 'title' in data:
            notification.title = data['title']
        if 'message' in data:
            notification.message = data['message']
        if 'type' in data:
            notification.type = data['type']
        if 'priority' in data:
            notification.priority = data['priority']
        if 'is_active' in data:
            notification.is_active = data['is_active']
        if 'expires_at' in data:
            if data['expires_at']:
                try:
                    # Frontend sends ISO string with timezone (from toISOString())
                    # Parse it directly - it's already in UTC format
                    date_str = data['expires_at']
                    if date_str.endswith('Z'):
                        notification.expires_at = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
                    elif '+' in date_str or date_str.count('-') >= 3:
                        # It's an ISO string with timezone
                        notification.expires_at = datetime.fromisoformat(date_str)
                    else:
                        # Fallback: treat as naive datetime and assume UTC
                        notification.expires_at = datetime.fromisoformat(date_str).replace(tzinfo=timezone.utc)
                except Exception as e:
                    print(f"Error parsing expires_at: {e}")
                    notification.expires_at = None
            else:
                notification.expires_at = None
        
        notification.updated_at = datetime.now(timezone.utc)
        db.session.commit()
        
        return jsonify(notification.to_dict())
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/notifications/<int:notification_id>', methods=['DELETE'])
@login_required
def admin_api_delete_notification(notification_id):
    """Delete a notification"""
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        notification = Notification.query.get_or_404(notification_id)
        db.session.delete(notification)
        db.session.commit()
        
        return jsonify({"message": "Notification deleted successfully"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/notifications/<int:notification_id>/status', methods=['PATCH'])
@login_required
def admin_api_toggle_notification_status(notification_id):
    """Toggle notification active status"""
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        notification = Notification.query.get_or_404(notification_id)
        data = request.get_json()
        
        if 'is_active' in data:
            notification.is_active = data['is_active']
        
        notification.updated_at = datetime.now(timezone.utc)
        db.session.commit()
        
        return jsonify(notification.to_dict())
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@views.route('/api/notifications', methods=['GET'])
@login_required
def api_user_notifications():
    """Get active notifications for users"""
    try:
        now = datetime.now(timezone.utc)
        notifications = Notification.query.filter(
            Notification.is_active == True,
            db.or_(
                Notification.expires_at.is_(None),
                Notification.expires_at > now
            )
        ).order_by(
            # Sort by priority (High > Medium > Low), then by created_at (newest first)
            db.case(
                (Notification.priority == 'High', 3),
                (Notification.priority == 'Medium', 2),
                (Notification.priority == 'Low', 1),
                else_=0
            ).desc(),
            Notification.created_at.desc()
        ).all()
        
        return jsonify([n.to_dict() for n in notifications])
    except Exception as e:
        print(f"Error fetching user notifications: {str(e)}")
        return jsonify([])  # Return empty array on error instead of 500

@views.route('/api/admin/seasonal-content')
@login_required
def admin_api_seasonal_content():
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    # Mock data for seasonal content
    seasonal_content = [
        {
            "id": 1,
            "season": "Spring",
            "description": "Spring planting guide and tips",
            "month": "March",
            "region": "Northern Hemisphere",
            "updated_at": datetime.now().isoformat()
        },
        {
            "id": 2,
            "season": "Summer",
            "description": "Summer garden maintenance",
            "month": "June",
            "region": "Northern Hemisphere", 
            "updated_at": datetime.now().isoformat()
        }
    ]
    return jsonify(seasonal_content)

@views.route('/api/admin/feedback')
@login_required
def admin_api_feedback():
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        # Get all feedback from database
        feedbacks = Feedback.query.order_by(Feedback.created_at.desc()).all()
        feedback_list = [feedback.to_dict() for feedback in feedbacks]
        
        return jsonify(feedback_list)
    except Exception as e:
        return jsonify({"error": "Failed to fetch feedback. Please try again."}), 500

@views.route('/api/admin/reports')
@login_required
def admin_api_reports():
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    # Mock data for reports
    reports = []
    return jsonify(reports)

@views.route('/api/admin/reports/generate', methods=['POST'])
@login_required
def admin_api_generate_report():
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        data = request.get_json()
        report_type = data.get('type')
        date_range = data.get('dateRange', '30d')
        
        # Mock report generation
        report = {
            "id": 1,
            "name": f"{report_type.replace('-', ' ').title()} Report",
            "description": f"Generated report for {report_type} covering {date_range}",
            "type": report_type,
            "size": "2.3 MB",
            "created_at": datetime.now().isoformat()
        }
        
        return jsonify({"message": "Report generated successfully", "report": report})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/rollback/<content_type>/<int:content_id>', methods=['POST'])
@login_required
def admin_api_rollback_content(content_type, content_id):
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        # Mock rollback functionality
        return jsonify({"message": f"Content {content_type} with ID {content_id} rolled back successfully"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Additional admin endpoints for new pages
@views.route('/api/admin/subscription/stats')
@login_required
def admin_api_subscription_stats():
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        from datetime import datetime, timedelta, timezone
        from website.models import UserSubscription
        
        total_subscribers = User.query.filter_by(subscribed=True).count()
        active_subscriptions = User.query.filter_by(subscribed=True, is_active=True).count()
        monthly_revenue = total_subscribers * 150.00  # 150 PHP per subscriber
        total_users = User.query.count()
        subscription_rate = (total_subscribers / total_users * 100) if total_users > 0 else 0
        
        # Calculate Average Revenue Per User (ARPU)
        # ARPU = Total Revenue / Total Subscribers
        average_revenue_per_user = (monthly_revenue / total_subscribers) if total_subscribers > 0 else 0
        
        # Calculate new subscribers this month
        now = datetime.now(timezone.utc)
        start_of_month = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        new_subscribers_this_month = User.query.filter(
            User.subscribed == True,
            User.created_at >= start_of_month
        ).count()
        
        # Calculate churn rate (cancelled subscriptions / total subscriptions)
        total_subscription_records = UserSubscription.query.count()
        cancelled_subscriptions = UserSubscription.query.filter_by(status='cancelled').count()
        churn_rate = (cancelled_subscriptions / total_subscription_records * 100) if total_subscription_records > 0 else 0
        
        return jsonify({
            "totalSubscribers": total_subscribers,
            "activeSubscriptions": active_subscriptions,
            "monthlyRevenue": round(monthly_revenue, 2),
            "subscriptionRate": round(subscription_rate, 1),
            "averageRevenuePerUser": round(average_revenue_per_user, 2),
            "churnRate": round(churn_rate, 1),
            "newSubscribersThisMonth": new_subscribers_this_month,
            "totalRevenue": round(monthly_revenue, 2)  # For compatibility
        })
    except Exception as e:
        print(f"Error calculating subscription stats: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/subscription/subscribers')
@login_required
def admin_api_subscription_subscribers():
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        users = User.query.all()
        subscribers_data = []
        for user in users:
            subscribers_data.append({
                "id": user.id,
                "email": user.email,
                "firstname": user.firstname,
                "lastname": user.lastname,
                "full_name": user.full_name,
                "subscribed": user.subscribed,
                "is_active": user.is_active,
                "created_at": user.created_at.isoformat() if user.created_at else None
            })
        return jsonify(subscribers_data)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@views.route('/api/subscription/upgrade', methods=['POST'])
@login_required
def user_subscribe():
    """Handle user subscription upgrade"""
    try:
        data = request.get_json() or {}
        plan_type = data.get('plan_type', 'premium')
        payment_method = data.get('payment_method', 'demo')
        
        # For demo purposes, we'll simulate successful subscription
        print(f"💰 SUBSCRIPTION: User {current_user.id} upgrading to {plan_type} via {payment_method}")
        
        # Update user subscription status
        current_user.subscribed = True
        
        # Create subscription record in user_subscriptions table
        from website.models import UserSubscription, SubscriptionPlan
        
        # Get the premium plan
        premium_plan = SubscriptionPlan.query.filter_by(plan_name='Premium Plan').first()
        if not premium_plan:
            # Create premium plan if it doesn't exist
            premium_plan = SubscriptionPlan(
                plan_name='Premium Plan',
                plan_type='premium',
                price=150.00,
                currency='PHP',
                grid_planner_size='6x6',
                free_ai_analyses=20,
                free_plant_analyses=10,
                free_soil_analyses=10,
                additional_grid_cost=20.00,
                additional_ai_cost=25.00
            )
            db.session.add(premium_plan)
            db.session.flush()  # Get the ID
        
        # Create user subscription record
        from datetime import datetime, timedelta, timezone
        now = datetime.now(timezone.utc)
        subscription = UserSubscription(
            user_id=current_user.id,
            plan_id=premium_plan.id,
            start_date=now,
            end_date=now + timedelta(days=30),  # 30 days for demo
            status='active',
            payment_status='paid',
            total_paid=150.00,
            created_at=now  # Explicitly set created_at to current UTC time
        )
        db.session.add(subscription)
        
        # Log the subscription activity
        log_subscription_activity(
            user_id=current_user.id,
            user_name=current_user.full_name,
            action='subscription_created',
            plan_name='Premium Plan',
            amount=150.00,
            currency='PHP',
            payment_method=payment_method,
            status='active',
            subscription_id=subscription.id
        )
        
        # Log the user activity
        log_activity(
            user_id=current_user.id,
            user_name=current_user.full_name,
            action='subscription_upgrade',
            description='User upgraded to Premium plan'
        )
        
        # Expand all existing gardens from 3x3 to 6x6, preserving existing plants
        from website.models import Garden
        user_gardens = Garden.query.filter_by(user_id=current_user.id).all()
        
        for garden in user_gardens:
            # Only expand if garden is currently 3x3
            if garden.grid_size == '3x3' or garden.base_grid_spaces == 9:
                # Log the garden changes
                log_history_change(
                    table_name='garden',
                    record_id=garden.id,
                    action='UPDATE',
                    old_values={'grid_size': garden.grid_size, 'base_grid_spaces': garden.base_grid_spaces},
                    new_values={'grid_size': '6x6', 'base_grid_spaces': 36},
                    changed_by=f'user_{current_user.id}'
                )
                
                # Expand the grid from 3x3 to 6x6 (preserves existing plants)
                expand_garden_grid_to_premium(garden.id)
                
                # Update garden metadata
                garden.grid_size = '6x6'
                garden.base_grid_spaces = 36
                
                print(f"🌱 Expanded garden {garden.id} ({garden.name}) from 3x3 to 6x6")
        
        # Reset free analysis usage when upgrading to premium
        # This gives premium users a fresh start with 10 free analyses
        usage = _get_or_create_ai_usage(current_user.id)
        old_free_used = usage.free_analyses_used or 0
        usage.free_analyses_used = 0
        print(f"🔄 RESET AI USAGE: User {current_user.id} upgraded to premium - reset free_analyses_used from {old_free_used} to 0")
        
        db.session.commit()
        
        # Clear auth status cache to force refresh
        # Import the auth module to access the cache
        import website.auth as auth_module
        if hasattr(auth_module.auth_status, '_cache'):
            cache_key = f"auth_status_{current_user.id}"
            if cache_key in auth_module.auth_status._cache:
                del auth_module.auth_status._cache[cache_key]
                print(f"🔐 Cleared auth status cache for user {current_user.id}")
        
        print(f"💰 SUBSCRIPTION: Successfully upgraded user {current_user.id} to premium")
        
        return jsonify({
            "success": True,
            "message": "Subscription upgraded successfully!",
            "subscription": {
                "plan": plan_type,
                "status": "active",
                "start_date": subscription.start_date.isoformat(),
                "end_date": subscription.end_date.isoformat()
            }
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"Error processing subscription: {str(e)}")
        return jsonify({"success": False, "error": f"Subscription failed: {str(e)}"}), 500

@views.route('/api/subscription/downgrade', methods=['POST'])
@login_required
def user_downgrade_subscription():
    """Handle user subscription downgrade"""
    try:
        print(f"💰 SUBSCRIPTION DOWNGRADE: User {current_user.id} downgrading subscription")
        
        # Update user subscription status to basic
        current_user.subscribed = False
        current_user.subscription_plan = 'basic'
        
        # Update active subscription status to cancelled
        from website.models import UserSubscription
        active_subscription = UserSubscription.query.filter_by(
            user_id=current_user.id,
            status='active'
        ).first()
        
        if active_subscription:
            active_subscription.status = 'cancelled'
            active_subscription.updated_at = datetime.now(timezone.utc)
        
        # Log the subscription downgrade
        log_subscription_activity(
            user_id=current_user.id,
            user_name=current_user.full_name,
            action='subscription_downgraded',
            plan_name='Premium Plan',
            amount=150.00,
            currency='PHP',
            payment_method='demo',
            status='cancelled',
            subscription_id=active_subscription.id if active_subscription else None
        )
        
        # Log the user activity
        log_activity(
            user_id=current_user.id,
            user_name=current_user.full_name,
            action='subscription_downgraded',
            description='User downgraded from Premium subscription to basic plan'
        )
        
        # Revert garden features to basic
        from website.models import Garden, GridSpace
        user_gardens = Garden.query.filter_by(user_id=current_user.id).all()
        
        for garden in user_gardens:
            # Log the garden changes
            log_history_change(
                table_name='garden',
                record_id=garden.id,
                action='UPDATE',
                old_values={'grid_size': garden.grid_size, 'base_grid_spaces': garden.base_grid_spaces},
                new_values={'grid_size': '3x3', 'base_grid_spaces': 9},
                changed_by=f'user_{current_user.id}'
            )
            
            # Revert to 3x3 grid for basic plan
            garden.grid_size = '3x3'
            garden.base_grid_spaces = 9
            garden.additional_spaces_purchased = 0
            
            # Remove any plants beyond the basic 9 spaces
            grid_spaces = GridSpace.query.filter_by(garden_id=garden.id).all()
            
            # Keep only the first 9 spaces (3x3 grid)
            spaces_to_remove = grid_spaces[9:]  # Remove spaces beyond 9
            for space in spaces_to_remove:
                if space.plant_id:  # If there's a plant in this space
                    # Remove the plant from the space
                    space.plant_id = None
                    space.planting_date = None
                    space.notes = None
                db.session.delete(space)
            
            # Update used grid spaces count
            remaining_spaces = GridSpace.query.filter_by(garden_id=garden.id).all()
            garden.used_grid_spaces = len([s for s in remaining_spaces if s.plant_id is not None])
        
        db.session.commit()
        
        # Clear auth status cache to force refresh
        import website.auth as auth_module
        if hasattr(auth_module.auth_status, '_cache'):
            cache_key = f"auth_status_{current_user.id}"
            if cache_key in auth_module.auth_status._cache:
                del auth_module.auth_status._cache[cache_key]
                print(f"🔐 Cleared auth status cache for user {current_user.id}")
        
        print(f"✅ SUBSCRIPTION DOWNGRADED: User {current_user.id} reverted to basic plan")
        
        return jsonify({
            "success": True,
            "message": "Subscription downgraded successfully. You will retain Premium access until the end of your billing period.",
            "subscription_plan": "basic"
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"❌ SUBSCRIPTION DOWNGRADE ERROR: {str(e)}")
        
        return jsonify({"success": False, "error": f"Subscription downgrade failed: {str(e)}"}), 500

@views.route('/api/subscription/details', methods=['GET'])
@login_required
def get_subscription_details():
    """Get current user's subscription details"""
    try:
        from website.models import UserSubscription, SubscriptionPlan
        from datetime import datetime, timedelta, timezone
        
        # Get the user's active subscription
        active_subscription = UserSubscription.query.filter_by(
            user_id=current_user.id,
            status='active'
        ).order_by(UserSubscription.created_at.desc()).first()
        
        if not active_subscription:
            return jsonify({
                "success": False,
                "message": "No active subscription found"
            }), 404
        
        # Get the plan details
        plan = SubscriptionPlan.query.get(active_subscription.plan_id)
        
        if not plan:
            return jsonify({
                "success": False,
                "message": "Plan not found"
            }), 404
        
        # Get start_date (use created_at if start_date is not set)
        start_date = active_subscription.start_date or active_subscription.created_at
        if not start_date:
            return jsonify({
                "success": False,
                "message": "Subscription start date not found"
            }), 404
        
        # Ensure start_date is timezone-aware
        if start_date.tzinfo is None:
            start_date = start_date.replace(tzinfo=timezone.utc)
        
        # Calculate next billing date: add 1 month to start_date
        # Handle month boundaries correctly
        year = start_date.year
        month = start_date.month
        day = start_date.day
        
        # Add 1 month
        if month == 12:
            year += 1
            month = 1
        else:
            month += 1
        
        # Handle edge case: if day doesn't exist in next month (e.g., Jan 31 -> Feb 31)
        # Use the last day of the month instead
        try:
            next_billing = datetime(year, month, day, tzinfo=timezone.utc)
        except ValueError:
            # Day doesn't exist in next month, use last day of month
            # Get first day of month after next month, then subtract 1 day
            if month == 12:
                next_month = datetime(year + 1, 1, 1, tzinfo=timezone.utc)
            else:
                next_month = datetime(year, month + 1, 1, tzinfo=timezone.utc)
            # Last day of target month
            last_day = (next_month - timedelta(days=1)).day
            next_billing = datetime(year, month, last_day, tzinfo=timezone.utc)
        
        return jsonify({
            "success": True,
            "subscription": {
                "plan": plan.plan_name,
                "status": active_subscription.status,
                "startDate": start_date.isoformat(),
                "nextBillingDate": next_billing.isoformat(),
                "amount": float(plan.price),
                "currency": plan.currency or "PHP",
                "paymentMethod": "GCash",  # Default, can be enhanced later
                "autoRenew": active_subscription.status == 'active',
                "features": {
                    "gridPlanner6x6": plan.grid_planner_size == '6x6',
                    "aiAnalyses20": plan.free_ai_analyses >= 20,
                    "plantAnalyses10": plan.free_plant_analyses >= 10,
                    "soilAnalyses10": plan.free_soil_analyses >= 10,
                    "advancedPlantId": True,
                    "detailedSoilAnalysis": True,
                    "personalizedRecommendations": True,
                    "prioritySupport": True
                }
            }
        })
    except Exception as e:
        print(f"Error fetching subscription details: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@views.route('/api/subscription/cancel', methods=['POST'])
@login_required
def user_cancel_subscription():
    """Handle user subscription cancellation"""
    try:
        print(f"💰 SUBSCRIPTION CANCELLATION: User {current_user.id} cancelling subscription")
        
        # Update user subscription status to basic
        current_user.subscribed = False
        current_user.subscription_plan = 'basic'
        
        # Update active subscription status to cancelled
        from website.models import UserSubscription
        active_subscription = UserSubscription.query.filter_by(
            user_id=current_user.id,
            status='active'
        ).first()
        
        if active_subscription:
            active_subscription.status = 'cancelled'
            active_subscription.updated_at = datetime.now(timezone.utc)
        
        # Log the subscription cancellation
        log_subscription_activity(
            user_id=current_user.id,
            user_name=current_user.full_name,
            action='subscription_cancelled',
            plan_name='Premium Plan',
            amount=150.00,
            currency='PHP',
            payment_method='demo',
            status='cancelled',
            subscription_id=active_subscription.id if active_subscription else None
        )
        
        # Log the user activity
        log_activity(
            user_id=current_user.id,
            user_name=current_user.full_name,
            action='subscription_cancelled',
            description='User cancelled Premium subscription and reverted to basic plan'
        )
        
        # Revert garden features to basic
        from website.models import Garden
        user_gardens = Garden.query.filter_by(user_id=current_user.id).all()
        
        for garden in user_gardens:
            # Log the garden changes
            log_history_change(
                table_name='garden',
                record_id=garden.id,
                action='UPDATE',
                old_values={'grid_size': garden.grid_size, 'base_grid_spaces': garden.base_grid_spaces},
                new_values={'grid_size': '3x3', 'base_grid_spaces': 9},
                changed_by=f'user_{current_user.id}'
            )
            
            # Revert to 3x3 grid for basic plan
            garden.grid_size = '3x3'
            garden.base_grid_spaces = 9
            garden.additional_spaces_purchased = 0
            
            # Remove any plants beyond the basic 9 spaces
            from website.models import GridSpace
            grid_spaces = GridSpace.query.filter_by(garden_id=garden.id).all()
            
            # Keep only the first 9 spaces (3x3 grid)
            spaces_to_remove = grid_spaces[9:]  # Remove spaces beyond 9
            for space in spaces_to_remove:
                if space.plant_id:  # If there's a plant in this space
                    # Remove the plant from the space
                    space.plant_id = None
                    space.planting_date = None
                    space.notes = None
                db.session.delete(space)
            
            # Update used grid spaces count
            remaining_spaces = GridSpace.query.filter_by(garden_id=garden.id).all()
            garden.used_grid_spaces = len([s for s in remaining_spaces if s.plant_id is not None])
        
        db.session.commit()
        
        print(f"✅ SUBSCRIPTION CANCELLED: User {current_user.id} reverted to basic plan")
        
        return jsonify({
            "success": True,
            "message": "Subscription cancelled successfully. You have been reverted to the basic plan.",
            "subscription_plan": "basic"
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"❌ SUBSCRIPTION CANCELLATION ERROR: {str(e)}")
        
        return jsonify({"success": False, "error": f"Subscription cancellation failed: {str(e)}"}), 500

# Reports API endpoints
@views.route('/api/admin/activity-logs')
@login_required
def admin_api_activity_logs():
    """Get activity logs for admin reports"""
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        date_range = request.args.get('date_range', '7d')
        
        # Calculate date filter
        from datetime import datetime, timedelta
        now = datetime.now()
        if date_range == '1d':
            start_date = now - timedelta(days=1)
        elif date_range == '7d':
            start_date = now - timedelta(days=7)
        elif date_range == '30d':
            start_date = now - timedelta(days=30)
        elif date_range == '90d':
            start_date = now - timedelta(days=90)
        else:
            start_date = now - timedelta(days=7)
        
        # Get real activity logs from memory cache
        activity_logs = _ACTIVITY_LOGS.copy()
        
        # Add some mock data if no real logs exist yet
        if not activity_logs:
            activity_logs = [
                {
                    'id': 1,
                    'user_id': 1,
                    'user_name': 'John Doe',
                    'action': 'user_login',
                    'description': 'User logged in successfully',
                    'ip_address': '192.168.1.100',
                    'user_agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)',
                    'timestamp': (now - timedelta(minutes=30)).isoformat(),
                    'status': 'success'
                },
                {
                    'id': 2,
                    'user_id': 2,
                    'user_name': 'Jane Smith',
                    'action': 'subscription_upgrade',
                    'description': 'User upgraded to Premium plan',
                    'ip_address': '192.168.1.101',
                    'user_agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7)',
                    'timestamp': (now - timedelta(hours=2)).isoformat(),
                    'status': 'success'
                }
            ]
        
        return jsonify({
            "success": True,
            "logs": activity_logs,
            "total": len(activity_logs),
            "date_range": date_range
        })
        
    except Exception as e:
        print(f"Error fetching activity logs: {str(e)}")
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/history-logs')
@login_required
def admin_api_history_logs():
    """Get history logs for admin reports"""
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        date_range = request.args.get('date_range', '7d')
        
        # Calculate date filter
        from datetime import datetime, timedelta
        now = datetime.now()
        if date_range == '1d':
            start_date = now - timedelta(days=1)
        elif date_range == '7d':
            start_date = now - timedelta(days=7)
        elif date_range == '30d':
            start_date = now - timedelta(days=30)
        elif date_range == '90d':
            start_date = now - timedelta(days=90)
        else:
            start_date = now - timedelta(days=7)
        
        # Get real history logs from memory cache
        history_logs = _HISTORY_LOGS.copy()
        
        # Filter by date range
        filtered_logs = []
        for log in history_logs:
            try:
                log_timestamp = datetime.fromisoformat(log.get('timestamp', ''))
                if log_timestamp >= start_date:
                    filtered_logs.append(log)
            except (ValueError, TypeError):
                # If timestamp parsing fails, include the log anyway
                filtered_logs.append(log)
        
        # Sort by timestamp (newest first)
        filtered_logs.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
        
        # Add some mock data if no real logs exist yet
        if not filtered_logs:
            filtered_logs = [
                {
                    'id': 1,
                    'table_name': 'users',
                    'record_id': 1,
                    'action': 'UPDATE',
                    'old_values': {'full_name': 'John Doe'},
                    'new_values': {'full_name': 'John Smith'},
                    'changed_by': 'admin',
                    'timestamp': (now - timedelta(hours=1)).isoformat(),
                    'field_changes': ['full_name']
                }
            ]
        
        return jsonify({
            "success": True,
            "logs": filtered_logs,
            "total": len(filtered_logs),
            "date_range": date_range
        })
        
    except Exception as e:
        print(f"Error fetching history logs: {str(e)}")
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/subscription-logs')
@login_required
def admin_api_subscription_logs():
    """Get subscription logs for admin reports"""
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        date_range = request.args.get('date_range', '7d')
        
        # Calculate date filter
        from datetime import datetime, timedelta
        now = datetime.now()
        if date_range == '1d':
            start_date = now - timedelta(days=1)
        elif date_range == '7d':
            start_date = now - timedelta(days=7)
        elif date_range == '30d':
            start_date = now - timedelta(days=30)
        elif date_range == '90d':
            start_date = now - timedelta(days=90)
        else:
            start_date = now - timedelta(days=7)
        
        # Get real subscription logs from memory cache
        subscription_logs = _SUBSCRIPTION_LOGS.copy()
        
        # Add some mock data if no real logs exist yet
        if not subscription_logs:
            subscription_logs = [
                {
                    'id': 1,
                    'user_id': 2,
                    'user_name': 'Jane Smith',
                    'action': 'subscription_created',
                    'plan_name': 'Premium Plan',
                    'amount': 150.00,
                    'currency': 'PHP',
                    'payment_method': 'gcash',
                    'status': 'active',
                    'start_date': (now - timedelta(days=7)).isoformat(),
                    'end_date': (now + timedelta(days=23)).isoformat(),
                    'timestamp': (now - timedelta(days=7)).isoformat()
                }
            ]
        
        return jsonify({
            "success": True,
            "logs": subscription_logs,
            "total": len(subscription_logs),
            "date_range": date_range
        })
        
    except Exception as e:
        print(f"Error fetching subscription logs: {str(e)}")
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/reports/summary')
@login_required
def admin_api_reports_summary():
    """Get summary statistics for reports"""
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        date_range = request.args.get('date_range', '7d')
        
        # Calculate real summary statistics from logged data
        total_activities = len(_ACTIVITY_LOGS)
        total_history = len(_HISTORY_LOGS)
        total_subscriptions = len(_SUBSCRIPTION_LOGS)
        active_subscriptions = len([log for log in _SUBSCRIPTION_LOGS if log.get('status') == 'active'])
        cancelled_subscriptions = len([log for log in _SUBSCRIPTION_LOGS if log.get('status') == 'cancelled'])
        
        summary_stats = {
            'totalActivities': total_activities,
            'totalHistoryRecords': total_history,
            'totalSubscriptions': total_subscriptions,
            'activeSubscriptions': active_subscriptions,
            'cancelledSubscriptions': cancelled_subscriptions,
            'date_range': date_range
        }
        
        return jsonify({
            "success": True,
            **summary_stats
        })
        
    except Exception as e:
        print(f"Error fetching reports summary: {str(e)}")
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/subscription/<int:user_id>/toggle', methods=['PATCH'])
@login_required
def admin_api_toggle_subscription(user_id):
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        user = User.query.get_or_404(user_id)
        data = request.get_json()
        new_subscribed_status = data.get('subscribed', not user.subscribed)
        was_subscribed = user.subscribed
        
        user.subscribed = new_subscribed_status
        
        # Update subscription plan based on subscription status
        if user.subscribed:
            user.subscription_plan = 'premium'
            # Reset free analysis usage when upgrading to premium
            # This gives premium users a fresh start with 10 free analyses
            usage = _get_or_create_ai_usage(user_id)
            old_free_used = usage.free_analyses_used or 0
            usage.free_analyses_used = 0
            print(f"🔄 RESET AI USAGE: Admin upgraded user {user_id} to premium - reset free_analyses_used from {old_free_used} to 0")
        else:
            user.subscription_plan = 'basic'
        
        # If downgrading from premium to basic, revert garden features
        if was_subscribed and not new_subscribed_status:
            from website.models import UserSubscription, Garden, GridSpace
            # Update active subscription status to cancelled
            active_subscription = UserSubscription.query.filter_by(
                user_id=user_id,
                status='active'
            ).first()
            
            if active_subscription:
                active_subscription.status = 'cancelled'
                active_subscription.updated_at = datetime.now(timezone.utc)
            
            # Revert garden features to basic
            user_gardens = Garden.query.filter_by(user_id=user_id).all()
            
            for garden in user_gardens:
                # Log the garden changes
                log_history_change(
                    table_name='garden',
                    record_id=garden.id,
                    action='UPDATE',
                    old_values={'grid_size': garden.grid_size, 'base_grid_spaces': garden.base_grid_spaces},
                    new_values={'grid_size': '3x3', 'base_grid_spaces': 9},
                    changed_by=f'admin_{current_user.id}'
                )
                
                # Revert to 3x3 grid for basic plan
                garden.grid_size = '3x3'
                garden.base_grid_spaces = 9
                garden.additional_spaces_purchased = 0
                
                # Remove any plants beyond the basic 9 spaces
                grid_spaces = GridSpace.query.filter_by(garden_id=garden.id).all()
                
                # Keep only the first 9 spaces (3x3 grid)
                spaces_to_remove = grid_spaces[9:]  # Remove spaces beyond 9
                for space in spaces_to_remove:
                    if space.plant_id:  # If there's a plant in this space
                        # Remove the plant from the space
                        space.plant_id = None
                        space.planting_date = None
                        space.notes = None
                    db.session.delete(space)
                
                # Update used grid spaces count
                remaining_spaces = GridSpace.query.filter_by(garden_id=garden.id).all()
                garden.used_grid_spaces = len([s for s in remaining_spaces if s.plant_id is not None])
            
            # Log the subscription cancellation
            log_subscription_activity(
                user_id=user_id,
                user_name=user.full_name,
                action='subscription_cancelled',
                plan_name='Premium Plan',
                amount=150.00,
                currency='PHP',
                payment_method='admin_action',
                status='cancelled',
                subscription_id=active_subscription.id if active_subscription else None
            )
            
            # Log the user activity
            log_activity(
                user_id=user_id,
                user_name=user.full_name,
                action='subscription_cancelled',
                description=f'Admin {current_user.full_name} cancelled Premium subscription and reverted to basic plan'
            )
        
        db.session.commit()
        return jsonify({"message": "Subscription status updated successfully"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/subscription/recent-activity')
@login_required
def admin_api_subscription_recent_activity():
    """Get recent subscription activities for admin dashboard"""
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        from website.models import UserSubscription, SubscriptionPlan, User
        from datetime import datetime, timedelta
        
        # Get recent subscriptions (last 10, ordered by created_at)
        recent_subscriptions = UserSubscription.query.order_by(
            UserSubscription.created_at.desc()
        ).limit(10).all()
        
        activities = []
        for sub in recent_subscriptions:
            user = User.query.get(sub.user_id)
            plan = SubscriptionPlan.query.get(sub.plan_id)
            
            if not user or not plan:
                continue
            
            # Always show subscription creation event (using created_at)
            if sub.created_at:
                created_timestamp = sub.created_at
                if created_timestamp.tzinfo is None:
                    created_timestamp = created_timestamp.replace(tzinfo=timezone.utc)
                
                activities.append({
                    'id': f"{sub.id}_created",
                    'type': 'subscription_created',
                    'message': f'New {plan.plan_name} subscription activated',
                    'user': f"{user.firstname} {user.lastname}",
                    'user_email': user.email,
                    'plan_name': plan.plan_name,
                    'status': 'active',
                    'payment_status': sub.payment_status,
                    'timestamp': created_timestamp.isoformat(),
                    'icon_type': 'crown',
                    'color': 'text-yellow-600'
                })
            
            # If subscription was cancelled, also show cancellation event (using updated_at)
            if sub.status == 'cancelled' and sub.updated_at and sub.updated_at != sub.created_at:
                cancelled_timestamp = sub.updated_at
                if cancelled_timestamp.tzinfo is None:
                    cancelled_timestamp = cancelled_timestamp.replace(tzinfo=timezone.utc)
                
                activities.append({
                    'id': f"{sub.id}_cancelled",
                    'type': 'subscription_cancelled',
                    'message': f'{plan.plan_name} subscription cancelled',
                    'user': f"{user.firstname} {user.lastname}",
                    'user_email': user.email,
                    'plan_name': plan.plan_name,
                    'status': 'cancelled',
                    'payment_status': sub.payment_status,
                    'timestamp': cancelled_timestamp.isoformat(),
                    'icon_type': 'xcircle',
                    'color': 'text-red-600'
                })
            
            # If payment failed, show payment failure event
            if sub.payment_status == 'failed' and sub.updated_at and sub.updated_at != sub.created_at:
                failed_timestamp = sub.updated_at
                if failed_timestamp.tzinfo is None:
                    failed_timestamp = failed_timestamp.replace(tzinfo=timezone.utc)
                
                activities.append({
                    'id': f"{sub.id}_failed",
                    'type': 'payment_failed',
                    'message': f'Payment failed for {plan.plan_name} subscription',
                    'user': f"{user.firstname} {user.lastname}",
                    'user_email': user.email,
                    'plan_name': plan.plan_name,
                    'status': sub.status,
                    'payment_status': 'failed',
                    'timestamp': failed_timestamp.isoformat(),
                    'icon_type': 'alert',
                    'color': 'text-orange-600'
                })
        
        # Sort activities by timestamp (most recent first)
        activities.sort(key=lambda x: x['timestamp'], reverse=True)
        
        # Limit to 10 most recent activities
        activities = activities[:10]
        
        return jsonify({
            "success": True,
            "activities": activities,
            "total": len(activities)
        })
    except Exception as e:
        print(f"Error fetching recent subscription activity: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/subscription/plans')
@login_required
def admin_api_subscription_plans():
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        from website.models import SubscriptionPlan
        plans = SubscriptionPlan.query.filter_by(is_active=True).all()
        return jsonify([plan.to_dict() for plan in plans])
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/subscription/plans/<int:plan_id>', methods=['PUT'])
@login_required
def admin_api_update_subscription_plan(plan_id):
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        from website.models import SubscriptionPlan
        plan = SubscriptionPlan.query.get_or_404(plan_id)
        data = request.get_json()
        
        if 'plan_name' in data:
            plan.plan_name = data['plan_name']
        if 'price' in data:
            plan.price = data['price']
        if 'currency' in data:
            plan.currency = data['currency']
        if 'grid_planner_size' in data:
            plan.grid_planner_size = data['grid_planner_size']
        if 'free_ai_analyses' in data:
            plan.free_ai_analyses = data['free_ai_analyses']
        if 'free_plant_analyses' in data:
            plan.free_plant_analyses = data['free_plant_analyses']
        if 'free_soil_analyses' in data:
            plan.free_soil_analyses = data['free_soil_analyses']
        if 'additional_grid_cost' in data:
            plan.additional_grid_cost = data['additional_grid_cost']
        if 'additional_ai_cost' in data:
            plan.additional_ai_cost = data['additional_ai_cost']
        
        db.session.commit()
        return jsonify({"message": "Plan updated successfully", "plan": plan.to_dict()})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/subscription/analytics/<plan_identifier>')
@login_required
def admin_api_subscription_analytics(plan_identifier):
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        from website.models import SubscriptionPlan, UserSubscription, User
        from sqlalchemy import func
        from datetime import datetime, timedelta
        
        # Try to get plan by ID first, then by plan_type
        try:
            plan_id = int(plan_identifier)
            plan = SubscriptionPlan.query.get(plan_id)
        except ValueError:
            # If not a number, treat as plan_type
            plan = SubscriptionPlan.query.filter_by(plan_type=plan_identifier).first()
        
        if not plan:
            # Return default analytics if plan not found
            total_subscribers = User.query.filter_by(subscribed=True).count()
            return jsonify({
                "totalSubscribers": total_subscribers,
                "monthlyRevenue": total_subscribers * 150,
                "churnRate": 3.2,
                "conversionRate": 12.5,
                "averageRevenuePerUser": 150,
                "growthRate": 15.3
            })
        
        # Get subscribers for this plan
        subscriptions = UserSubscription.query.filter_by(plan_id=plan.id, status='active').all()
        total_subscribers = len(subscriptions)
        
        # Calculate monthly revenue
        monthly_revenue = total_subscribers * float(plan.price)
        
        # Calculate churn rate (simplified)
        total_users = User.query.count()
        subscribed_users = User.query.filter_by(subscribed=True).count()
        churn_rate = 3.2  # Mock value
        
        # Calculate conversion rate
        conversion_rate = (subscribed_users / total_users * 100) if total_users > 0 else 0
        
        # Calculate ARPU
        arpu = float(plan.price)
        
        # Calculate growth rate (mock)
        growth_rate = 15.3
        
        return jsonify({
            "totalSubscribers": total_subscribers,
            "monthlyRevenue": round(monthly_revenue, 2),
            "churnRate": round(churn_rate, 1),
            "conversionRate": round(conversion_rate, 1),
            "averageRevenuePerUser": round(arpu, 2),
            "growthRate": round(growth_rate, 1)
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Learning Path Management API endpoints
# Note: Learning paths are fixed (Beginner, Intermediate, Expert) - no creation/deletion

@views.route('/api/admin/learning-paths/<int:path_id>', methods=['PUT'])
@login_required
def admin_api_update_learning_path(path_id):
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        data = request.get_json()
        
        # Validate required fields
        required_fields = ['title', 'description', 'difficulty', 'duration', 'modules_count']
        for field in required_fields:
            if not data.get(field):
                return jsonify({"error": f"Missing required field: {field}"}), 400
        
        # In a real app, you would:
        # 1. Update the LearningPath record in the database
        # 2. Update the corresponding learning path file (BeginnerLearningPath.jsx, etc.)
        # 3. Sync the changes with the user interface
        # 4. Clear any cached data
        
        # For now, return success with the updated learning path data
        learning_path = {
            "id": path_id,
            "title": data.get('title'),
            "description": data.get('description'),
            "difficulty": data.get('difficulty'),
            "duration": data.get('duration'),
            "modules_count": data.get('modules_count'),
            "is_active": data.get('is_active', True),
            "video": data.get('video'),
            "updated_at": datetime.now().isoformat()
        }
        
        # TODO: In a real implementation, you would:
        # - Update database record
        # - Update the corresponding learning path file
        # - Clear user progress cache if needed
        # - Notify users of content updates
        
        return jsonify({
            "message": "Learning path updated successfully", 
            "learning_path": learning_path,
            "note": "In production, this would update the actual learning path files and database"
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Learning paths cannot be deleted - they are fixed (Beginner, Intermediate, Expert)

@views.route('/api/admin/learning-paths/<path_id>/status', methods=['PATCH'])
@login_required
def admin_api_toggle_learning_path_status(path_id):
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        data = request.get_json()
        is_active = data.get('is_active', True)
        
        # Map path_id to difficulty
        difficulty_map = {
            '1': 'Beginner',
            '2': 'Intermediate', 
            '3': 'Expert',
            'beginner': 'Beginner',
            'intermediate': 'Intermediate',
            'expert': 'Expert'
        }
        
        difficulty = difficulty_map.get(str(path_id).lower())
        if not difficulty:
            return jsonify({"error": "Invalid learning path ID"}), 400
        
        # Update all LearningPathContent records for this difficulty
        # This effectively activates/deactivates the entire learning path
        updated_count = LearningPathContent.query.filter_by(
            path_difficulty=difficulty
        ).update({'is_active': is_active})
        
        db.session.commit()
        
        print(f"✅ Updated {updated_count} modules for {difficulty} learning path - is_active: {is_active}")
        
        return jsonify({
            "success": True,
            "message": f"Learning path {difficulty} {'activated' if is_active else 'deactivated'} successfully",
            "updated_modules": updated_count,
            "is_active": is_active
        })
    except Exception as e:
        db.session.rollback()
        print(f"❌ Error updating learning path status: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# Save Module API (POST for create)
@views.route('/api/admin/learning-paths/<difficulty>/modules', methods=['POST'])
@login_required
def admin_api_create_module(difficulty):
    """Create a new module for a learning path"""
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        data = request.get_json()
        module_id = data.get('id')
        module_data = data.get('module')
        
        if not module_id or not module_data:
            return jsonify({"error": "Module ID and data required"}), 400
        
        # Check if module already exists
        existing = LearningPathContent.query.filter_by(
            path_difficulty=difficulty,
            module_id=module_id,
            content_type='module'
        ).first()
        
        if existing:
            return jsonify({"error": "Module already exists. Use PUT to update."}), 400
        
        # Create new module record
        new_module = LearningPathContent(
            path_difficulty=difficulty,
            module_id=module_id,
            content_type='module',
            content_id=0,
            title=module_data.get('title', ''),
            content=json.dumps(module_data),
            is_active=True
        )
        db.session.add(new_module)
        db.session.commit()
        
        return jsonify({
            "message": "Module created successfully",
            "module_id": module_id
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

# Update Module API (PUT for update)
@views.route('/api/admin/learning-paths/<difficulty>/modules', methods=['PUT'])
@login_required
def admin_api_update_module(difficulty):
    """Update an existing module for a learning path"""
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        data = request.get_json()
        module_id = data.get('id')
        module_data = data.get('module')
        
        if not module_id or not module_data:
            return jsonify({"error": "Module ID and data required"}), 400
        
        # Find existing module
        existing = LearningPathContent.query.filter_by(
            path_difficulty=difficulty,
            module_id=module_id,
            content_type='module'
        ).first()
        
        if existing:
            # Update existing module
            existing.content = json.dumps(module_data)
            existing.title = module_data.get('title', '')
            existing.updated_at = datetime.now(timezone.utc)
        else:
            # Create if doesn't exist (for backward compatibility)
            existing = LearningPathContent(
                path_difficulty=difficulty,
                module_id=module_id,
                content_type='module',
                content_id=0,
                title=module_data.get('title', ''),
                content=json.dumps(module_data),
                is_active=True
            )
            db.session.add(existing)
        
        db.session.commit()
        
        return jsonify({
            "message": "Module updated successfully",
            "module_id": module_id
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

# Delete Module API
@views.route('/api/admin/learning-paths/<difficulty>/modules/<module_id>', methods=['DELETE'])
@login_required
def admin_api_delete_module(difficulty, module_id):
    """Delete a module from a learning path"""
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        # Delete the module record
        module = LearningPathContent.query.filter_by(
            path_difficulty=difficulty,
            module_id=module_id,
            content_type='module'
        ).first()
        
        if module:
            db.session.delete(module)
            db.session.commit()
            return jsonify({"message": "Module deleted successfully"})
        else:
            return jsonify({"error": "Module not found"}), 404
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@views.route('/api/admin/upload', methods=['POST'])
@login_required
def admin_api_upload_file():
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        file_type = request.form.get('type', 'unknown')
        
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        # Validate file type
        allowed_image_types = {'image/jpeg', 'image/jpg', 'image/png', 'image/gif', 'image/webp'}
        allowed_video_types = {'video/mp4', 'video/avi', 'video/mov', 'video/wmv', 'video/webm'}
        
        if file_type == 'image' and file.content_type not in allowed_image_types:
            return jsonify({"error": "Invalid image format. Allowed: JPEG, PNG, GIF, WebP"}), 400
        
        if file_type == 'video' and file.content_type not in allowed_video_types:
            return jsonify({"error": "Invalid video format. Allowed: MP4, AVI, MOV, WMV, WebM"}), 400
        
        # Create upload directory if it doesn't exist
        upload_dir = os.path.join(os.getcwd(), 'uploads', 'learning_paths', file_type)
        os.makedirs(upload_dir, exist_ok=True)
        
        # Generate unique filename
        import time
        import uuid
        timestamp = int(time.time())
        unique_id = str(uuid.uuid4())[:8]
        
        # Get file extension
        file_extension = os.path.splitext(file.filename)[1].lower()
        if not file_extension:
            # Default extensions based on content type
            if file.content_type.startswith('image/'):
                file_extension = '.jpg'
            elif file.content_type.startswith('video/'):
                file_extension = '.mp4'
            else:
                file_extension = '.bin'
        
        filename = f"learning_{timestamp}_{unique_id}{file_extension}"
        file_path = os.path.join(upload_dir, filename)
        
        # Save the file
        file.save(file_path)
        
        # Create relative URL for frontend
        file_url = f"/uploads/learning_paths/{file_type}/{filename}"
        
        # Get file size
        file_size = os.path.getsize(file_path)
        
        # Save to database if this is for learning path content
        if 'learning_path' in request.form:
            
            # Get the learning path details from form
            path_difficulty = request.form.get('path_difficulty', 'Beginner')
            module_id = request.form.get('module_id', 'plant-basics')
            content_type = request.form.get('content_type', 'quiz_question')
            content_id = request.form.get('content_id', 1)
            question_number = request.form.get('question_number')
            title = request.form.get('title', '')
            description = request.form.get('description', '')
            
            # Create learning path content record
            content = LearningPathContent(
                path_difficulty=path_difficulty,
                module_id=module_id,
                content_type=content_type,
                content_id=int(content_id) if content_id else 1,
                title=title,
                media_type=file_type,
                media_url=file_url,
                media_description=description,
                question_number=int(question_number) if question_number else None
            )
            
            db.session.add(content)
            db.session.commit()
            
            print(f"✅ Learning path content saved: {content_type} for {module_id} (Question {question_number})")
        
        return jsonify({
            "message": "File uploaded successfully",
            "fileUrl": file_url,
            "filename": filename,
            "size": file_size,
            "type": file_type,
            "id": f"{timestamp}_{unique_id}"
        })
        
    except Exception as e:
        print(f"Error uploading file: {str(e)}")
        return jsonify({"error": f"Upload failed: {str(e)}"}), 500

@views.route('/api/admin/create-user', methods=['POST'])
@login_required
def admin_create_user():
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        data = request.get_json()
        email = data.get('email')
        firstname = data.get('firstname')
        lastname = data.get('lastname')
        contact = data.get('contact')
        password = data.get('password')
        
        # Validation
        if not email or not firstname or not lastname or not contact or not password:
            return jsonify({"success": False, "message": "All fields are required"}), 400
        
        if len(password) < 7:
            return jsonify({"success": False, "message": "Password must be at least 7 characters"}), 400
        
        if User.query.filter_by(email=email).first():
            return jsonify({"success": False, "message": "Email already exists"}), 400
        
        # Create new user
        new_user = User(
            email=email,
            firstname=firstname,
            lastname=lastname,
            contact=contact,
            is_active=True,  # Admin-created users are immediately active
            email_verified=True  # Admin-created users are pre-verified
        )
        new_user.set_password(password)
        
        db.session.add(new_user)
        db.session.commit()
        
        return jsonify({
            "success": True,
            "message": "User created successfully",
            "user_id": new_user.id
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "message": f"Error creating user: {str(e)}"}), 500

@views.route('/api/admin/create-admin', methods=['POST'])
@login_required
def admin_create_admin():
    if not current_user.is_admin():
        return jsonify({"error": "Admin access required"}), 403
    
    try:
        data = request.get_json()
        username = data.get('username')
        email = data.get('email')
        full_name = data.get('full_name')
        password = data.get('password')
        
        # Validation
        if not username or not email or not full_name or not password:
            return jsonify({"success": False, "message": "All fields are required"}), 400
        
        if len(username) < 3:
            return jsonify({"success": False, "message": "Username must be at least 3 characters"}), 400
        
        if len(password) < 7:
            return jsonify({"success": False, "message": "Password must be at least 7 characters"}), 400
        
        if Admin.query.filter_by(username=username).first():
            return jsonify({"success": False, "message": "Username already exists"}), 400
        
        if Admin.query.filter_by(email=email).first():
            return jsonify({"success": False, "message": "Email already exists"}), 400
        
        # Create new admin (all admins are regular admins, no super admin)
        new_admin = Admin(
            username=username,
            email=email,
            full_name=full_name,
            is_super_admin=False  # All created admins are regular admins
        )
        new_admin.set_password(password)
        
        db.session.add(new_admin)
        db.session.commit()
        
        return jsonify({
            "success": True,
            "message": "Admin created successfully",
            "admin_id": new_admin.id
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "message": f"Error creating admin: {str(e)}"}), 500